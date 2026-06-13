"""导出服务 - 将分析结果、报告、对话记录导出为Word文档"""
import io
import logging
from datetime import datetime
from typing import Any, Dict, List

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger(__name__)


def _add_heading(doc: Document, text: str, level: int = 1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return heading


def _add_paragraph(doc: Document, text: str, bold: bool = False, style: str = None):
    """添加段落"""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p


def _add_table(doc: Document, headers: List[str], rows: List[List[str]]):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # 数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    return table


def _add_list(doc: Document, items: List[str], ordered: bool = False):
    """添加列表"""
    for item in items:
        prefix = f"{items.index(item) + 1}. " if ordered else "• "
        _add_paragraph(doc, f"{prefix}{item}")


def _add_signal_badge(doc: Document, signal: str):
    """添加信号标识"""
    signal_map = {
        "buy": "买入",
        "sell": "卖出",
        "hold": "持有",
    }
    signal_text = signal_map.get(signal, signal)
    p = doc.add_paragraph()
    run = p.add_run(f"建议信号: {signal_text}")
    run.bold = True
    run.font.size = Pt(12)
    if signal == "buy":
        run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)
    elif signal == "sell":
        run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
    else:
        run.font.color.rgb = RGBColor(0x25, 0x6E, 0xEB)


def export_analysis_to_word(analysis_data: Dict[str, Any]) -> bytes:
    """将分析结果导出为Word文档"""
    doc = Document()

    # 文档标题
    symbol = analysis_data.get("symbol", "")
    company_name = analysis_data.get("company_name", "")
    title = f"{symbol} {company_name} 分析报告" if company_name else f"{symbol} 分析报告"
    _add_heading(doc, title, level=0)

    # 基本信息
    _add_heading(doc, "基本信息", level=1)
    current_price = analysis_data.get("current_price", "N/A")
    _add_paragraph(doc, f"股票代码: {symbol}")
    _add_paragraph(doc, f"公司名称: {company_name or 'N/A'}")
    _add_paragraph(doc, f"当前价格: ¥{current_price}")
    _add_paragraph(doc, f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # 推荐建议
    recommendation = analysis_data.get("recommendation", {})
    if recommendation:
        _add_heading(doc, "投资建议", level=1)
        _add_signal_badge(doc, recommendation.get("signal", "hold"))
        confidence = recommendation.get("confidence", 0)
        _add_paragraph(doc, f"置信度: {confidence * 100:.0f}%")
        reasoning = recommendation.get("reasoning", "")
        if reasoning:
            _add_heading(doc, "综合推理", level=2)
            _add_paragraph(doc, reasoning)

    # 各智能体分析结果
    agent_results = analysis_data.get("agent_results", {})
    if agent_results:
        _add_heading(doc, "智能体分析详情", level=1)
        for key, agent in agent_results.items():
            agent_name = agent.get("agent_name", key)
            _add_heading(doc, agent_name, level=2)

            signal = agent.get("signal", "")
            if signal:
                _add_signal_badge(doc, signal)

            analysis_text = agent.get("analysis", "")
            if analysis_text:
                _add_paragraph(doc, analysis_text)

            key_findings = agent.get("key_findings", [])
            if key_findings:
                _add_paragraph(doc, "关键发现:", bold=True)
                _add_list(doc, key_findings)

            risk_factors = agent.get("risk_factors", [])
            if risk_factors:
                _add_paragraph(doc, "风险因素:", bold=True)
                _add_list(doc, risk_factors)

    # 处理信息
    processing_time = analysis_data.get("processing_time")
    if processing_time:
        _add_heading(doc, "处理信息", level=1)
        _add_paragraph(doc, f"处理耗时: {processing_time:.1f}秒")

    # 写入字节流
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_report_to_word(report_data: Dict[str, Any]) -> bytes:
    """将报告导出为Word文档"""
    doc = Document()

    # 文档标题
    title = report_data.get("title", "金融分析报告")
    _add_heading(doc, title, level=0)

    # 报告元信息
    _add_heading(doc, "报告信息", level=1)
    report_type = report_data.get("report_type", "")
    report_type_map = {
        "morning_daily": "金融晨报",
        "stock_research": "个股研报",
        "risk_weekly": "风控周报",
        "portfolio_monthly": "组合月报",
        "event_flash": "事件快报",
    }
    _add_paragraph(doc, f"报告类型: {report_type_map.get(report_type, report_type)}")
    generated_at = report_data.get("generated_at", "")
    if generated_at:
        _add_paragraph(doc, f"生成时间: {generated_at}")
    _add_paragraph(doc, f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    symbols = report_data.get("symbols", [])
    if symbols:
        _add_paragraph(doc, f"相关股票: {', '.join(symbols)}")

    # 摘要
    summary = report_data.get("summary", "")
    if summary:
        _add_heading(doc, "摘要", level=1)
        _add_paragraph(doc, summary)

    # 正文内容
    content = report_data.get("content", "")
    if content:
        _add_heading(doc, "报告正文", level=1)
        # 按Markdown标题分段
        lines = content.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith("### "):
                _add_heading(doc, line[4:], level=3)
            elif line.startswith("## "):
                _add_heading(doc, line[3:], level=2)
            elif line.startswith("# "):
                _add_heading(doc, line[2:], level=1)
            elif line.startswith("- ") or line.startswith("• "):
                _add_paragraph(doc, f"• {line[2:]}")
            elif line.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")):
                _add_paragraph(doc, line)
            else:
                _add_paragraph(doc, line)

    # 关键发现
    key_findings = report_data.get("key_findings", [])
    if key_findings:
        _add_heading(doc, "关键发现", level=1)
        _add_list(doc, key_findings)

    # 风险因素
    risk_factors = report_data.get("risk_factors", [])
    if risk_factors:
        _add_heading(doc, "风险因素", level=1)
        _add_list(doc, risk_factors)

    # 写入字节流
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_chat_to_word(messages: List[Dict[str, Any]]) -> bytes:
    """将对话记录导出为Word文档"""
    doc = Document()

    # 文档标题
    _add_heading(doc, "FinAgent Pro 对话记录", level=0)
    _add_paragraph(doc, f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    _add_paragraph(doc, f"消息数量: {len(messages)}")

    # 对话内容
    _add_heading(doc, "对话详情", level=1)

    for msg in messages:
        role = msg.get("role", "unknown")
        content = msg.get("content", "")
        timestamp = msg.get("timestamp", "")

        role_map = {
            "user": "用户",
            "assistant": "AI助手",
            "system": "系统",
        }
        role_label = role_map.get(role, role)

        # 添加消息头
        header_text = f"[{role_label}] {timestamp}" if timestamp else f"[{role_label}]"
        p = doc.add_paragraph()
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(11)
        if role == "user":
            run.font.color.rgb = RGBColor(0x25, 0x6E, 0xEB)
        elif role == "assistant":
            run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)

        # 添加消息内容
        if content:
            # 简单处理Markdown格式
            lines = content.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if line.startswith("### "):
                    _add_heading(doc, line[4:], level=3)
                elif line.startswith("## "):
                    _add_heading(doc, line[3:], level=2)
                elif line.startswith("# "):
                    _add_heading(doc, line[2:], level=1)
                elif line.startswith("- ") or line.startswith("• "):
                    _add_paragraph(doc, f"• {line[2:]}")
                else:
                    _add_paragraph(doc, line)

        # 添加分隔
        doc.add_paragraph("─" * 40)

    # 写入字节流
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# -*- coding: utf-8 -*-
"""生成 FinAgent Pro 解决方案描述 Word 文档"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = r"D:\AFAC2026金融智能创新大赛\FinAgent-Pro\docs\FinAgentPro解决方案描述.docx"

DARK_BLUE = RGBColor(0x1A, 0x3C, 0x6E)
ACCENT_BLUE = RGBColor(0x2B, 0x5C, 0x9E)
LIGHT_BLUE_BG = "D6E4F0"
WHITE = "FFFFFF"
DARK_BG = "1A3C6E"
MID_BG = "2B5C9E"
LIGHT_BG = "E8F0FE"
GRAY_TEXT = RGBColor(0x55, 0x55, 0x55)

FONT_NAME = "微软雅黑"


def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_run(paragraph, text, bold=False, size=11, color=None, font_name=FONT_NAME):
    """添加格式化 run"""
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color
    return run


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.15):
    """设置段落间距"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = line_spacing


def add_header(doc):
    """添加页眉"""
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "FinAgent Pro 解决方案描述", bold=True, size=9, color=DARK_BLUE)
    # 页眉下方加线
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="1A3C6E"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_footer(doc):
    """添加页脚（页码）"""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 添加页码字段
    run = p.add_run()
    run.font.size = Pt(9)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)
    run2 = p.add_run()
    run2.font.size = Pt(9)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._element.append(instrText)
    run3 = p.add_run()
    run3.font.size = Pt(9)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fldChar2)


def create_cover_page(doc):
    """创建封面页"""
    # 多个空行将标题推到页面中部偏上
    for _ in range(6):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0)

    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=12, line_spacing=1.0)
    add_run(p, "FinAgent Pro", bold=True, size=36, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.0)
    add_run(p, "解决方案描述", bold=True, size=28, color=DARK_BLUE)

    # 装饰线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=12, after=12, line_spacing=1.0)
    add_run(p, "━" * 30, size=12, color=ACCENT_BLUE)

    # 副标题信息
    for text in ["AFAC2026金融智能创新大赛 初创组", "智融先锋团队", "2026年6月"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=0, after=4, line_spacing=1.5)
        add_run(p, text, size=14, color=GRAY_TEXT)

    # 分页
    doc.add_page_break()


def add_section_heading(doc, text, level=1):
    """添加章节标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        run.font.color.rgb = DARK_BLUE
        run.font.size = Pt(18 if level == 1 else 14)
    set_paragraph_spacing(heading, before=18, after=10)
    return heading


def add_body_paragraph(doc, text, bold_prefix=None, indent=False):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    set_paragraph_spacing(p, before=2, after=4, line_spacing=1.5)
    if bold_prefix:
        add_run(p, bold_prefix, bold=True, size=11)
    add_run(p, text, size=11)
    return p


def add_architecture_table(doc):
    """添加三层架构高亮表格"""
    table = doc.add_table(rows=5, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 设置表格宽度
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(16)

    layers = [
        ("第三层：主动智能", "定时巡检(APScheduler) · 事件驱动(WebSocket) · 阈值预警", DARK_BG, WHITE),
        ("第二层：协同调度", "Master Orchestrator 三阶段流水线 · 加权协商 + LLM增强推理", MID_BG, WHITE),
        ("第一层：智能体引擎", "市场分析 · 新闻舆情 · 风控合规 · 投资策略 · 报告生成 · 执行监控\nReAct推理引擎 + 三层记忆系统", "3A6FA0", WHITE),
        ("基座层", "GLM-5.1 / DeepSeek-v4-pro / SenseNova\nAKShare + Tushare + 东方财富", "5A8DB8", WHITE),
    ]

    # 标题行
    title_cell = table.rows[0].cells[0]
    set_cell_shading(title_cell, LIGHT_BLUE_BG)
    tp = title_cell.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(tp, before=6, after=6)
    add_run(tp, "FinAgent Pro 三层架构总览", bold=True, size=13, color=DARK_BLUE)

    for i, (title, desc, bg_color, fg_color) in enumerate(layers):
        cell = table.rows[i + 1].cells[0]
        set_cell_shading(cell, bg_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=8, after=2)
        add_run(p, title, bold=True, size=12, color=RGBColor.from_string(fg_color))

        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p2, before=2, after=8)
        add_run(p2, desc, size=10, color=RGBColor.from_string(fg_color))

    # 设置表格边框
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="8" w:space="0" w:color="1A3C6E"/>'
        f'<w:left w:val="single" w:sz="8" w:space="0" w:color="1A3C6E"/>'
        f'<w:bottom w:val="single" w:sz="8" w:space="0" w:color="1A3C6E"/>'
        f'<w:right w:val="single" w:sz="8" w:space="0" w:color="1A3C6E"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="1A3C6E"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="1A3C6E"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    doc.add_paragraph()  # 间距


def add_ascii_diagram(doc):
    """添加 ASCII 架构图"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=6, after=6, line_spacing=1.0)

    diagram = (
        "┌─────────────────────────────────────────────────────┐\n"
        "│           第三层：主动智能                              │\n"
        "│  定时巡检(APScheduler) │ 事件驱动(WebSocket) │ 阈值预警  │\n"
        "├─────────────────────────────────────────────────────┤\n"
        "│           第二层：协同调度                              │\n"
        "│        Master Orchestrator (三阶段流水线)              │\n"
        "│  Phase1: 并行(市场+新闻) → Phase2: 风控 → Phase3: 策略  │\n"
        "│           加权协商 + LLM增强推理                       │\n"
        "├─────────────────────────────────────────────────────┤\n"
        "│           第一层：智能体引擎                            │\n"
        "│  ┌──────┐ ┌──────┐ ┌──────┐ ┌──────┐ ┌──────┐ ┌──────┐│\n"
        "│  │市场  │ │新闻  │ │风控  │ │策略  │ │报告  │ │执行  ││\n"
        "│  │分析  │ │舆情  │ │合规  │ │投资  │ │生成  │ │监控  ││\n"
        "│  └──────┘ └──────┘ └──────┘ └──────┘ └──────┘ └──────┘│\n"
        "│           ReAct推理引擎 + 三层记忆系统                  │\n"
        "├─────────────────────────────────────────────────────┤\n"
        "│           基座层                                      │\n"
        "│  GLM-5.1 / DeepSeek-v4-pro / SenseNova              │\n"
        "│  AKShare + Tushare + 东方财富                         │\n"
        "└─────────────────────────────────────────────────────┘"
    )

    run = add_run(p, diagram, size=8, color=DARK_BLUE, font_name="Consolas")
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    # 添加浅蓝底色
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{LIGHT_BLUE_BG}"/>')
    p._element.get_or_add_pPr().append(shd)


def add_innovation_table(doc):
    """添加核心创新点表格"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=6)

    innovations = [
        ("1", "ReAct推理引擎", "思考→行动→观察循环，让AI像分析师一样推理"),
        ("2", "持久记忆系统", "工作记忆+情节记忆+语义记忆，记住上下文"),
        ("3", "主动智能", "定时巡检+事件驱动+阈值预警，从被动到主动"),
        ("4", "合规内嵌", "监管规则引擎+审计日志，每步可追溯"),
        ("5", "多智能体协商", "加权协商+LLM增强推理，综合多方意见"),
        ("6", "思维链可视化", "AI决策过程实时可见，黑盒变白盒"),
    ]

    table = doc.add_table(rows=len(innovations) + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置列宽
    for row in table.rows:
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(3.5)
        row.cells[2].width = Cm(11)

    # 表头
    headers = ["序号", "创新点", "说明"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, DARK_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=4, after=4)
        add_run(p, h, bold=True, size=10, color=RGBColor.from_string(WHITE))

    # 数据行
    for row_idx, (num, name, desc) in enumerate(innovations):
        bg = LIGHT_BG if row_idx % 2 == 0 else WHITE
        for col_idx, text in enumerate([num, name, desc]):
            cell = table.rows[row_idx + 1].cells[col_idx]
            set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx < 2 else WD_ALIGN_PARAGRAPH.LEFT
            p.alignment = align
            set_paragraph_spacing(p, before=3, after=3)
            is_bold = col_idx == 1
            add_run(p, text, bold=is_bold, size=10, color=DARK_BLUE if col_idx == 1 else RGBColor(0x33, 0x33, 0x33))

    # 表格边框
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="6" w:space="0" w:color="1A3C6E"/>'
        f'<w:left w:val="single" w:sz="6" w:space="0" w:color="1A3C6E"/>'
        f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="1A3C6E"/>'
        f'<w:right w:val="single" w:sz="6" w:space="0" w:color="1A3C6E"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="1A3C6E"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="1A3C6E"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def main():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 添加页眉页脚
    add_header(doc)
    add_footer(doc)

    # 封面
    create_cover_page(doc)

    # ========== 概述 ==========
    add_section_heading(doc, "一、方案概述")
    add_body_paragraph(doc,
        "FinAgent Pro构建了6大专业智能体协同工作的金融数字员工平台，通过三层架构解决金融行业痛点。"
        "平台以国产大模型为基座，融合ReAct推理引擎、持久记忆系统与主动智能机制，"
        "为金融机构提供从市场分析到投资决策的全链路智能支持。")
    add_body_paragraph(doc,
        "传统金融分析面临三大痛点：一是信息孤岛，市场数据、新闻舆情、风控合规分散在不同系统；"
        "二是被动响应，分析师只能被动等待信息更新，无法主动发现风险；"
        "三是合规黑盒，AI决策过程不透明，难以满足监管审计要求。"
        "FinAgent Pro通过三层架构设计，系统性解决以上问题。")

    # ========== 三层架构 ==========
    add_section_heading(doc, "二、三层架构设计")

    # 架构总览表格
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=8)
    add_run(p, "▎ 架构总览", bold=True, size=13, color=DARK_BLUE)
    add_architecture_table(doc)

    # 第一层
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4)
    add_run(p, "▎ 第一层：智能体引擎", bold=True, size=13, color=DARK_BLUE)

    add_body_paragraph(doc,
        '每个智能体内置ReAct推理引擎，执行「思考→行动→观察」循环。'
        '不同于传统的单次调用模式，ReAct引擎让智能体能够像人类分析师一样，'
        '先思考需要什么信息，再采取行动获取，最后观察结果并决定下一步。', indent=True)

    # 智能体详情表格
    agent_table = doc.add_table(rows=7, cols=3)
    agent_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in agent_table.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(5)
        row.cells[2].width = Cm(8.5)

    agent_headers = ["智能体", "核心能力", "技术指标"]
    for i, h in enumerate(agent_headers):
        cell = agent_table.rows[0].cells[i]
        set_cell_shading(cell, DARK_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=4, after=4)
        add_run(p, h, bold=True, size=10, color=RGBColor.from_string(WHITE))

    agents = [
        ("市场分析智能体", "技术指标与估值分析", "RSI/MACD/KDJ技术指标，PE/PB/ROE估值指标"),
        ("新闻舆情智能体", "中文新闻抓取与情感分析", "NLP情感分析，多源新闻聚合"),
        ("风控合规智能体", "四维风险评估与合规检查", "波动率/集中度/流动性四维风险，合规规则引擎"),
        ("投资策略智能体", "仓位建议与目标价生成", "多因子模型，风险收益优化"),
        ("报告生成智能体", "自动输出专业报告", "晨报/研报/风控周报模板化生成"),
        ("执行监控智能体", "定时巡检与主动预警", "APScheduler调度，WebSocket推送"),
    ]
    for row_idx, (name, ability, detail) in enumerate(agents):
        bg = LIGHT_BG if row_idx % 2 == 0 else WHITE
        for col_idx, text in enumerate([name, ability, detail]):
            cell = agent_table.rows[row_idx + 1].cells[col_idx]
            set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, before=3, after=3)
            is_bold = col_idx == 0
            add_run(p, text, bold=is_bold, size=10, color=DARK_BLUE if col_idx == 0 else RGBColor(0x33, 0x33, 0x33))

    # 表格边框
    tbl = agent_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="6" w:space="0" w:color="1A3C6E"/>'
        f'<w:left w:val="single" w:sz="6" w:space="0" w:color="1A3C6E"/>'
        f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="1A3C6E"/>'
        f'<w:right w:val="single" w:sz="6" w:space="0" w:color="1A3C6E"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="1A3C6E"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="1A3C6E"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    doc.add_paragraph()

    # 第二层
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4)
    add_run(p, "▎ 第二层：协同调度", bold=True, size=13, color=DARK_BLUE)

    add_body_paragraph(doc,
        "Master Orchestrator采用三阶段流水线调度，确保智能体间的高效协同：", indent=True)

    phases = [
        ("Phase 1（并行执行）：", "市场分析智能体和新闻舆情智能体同时启动，"
         "分别计算技术指标/估值指标和抓取新闻/情感分析，互不依赖，最大化并行效率。"),
        ("Phase 2（串行执行）：", "风控合规智能体依赖Phase 1的结果，"
         "在获取市场数据和舆情信息后，评估波动率/集中度/流动性四维风险，并执行合规规则引擎检查。"),
        ("Phase 3（策略生成）：", "投资策略智能体综合前两个阶段的分析结果，"
         "生成仓位建议和目标价，报告生成智能体自动输出晨报/研报/风控周报。"),
    ]
    for prefix, text in phases:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        set_paragraph_spacing(p, before=2, after=2, line_spacing=1.5)
        add_run(p, prefix, bold=True, size=11, color=ACCENT_BLUE)
        add_run(p, text, size=11)

    add_body_paragraph(doc,
        "6大智能体通过加权协商算法（Σsignal×confidence / Σconfidence）综合出最终建议，"
        "并由LLM增强推理生成专业分析文本，确保结论既有数据支撑又有专业解读。", indent=True)

    # 第三层
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4)
    add_run(p, "▎ 第三层：主动智能", bold=True, size=13, color=DARK_BLUE)

    add_body_paragraph(doc,
        "传统金融系统是被动响应式的——用户问什么，系统答什么。"
        "FinAgent Pro的第三层主动智能，让数字员工从被动变为主动：", indent=True)

    pro_features = [
        ("定时巡检（APScheduler）：", "每日8:30自动生成晨报，17:00自动生成收盘总结，"
         "无需人工触发，数字员工主动完成日常分析任务。"),
        ("事件驱动（WebSocket）：", "持仓股突发利空消息时，系统自动触发分析流程，"
         "从新闻抓取到风险评估到策略调整，全链路自动执行并实时推送结果。"),
        ("阈值预警：", "当技术指标越界（如RSI超买超卖）、风险指标超限时，"
         "系统主动发现并推送预警，不等用户查询。"),
    ]
    for prefix, text in pro_features:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        set_paragraph_spacing(p, before=2, after=2, line_spacing=1.5)
        add_run(p, prefix, bold=True, size=11, color=ACCENT_BLUE)
        add_run(p, text, size=11)

    add_body_paragraph(doc,
        "三层记忆系统（工作记忆+情节记忆+语义记忆）让数字员工记住上下文，"
        "工作记忆保存当前会话信息，情节记忆记录历史交互事件，语义记忆存储长期知识。"
        "合规规则引擎确保每步操作可审计可追溯，满足金融监管要求。", indent=True)

    add_body_paragraph(doc,
        "以国产大模型GLM-5.1/DeepSeek/SenseNova为基座，数据不出境，安全可信。"
        "通过AKShare、Tushare、东方财富等多源数据接入，确保数据的全面性和实时性。", indent=True)

    # ========== 架构图 ==========
    add_section_heading(doc, "三、系统架构图")
    add_ascii_diagram(doc)

    # ========== 核心创新点 ==========
    add_section_heading(doc, "四、核心创新点")
    add_body_paragraph(doc,
        "FinAgent Pro的核心创新体现在以下六个方面：")
    add_innovation_table(doc)

    # 保存
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"文档已生成: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()

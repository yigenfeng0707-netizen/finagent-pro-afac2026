#!/usr/bin/env python3
"""生成 FinAgent Pro 用户使用手册 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# 辅助函数
# ============================================================

def set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_with_header(doc, headers, rows, col_widths=None):
    """添加带表头样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2B579A")

    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "E8EEF7")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def add_code_block(doc, code_text):
    """添加代码块（Courier New + 灰色底色）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    # 添加底色
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F2F2F2"/>')
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p

def add_bold_text(paragraph, text):
    """在段落中添加加粗文本"""
    run = paragraph.add_run(text)
    run.bold = True
    return run

def add_normal_text(paragraph, text):
    """在段落中添加普通文本"""
    run = paragraph.add_run(text)
    return run

def add_bullet(doc, text, bold_prefix=None, level=0):
    """添加项目符号列表"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    if bold_prefix:
        add_bold_text(p, bold_prefix)
        add_normal_text(p, text)
    else:
        add_normal_text(p, text)
    return p


# ============================================================
# 创建文档
# ============================================================

doc = Document()

# ---- 全局默认字体 ----
style = doc.styles['Normal']
font = style.font
font.name = 'SimSun'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

# ---- 页面设置 ----
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

# ---- 页眉 ----
header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
hr = hp.add_run("FinAgent Pro 用户使用手册")
hr.font.size = Pt(9)
hr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
hr.font.name = 'SimSun'
hr.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

# ---- 页脚（页码） ----
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
# 插入页码字段
run = fp.add_run()
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run._r.append(fldChar1)
run2 = fp.add_run()
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
run2._r.append(instrText)
run3 = fp.add_run()
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run3._r.append(fldChar2)

# ---- 自定义标题样式 ----
for level in range(1, 4):
    style_name = f'Heading {level}'
    hs = doc.styles[style_name]
    hs.font.name = 'SimHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    if level == 1:
        hs.font.size = Pt(18)
        hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(14)
        hs.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    elif level == 3:
        hs.font.size = Pt(12)
        hs.font.color.rgb = RGBColor(0x3A, 0x6E, 0xA5)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)


# ============================================================
# 封面页
# ============================================================

# 空行撑开
for _ in range(6):
    doc.add_paragraph('')

# 主标题
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FinAgent Pro')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)
run.bold = True
run.font.name = 'SimHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('金融自主智能体平台')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
run.font.name = 'SimHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

doc.add_paragraph('')

# 分隔线
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━' * 30)
run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
run.font.size = Pt(14)

doc.add_paragraph('')

# 副标题
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('用户使用手册 V1.0')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x3A, 0x6E, 0xA5)
run.font.name = 'SimHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

for _ in range(3):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AFAC2026金融智能创新大赛参赛项目')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run.font.name = 'SimSun'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026年6月')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run.font.name = 'SimSun'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

# 分页
doc.add_page_break()


# ============================================================
# 目录页
# ============================================================

doc.add_heading('目  录', level=1)

toc_items = [
    ("第1章  产品概述", 1),
    ("  1.1  产品简介", 2),
    ("  1.2  核心特性", 2),
    ("  1.3  适用场景", 2),
    ("第2章  系统要求与安装", 1),
    ("  2.1  环境要求", 2),
    ("  2.2  后端安装", 2),
    ("  2.3  前端安装", 2),
    ("  2.4  Docker部署", 2),
    ("  2.5  环境变量说明", 2),
    ("第3章  功能使用指南", 1),
    ("  3.1  工作台", 2),
    ("  3.2  股票分析", 2),
    ("  3.3  数字员工对话", 2),
    ("  3.4  智能体管理", 2),
    ("  3.5  报告中心", 2),
    ("  3.6  预警中心", 2),
    ("第4章  API接口文档", 1),
    ("  4.1  认证方式", 2),
    ("  4.2  股票分析接口", 2),
    ("  4.3  对话接口", 2),
    ("  4.4  市场数据接口", 2),
    ("  4.5  报告接口", 2),
    ("  4.6  预警接口", 2),
    ("  4.7  定时任务接口", 2),
    ("  4.8  WebSocket接口", 2),
    ("第5章  智能体详解", 1),
    ("  5.1  市场分析智能体", 2),
    ("  5.2  新闻舆情智能体", 2),
    ("  5.3  风控合规智能体", 2),
    ("  5.4  投资策略智能体", 2),
    ("  5.5  报告生成智能体", 2),
    ("  5.6  执行监控智能体", 2),
    ("第6章  合规与安全", 1),
    ("  6.1  合规规则", 2),
    ("  6.2  审计日志", 2),
    ("  6.3  数据安全", 2),
    ("第7章  常见问题（FAQ）", 1),
    ("附录A  信号说明", 1),
    ("附录B  风险等级说明", 1),
]

for item, level in toc_items:
    p = doc.add_paragraph()
    if level == 1:
        run = p.add_run(item)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'SimHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    else:
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.name = 'SimSun'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()


# ============================================================
# 第1章 产品概述
# ============================================================

doc.add_heading('第1章  产品概述', level=1)

doc.add_heading('1.1  产品简介', level=2)

p = doc.add_paragraph()
add_normal_text(p, 'FinAgent Pro 是一款面向金融行业的')
add_bold_text(p, '自主智能体平台')
add_normal_text(p, '，基于大语言模型（LLM）与多智能体协作架构，为专业投资者提供全方位的智能投研服务。平台内置')
add_bold_text(p, '6大数字员工')
add_normal_text(p, '，分别覆盖市场分析、新闻舆情、风控合规、投资策略、报告生成与执行监控等核心业务场景。')

p = doc.add_paragraph()
add_normal_text(p, 'FinAgent Pro 采用 ReAct（Reasoning + Acting）推理框架，使每个智能体具备"思考-行动-观察"的闭环决策能力。同时，通过持久记忆机制，智能体可以积累历史分析经验，持续优化决策质量。')

p = doc.add_paragraph()
add_normal_text(p, '平台核心定位：让金融分析从"人工驱动"升级为"AI驱动"，实现从数据采集、分析推理、合规审查到报告输出的全链路智能化。')

doc.add_heading('1.2  核心特性', level=2)

features = [
    ("ReAct推理引擎：", "每个智能体基于 ReAct（Reasoning + Acting）范式，实现\u201c思考\u2192行动\u2192观察\u201d的闭环推理，确保每一步决策有据可依、可追溯。"),
    ("持久记忆机制：", "智能体具备长期记忆能力，可积累历史分析经验与用户偏好，持续优化决策质量，避免重复分析。"),
    ("主动智能推送：", "执行监控智能体定时巡检市场异动，主动触发预警和报告生成，从\u201c被动响应\u201d升级为\u201c主动服务\u201d。"),
    ("合规规则内嵌：", "内置单股集中度10%、行业集中度30%、创业板限制20%、ST股禁止等合规规则，所有投资建议均经过合规审查。"),
    ("多智能体协商：", "6大专业智能体并行分析、交叉验证，通过加权信号协商机制生成最终投资建议，避免单一视角偏差。"),
    ("思维链可视化：", "完整展示每个智能体的推理步骤（Thought\u2192Action\u2192Observation），让AI决策过程透明可解释。"),
]

for bold_part, normal_part in features:
    add_bullet(doc, normal_part, bold_prefix=bold_part)

doc.add_heading('1.3  适用场景', level=2)

scenarios = [
    ("券商研究所：", "自动化研报生成、个股深度分析、行业热点追踪，提升分析师工作效率3-5倍。"),
    ("基金公司：", "组合风险监控、合规审查自动化、投资建议生成，降低合规风险与运营成本。"),
    ("银行理财：", "理财产品风险评估、市场异动预警、客户画像分析，增强风控能力。"),
    ("个人投资者：", "智能选股辅助、实时预警提醒、个性化投资建议，降低信息不对称。"),
]

for bold_part, normal_part in scenarios:
    add_bullet(doc, normal_part, bold_prefix=bold_part)

doc.add_page_break()


# ============================================================
# 第2章 系统要求与安装
# ============================================================

doc.add_heading('第2章  系统要求与安装', level=1)

doc.add_heading('2.1  环境要求', level=2)

add_table_with_header(doc,
    ["组件", "最低版本", "推荐版本", "说明"],
    [
        ["Python", "3.12", "3.12+", "后端运行环境"],
        ["Node.js", "18.0", "20.x LTS", "前端构建环境"],
        ["Redis", "7.0", "7-alpine", "可选，缓存与消息队列"],
        ["Docker", "24.0", "最新版", "可选，容器化部署"],
        ["Docker Compose", "2.20", "最新版", "可选，编排容器服务"],
    ],
    col_widths=[3, 2.5, 2.5, 5]
)

doc.add_heading('2.2  后端安装', level=2)

p = doc.add_paragraph()
add_normal_text(p, '1. 克隆项目并进入后端目录：')
add_code_block(doc, 'git clone <repository_url>\ncd FinAgent-Pro/backend')

p = doc.add_paragraph()
add_normal_text(p, '2. 安装Python依赖：')
add_code_block(doc, 'pip install -r requirements.txt')

p = doc.add_paragraph()
add_normal_text(p, '3. 配置环境变量（复制示例文件并修改）：')
add_code_block(doc, 'cp .env.example .env\n# 编辑 .env 文件，填入API密钥等配置')

p = doc.add_paragraph()
add_normal_text(p, '4. 启动后端服务：')
add_code_block(doc, 'uvicorn main:app --host 0.0.0.0 --port 8000 --reload')

doc.add_heading('2.3  前端安装', level=2)

p = doc.add_paragraph()
add_normal_text(p, '1. 进入前端目录并安装依赖：')
add_code_block(doc, 'cd FinAgent-Pro/frontend\nnpm install')

p = doc.add_paragraph()
add_normal_text(p, '2. 启动开发服务器：')
add_code_block(doc, 'npm run dev')

p = doc.add_paragraph()
add_normal_text(p, '3. 访问 http://localhost:5173 即可使用前端界面。')

doc.add_heading('2.4  Docker部署', level=2)

p = doc.add_paragraph()
add_normal_text(p, '使用 Docker Compose 一键部署（包含后端API + Redis）：')
add_code_block(doc, 'docker-compose up -d')

p = doc.add_paragraph()
add_normal_text(p, '服务启动后，API地址为 http://localhost:8000，前端需单独启动或配置反向代理。')

doc.add_heading('2.5  环境变量说明', level=2)

add_table_with_header(doc,
    ["变量名", "默认值", "说明"],
    [
        ["LLM_PROVIDER", "glm", "主LLM提供商：glm / deepseek / sensenova"],
        ["LLM_ENABLED", "true", "是否启用LLM增强推理"],
        ["GLM_API_KEY", "-", "智谱AI API密钥（GLM-5.1主力模型）"],
        ["GLM_BASE_URL", "https://api.edgefn.net/v1", "智谱AI API地址"],
        ["GLM_MODEL", "GLM-5.1", "智谱AI模型名称"],
        ["DEEPSEEK_API_KEY", "-", "DeepSeek API密钥（深度推理备选）"],
        ["DEEPSEEK_BASE_URL", "https://api.deepseek.com", "DeepSeek API地址"],
        ["DEEPSEEK_MODEL", "deepseek-v4-pro", "DeepSeek模型名称"],
        ["SENSENOVA_API_KEY", "-", "商汤日日新API密钥（轻量快速模型）"],
        ["SENSENOVA_BASE_URL", "https://token.sensenova.cn/v1", "商汤日日新API地址"],
        ["SENSENOVA_MODEL", "sensenova-6.7-flash-lite", "商汤日日新模型名称"],
        ["TUSHARE_TOKEN", "-", "Tushare数据源Token"],
        ["NEWSAPI_KEY", "-", "新闻API密钥"],
        ["DATABASE_URL", "sqlite:///./finagent.db", "数据库连接字符串"],
        ["REDIS_URL", "redis://localhost:6379", "Redis连接地址"],
        ["DEBUG", "true", "调试模式"],
        ["SECRET_KEY", "-", "应用密钥（生产环境务必修改）"],
        ["API_ACCESS_KEY", "-", "API访问密钥"],
        ["API_ACCESS_ENABLED", "false", "是否启用API访问控制"],
        ["COMPLIANCE_ENABLED", "true", "是否启用合规检查"],
        ["MAX_SINGLE_STOCK_RATIO", "0.10", "单股集中度上限"],
        ["MAX_SECTOR_RATIO", "0.30", "行业集中度上限"],
        ["MAX_GEM_RATIO", "0.20", "创业板持仓上限"],
        ["SCHEDULER_ENABLED", "true", "是否启用定时任务"],
        ["MORNING_REPORT_TIME", "08:30", "晨报生成时间"],
        ["EVENING_REPORT_TIME", "17:00", "晚报生成时间"],
    ],
    col_widths=[4.5, 4.5, 5]
)

doc.add_page_break()


# ============================================================
# 第3章 功能使用指南
# ============================================================

doc.add_heading('第3章  功能使用指南', level=1)

doc.add_heading('3.1  工作台', level=2)

p = doc.add_paragraph()
add_normal_text(p, '工作台是FinAgent Pro的主界面，提供全局视角的金融数据概览与快速操作入口。主要包含以下模块：')

modules = [
    ("市场概览卡片：", "实时展示上证指数、深证成指、创业板指等主要指数的行情数据，包括涨跌幅、成交量等关键指标。"),
    ("快速分析入口：", "提供股票代码输入框，支持一键发起综合分析，快速获取投资建议。"),
    ("智能体状态面板：", "实时显示6大数字员工的运行状态（空闲/分析中/异常），方便监控系统健康度。"),
    ("预警摘要：", "汇总最近的高优先级预警信息，包括价格异动、风险预警、合规提醒等，支持点击查看详情。"),
]

for bold_part, normal_part in modules:
    add_bullet(doc, normal_part, bold_prefix=bold_part)

doc.add_heading('3.2  股票分析', level=2)

p = doc.add_paragraph()
add_bold_text(p, '操作步骤：')

p = doc.add_paragraph()
add_normal_text(p, '1. 在分析页面输入股票代码（如 600519），或通过对话界面使用自然语言提问（如"分析一下贵州茅台"）。')

p = doc.add_paragraph()
add_normal_text(p, '2. 选择分析类型：')

add_table_with_header(doc,
    ["分析类型", "说明", "调用智能体"],
    [
        ["综合分析（comprehensive）", "全维度深度分析", "市场+新闻+风控+策略"],
        ["快速分析（quick）", "仅市场技术面快速扫描", "市场"],
        ["技术分析（technical）", "侧重技术指标分析", "市场"],
        ["基本面分析（fundamental）", "侧重财务指标分析", "市场"],
    ],
    col_widths=[4, 5, 4]
)

p = doc.add_paragraph()
add_normal_text(p, '3. 查看分析结果：系统将展示6大智能体的分析结果，包括信号（买入/卖出/持有）、置信度（0-1）、关键发现与风险因素。')

p = doc.add_paragraph()
add_normal_text(p, '4. 解读信号与置信度：')

add_table_with_header(doc,
    ["信号", "含义", "置信度参考"],
    [
        ["strong_buy / 强烈买入", "多维度强烈看多", "≥ 0.8"],
        ["buy / 买入", "偏多，可考虑建仓", "0.5 - 0.8"],
        ["hold / 持有", "中性，建议观望", "0.3 - 0.5"],
        ["sell / 卖出", "偏空，建议减仓", "0.2 - 0.5"],
        ["strong_sell / 强烈卖出", "多维度强烈看空", "≥ 0.8"],
    ],
    col_widths=[4, 4.5, 3.5]
)

doc.add_heading('3.3  数字员工对话', level=2)

p = doc.add_paragraph()
add_normal_text(p, 'FinAgent Pro 提供')
add_bold_text(p, '对话式交互')
add_normal_text(p, '界面，用户可以使用自然语言与数字员工进行交流。')

p = doc.add_paragraph()
add_bold_text(p, '核心能力：')

dialog_features = [
    ("自然语言提问：", "支持中文自然语言输入，如\u201c贵州茅台最近怎么样？\u201d\u201c帮我看看600519的风险\u201d。"),
    ("智能意图识别：", "自动识别用户意图（分析股票/市场概览/风险检查/通用对话），路由到对应智能体。"),
    ("思维链实时展示：", "对话过程中实时展示智能体的推理步骤（Thought\u2192Action\u2192Observation），让AI思考过程透明可见。"),
    ("股票代码自动提取：", "支持6位数字代码和常见股票名称（如\u201c茅台\u201d\u2192600519，\u201c宁德时代\u201d\u2192300750）。"),
]

for bold_part, normal_part in dialog_features:
    add_bullet(doc, normal_part, bold_prefix=bold_part)

p = doc.add_paragraph()
add_bold_text(p, '示例对话：')

add_code_block(doc, '用户：分析一下贵州茅台\n\nAI：📊 600519 分析完成！\n建议: buy（买入）\n置信度: 72%\n理由: 基于市场分析、新闻舆情、风控合规3个专业智能体的综合分析...\n风险等级: medium（中等风险）')

doc.add_heading('3.4  智能体管理', level=2)

p = doc.add_paragraph()
add_normal_text(p, '智能体管理页面提供6大数字员工的运行状态监控与管理功能：')

add_table_with_header(doc,
    ["智能体", "类型标识", "核心功能", "状态监控"],
    [
        ["市场分析智能体", "market", "技术指标+基本面+资金流向", "✓"],
        ["新闻舆情智能体", "news", "新闻采集+情绪分析+事件提取", "✓"],
        ["风控合规智能体", "risk", "4维风控+合规规则引擎", "✓"],
        ["投资策略智能体", "strategy", "仓位建议+目标价+组合分析", "✓"],
        ["报告生成智能体", "report", "5种报告模板自动生成", "✓"],
        ["执行监控智能体", "execution", "定时巡检+主动预警", "✓"],
    ],
    col_widths=[3.5, 2.5, 5, 2]
)

p = doc.add_paragraph()
add_normal_text(p, '每个智能体可查看：执行统计（调用次数/成功率/平均耗时）、工具列表、最近执行记录等。')

doc.add_heading('3.5  报告中心', level=2)

p = doc.add_paragraph()
add_normal_text(p, '报告中心支持')
add_bold_text(p, '5种报告类型')
add_normal_text(p, '的一键生成：')

add_table_with_header(doc,
    ["报告类型", "标识", "说明", "生成方式"],
    [
        ["晨报", "morning_daily", "每日市场概览与重要资讯汇总", "定时自动 / 手动"],
        ["研报", "stock_research", "个股深度研究报告", "手动触发"],
        ["风控周报", "risk_weekly", "周度风险回顾与合规统计", "定时自动 / 手动"],
        ["组合月报", "portfolio_monthly", "月度投资组合表现分析", "定时自动 / 手动"],
        ["事件快报", "event_flash", "重大事件即时分析报告", "事件触发 / 手动"],
    ],
    col_widths=[2.5, 3, 5, 3]
)

doc.add_heading('3.6  预警中心', level=2)

p = doc.add_paragraph()
add_normal_text(p, '预警中心提供')
add_bold_text(p, '分级预警')
add_normal_text(p, '机制，确保用户及时获取关键信息：')

add_table_with_header(doc,
    ["预警等级", "标识", "颜色", "说明"],
    [
        ["低", "low", "绿色", "信息性提醒，无需立即行动"],
        ["中", "medium", "黄色", "需关注，建议择机处理"],
        ["高", "high", "橙色", "重要预警，建议尽快处理"],
        ["极高", "critical", "红色", "紧急预警，需立即处理"],
    ],
    col_widths=[2, 2.5, 2, 7]
)

p = doc.add_paragraph()
add_bold_text(p, '预警类型：')

alert_types = [
    ("价格预警（price）：", "股价突破关键价位、涨跌幅异常等。"),
    ("风险预警（risk）：", "波动率飙升、集中度超标、流动性不足等。"),
    ("新闻预警（news）：", "重大负面新闻、舆情恶化等。"),
    ("合规预警（compliance）：", "持仓违反合规规则、触及监管红线等。"),
]

for bold_part, normal_part in alert_types:
    add_bullet(doc, normal_part, bold_prefix=bold_part)

p = doc.add_paragraph()
add_normal_text(p, '预警支持实时推送（WebSocket）和页面展示，每条预警包含：预警类型、等级、标题、详细描述、关联股票、建议操作。')

doc.add_page_break()


# ============================================================
# 第4章 API接口文档
# ============================================================

doc.add_heading('第4章  API接口文档', level=1)

p = doc.add_paragraph()
add_normal_text(p, 'FinAgent Pro 后端基于 FastAPI 构建，所有接口基础路径为 ')
add_code_block(doc, 'http://localhost:8000/api')

doc.add_heading('4.1  认证方式', level=2)

p = doc.add_paragraph()
add_normal_text(p, '当启用API访问控制时（API_ACCESS_ENABLED=true），需在请求头中携带访问密钥：')
add_code_block(doc, 'Authorization: Bearer <API_ACCESS_KEY>')

p = doc.add_paragraph()
add_normal_text(p, '未启用访问控制时（默认），所有接口可直接访问。')

doc.add_heading('4.2  股票分析接口', level=2)

p = doc.add_paragraph()
add_bold_text(p, 'POST /api/analyze')
add_normal_text(p, '  ——  对指定股票进行多智能体综合分析')

p = doc.add_paragraph()
add_bold_text(p, '请求参数：')

add_table_with_header(doc,
    ["参数", "类型", "必填", "默认值", "说明"],
    [
        ["symbol", "string", "是", "-", "股票代码，如 600519"],
        ["analysis_type", "string", "否", "comprehensive", "分析类型：comprehensive/quick/technical/fundamental"],
        ["include_news", "boolean", "否", "true", "是否包含新闻舆情分析"],
        ["include_risk", "boolean", "否", "true", "是否包含风控合规分析"],
        ["include_strategy", "boolean", "否", "false", "是否包含投资策略建议"],
    ],
    col_widths=[3, 2, 1.5, 3, 5]
)

p = doc.add_paragraph()
add_bold_text(p, '请求示例：')
add_code_block(doc, 'POST /api/analyze\nContent-Type: application/json\n\n{\n  "symbol": "600519",\n  "analysis_type": "comprehensive",\n  "include_news": true,\n  "include_risk": true,\n  "include_strategy": true\n}')

p = doc.add_paragraph()
add_bold_text(p, '响应示例：')
add_code_block(doc, '{\n  "request_id": "a1b2c3d4",\n  "symbol": "600519",\n  "company_name": "贵州茅台",\n  "current_price": 1688.50,\n  "change_percent": 1.23,\n  "recommendation": {\n    "signal": "buy",\n    "confidence": 0.72,\n    "target_price": 1780.00,\n    "stop_loss": 1620.00,\n    "reasoning": "基于3个专业智能体的综合分析...",\n    "key_points": ["技术面金叉信号", "基本面PE合理", "资金持续流入"],\n    "risk_assessment": {\n      "overall_risk": "medium",\n      "risk_score": 0.45\n    }\n  },\n  "agent_results": { ... },\n  "processing_time": 3.52\n}')

doc.add_heading('4.3  对话接口', level=2)

p = doc.add_paragraph()
add_bold_text(p, 'POST /api/chat')
add_normal_text(p, '  ——  与金融数字员工进行对话交互')

p = doc.add_paragraph()
add_bold_text(p, '请求参数：')

add_table_with_header(doc,
    ["参数", "类型", "必填", "说明"],
    [
        ["message", "string", "是", "用户消息内容"],
        ["conversation_id", "string", "否", "会话ID，用于多轮对话"],
        ["context", "object", "否", "上下文信息"],
    ],
    col_widths=[3.5, 2, 1.5, 7]
)

p = doc.add_paragraph()
add_bold_text(p, '请求示例：')
add_code_block(doc, '{\n  "message": "分析一下贵州茅台",\n  "conversation_id": "conv_001"\n}')

p = doc.add_paragraph()
add_bold_text(p, '响应示例：')
add_code_block(doc, '{\n  "response": "📊 600519 分析完成！建议: buy，置信度: 72%...",\n  "agent_steps": [\n    {"thought": "用户想分析股票", "action": "extract_symbol", "observation": "提取到代码600519"},\n    {"thought": "需要综合分析", "action": "analyze_stock", "observation": "分析完成"}\n  ],\n  "related_stocks": ["600519"],\n  "suggestions": ["查看详细分析报告", "设置价格预警"],\n  "confidence": 0.72\n}')

doc.add_heading('4.4  市场数据接口', level=2)

p = doc.add_paragraph()
add_bold_text(p, 'GET /api/market/overview')
add_normal_text(p, '  ——  获取A股市场概览数据')

p = doc.add_paragraph()
add_bold_text(p, 'GET /api/stock/{symbol}')
add_normal_text(p, '  ——  获取单只股票的实时数据')

p = doc.add_paragraph()
add_bold_text(p, 'GET /api/stock/{symbol}/chart?period=6mo')
add_normal_text(p, '  ——  获取K线图数据')

p = doc.add_paragraph()
add_normal_text(p, 'period 参数支持：1mo / 3mo / 6mo / 1y / 2y')

doc.add_heading('4.5  报告接口', level=2)

p = doc.add_paragraph()
add_bold_text(p, 'GET /api/reports?report_type=&limit=20')
add_normal_text(p, '  ——  获取报告列表')

p = doc.add_paragraph()
add_bold_text(p, 'POST /api/reports/generate?report_type=morning_daily&symbol=600519')
add_normal_text(p, '  ——  生成指定类型的报告')

p = doc.add_paragraph()
add_normal_text(p, 'report_type 可选值：morning_daily / stock_research / risk_weekly / portfolio_monthly / event_flash')

doc.add_heading('4.6  预警接口', level=2)

p = doc.add_paragraph()
add_bold_text(p, 'GET /api/alerts?limit=20')
add_normal_text(p, '  ——  获取最近的预警信息列表')

p = doc.add_paragraph()
add_normal_text(p, '返回数据包含：alert_id、alert_type（price/risk/news/compliance）、severity（low/medium/high/critical）、title、message、symbol、timestamp、action_suggested。')

doc.add_heading('4.7  定时任务接口', level=2)

add_table_with_header(doc,
    ["方法", "路径", "说明"],
    [
        ["GET", "/api/tasks", "获取所有定时任务列表"],
        ["POST", "/api/tasks", "创建新的定时任务"],
        ["DELETE", "/api/tasks/{task_id}", "删除指定定时任务"],
    ],
    col_widths=[2.5, 4.5, 7]
)

p = doc.add_paragraph()
add_bold_text(p, '创建任务请求参数：')

add_table_with_header(doc,
    ["参数", "类型", "必填", "说明"],
    [
        ["task_type", "string", "是", "任务类型：morning_report / evening_report / risk_scan"],
        ["cron_expression", "string", "是", "Cron表达式，如 '0 8 * * 1-5'"],
        ["params", "object", "否", "任务参数"],
        ["enabled", "boolean", "否", "是否启用，默认 true"],
    ],
    col_widths=[3.5, 2, 1.5, 7]
)

doc.add_heading('4.8  WebSocket接口', level=2)

p = doc.add_paragraph()
add_bold_text(p, 'WS /api/ws')
add_normal_text(p, '  ——  实时推送通道')

p = doc.add_paragraph()
add_normal_text(p, '连接后可接收以下实时消息：')

ws_events = [
    ("智能体进度推送：", "分析过程中各智能体的执行状态更新。"),
    ("预警通知：", "实时推送新产生的预警信息。"),
    ("分析完成通知：", "股票分析完成后的结果推送。"),
]

for bold_part, normal_part in ws_events:
    add_bullet(doc, normal_part, bold_prefix=bold_part)

p = doc.add_paragraph()
add_bold_text(p, '连接示例：')
add_code_block(doc, 'const ws = new WebSocket("ws://localhost:8000/api/ws");\nws.onmessage = (event) => {\n  const data = JSON.parse(event.data);\n  console.log("收到推送:", data);\n};')

doc.add_page_break()


# ============================================================
# 第5章 智能体详解
# ============================================================

doc.add_heading('第5章  智能体详解', level=1)

doc.add_heading('5.1  市场分析智能体', level=2)

p = doc.add_paragraph()
add_normal_text(p, '市场分析智能体（Market Analysis Agent）是FinAgent Pro的核心分析引擎，负责从技术面、基本面和资金面三个维度对股票进行全面评估。')

p = doc.add_paragraph()
add_bold_text(p, '技术指标分析：')

add_table_with_header(doc,
    ["指标", "说明", "应用场景"],
    [
        ["RSI（相对强弱指数）", "衡量价格变动速度和幅度", ">70超买，<30超卖"],
        ["MACD（移动平均收敛散度）", "趋势跟踪指标", "金叉看多，死叉看空"],
        ["KDJ（随机指标）", "短期超买超卖判断", "K线上穿D线为买入信号"],
        ["布林带（Bollinger Bands）", "价格波动区间", "触及上轨超买，下轨超卖"],
        ["ATR（平均真实波幅）", "波动率度量", "用于计算止损位和目标价"],
    ],
    col_widths=[3.5, 4.5, 5]
)

p = doc.add_paragraph()
add_bold_text(p, '基本面分析：')

fundamental_items = [
    ("PE（市盈率）：", "评估估值水平，与行业均值对比。"),
    ("PB（市净率）：", "评估资产价值，适用于银行等重资产行业。"),
    ("ROE（净资产收益率）：", "评估盈利能力，持续高ROE是优质公司标志。"),
]

for bold_part, normal_part in fundamental_items:
    add_bullet(doc, normal_part, bold_prefix=bold_part)

p = doc.add_paragraph()
add_bold_text(p, '资金流向分析：')
add_normal_text(p, '追踪主力资金进出情况，判断资金态度。')

doc.add_heading('5.2  新闻舆情智能体', level=2)

p = doc.add_paragraph()
add_normal_text(p, '新闻舆情智能体（News Sentiment Agent）负责采集和分析与目标股票相关的新闻资讯和社交媒体情绪。')

news_features = [
    ("中文新闻采集：", "通过AKShare等数据源获取A股相关新闻，覆盖财经媒体、官方公告等。"),
    ("社交媒体情绪：", "分析社交平台讨论热度与情绪倾向（正面/负面/中性）。"),
    ("事件提取：", "自动识别新闻中的关键事件（业绩预告/高管变动/政策影响等），评估事件对股价的潜在影响。"),
]

for bold_part, normal_part in news_features:
    add_bullet(doc, normal_part, bold_prefix=bold_part)

doc.add_heading('5.3  风控合规智能体', level=2)

p = doc.add_paragraph()
add_normal_text(p, '风控合规智能体（Risk Compliance Agent）提供4维度风险评估与合规规则引擎。')

p = doc.add_paragraph()
add_bold_text(p, '4维风控体系：')

add_table_with_header(doc,
    ["维度", "评估内容", "关键指标"],
    [
        ["波动率风险", "价格波动程度评估", "ATR、历史波动率、VaR"],
        ["市场风险", "系统性风险暴露程度", "Beta系数、相关性分析"],
        ["舆情风险", "负面信息影响评估", "负面新闻数量、情绪恶化程度"],
        ["流动性风险", "交易流动性评估", "日均成交量、买卖价差"],
    ],
    col_widths=[3, 4.5, 5.5]
)

p = doc.add_paragraph()
add_bold_text(p, '合规规则引擎：')
add_normal_text(p, '自动检查投资建议是否符合合规要求，详见第6章。')

doc.add_heading('5.4  投资策略智能体', level=2)

p = doc.add_paragraph()
add_normal_text(p, '投资策略智能体（Strategy Agent）基于其他智能体的分析结果，生成具体的投资建议。')

strategy_features = [
    ("仓位建议：", "根据风险等级和置信度，建议合理的仓位比例。"),
    ("目标价与止损价：", "基于ATR计算目标价（当前价 + 3×ATR）和止损价（当前价 - 2×ATR）。"),
    ("组合分析：", "支持持仓组合的风险评估与配置优化，通过 /api/portfolio/analyze 接口调用。"),
]

for bold_part, normal_part in strategy_features:
    add_bullet(doc, normal_part, bold_prefix=bold_part)

doc.add_heading('5.5  报告生成智能体', level=2)

p = doc.add_paragraph()
add_normal_text(p, '报告生成智能体（Report Agent）支持5种专业报告模板的自动生成：')

report_types = [
    ("晨报（morning_daily）：", "每日开盘前自动生成，包含市场概览、重要资讯、异动股票。"),
    ("研报（stock_research）：", "个股深度研究报告，包含基本面、技术面、估值分析。"),
    ("风控周报（risk_weekly）：", "周度风险回顾，包含风险事件统计、合规检查结果。"),
    ("组合月报（portfolio_monthly）：", "月度组合表现分析，包含收益归因、风险调整收益。"),
    ("事件快报（event_flash）：", "重大事件即时分析，快速响应市场突发事件。"),
]

for bold_part, normal_part in report_types:
    add_bullet(doc, normal_part, bold_prefix=bold_part)

doc.add_heading('5.6  执行监控智能体', level=2)

p = doc.add_paragraph()
add_normal_text(p, '执行监控智能体（Execution Agent）是FinAgent Pro"主动智能"的核心实现，负责定时巡检和主动预警。')

exec_features = [
    ("定时巡检：", "按配置的时间表自动执行市场扫描、风险检查等任务。"),
    ("主动预警：", "发现异常情况（价格异动、风险超标等）时主动触发预警通知。"),
    ("任务调度：", "管理晨报/晚报的定时生成、定期风险扫描等周期性任务。"),
]

for bold_part, normal_part in exec_features:
    add_bullet(doc, normal_part, bold_prefix=bold_part)

doc.add_page_break()


# ============================================================
# 第6章 合规与安全
# ============================================================

doc.add_heading('第6章  合规与安全', level=1)

doc.add_heading('6.1  合规规则', level=2)

p = doc.add_paragraph()
add_normal_text(p, 'FinAgent Pro 内置完善的合规规则引擎，所有投资建议均经过合规审查。默认规则如下：')

add_table_with_header(doc,
    ["规则", "限制", "说明"],
    [
        ["单股集中度", "≤ 10%", "单一标的不超过组合总资产的10%"],
        ["行业集中度", "≤ 30%", "单一行业不超过组合总资产的30%"],
        ["创业板限制", "≤ 20%", "创业板股票合计不超过组合总资产的20%"],
        ["ST股禁止", "0%", "禁止买入ST/*ST股票"],
        ["科创板提示", "-", "科创板涨跌幅20%，需特别关注风险"],
    ],
    col_widths=[3, 2.5, 8.5]
)

p = doc.add_paragraph()
add_normal_text(p, '合规规则可通过环境变量动态配置（MAX_SINGLE_STOCK_RATIO、MAX_SECTOR_RATIO、MAX_GEM_RATIO），也可通过修改 compliance_service.py 添加自定义规则。')

doc.add_heading('6.2  审计日志', level=2)

p = doc.add_paragraph()
add_normal_text(p, 'FinAgent Pro 提供')
add_bold_text(p, '全链路审计日志')
add_normal_text(p, '，记录所有关键操作：')

audit_items = [
    "每次股票分析的请求与结果",
    "所有合规检查的通过/违规记录",
    "预警触发与处理记录",
    "报告生成记录",
    "定时任务执行记录",
    "API访问日志",
]

for item in audit_items:
    add_bullet(doc, item)

p = doc.add_paragraph()
add_normal_text(p, '审计日志可通过 ')
add_code_block(doc, 'GET /api/audit/log?limit=50')
add_normal_text(p, ' 接口查询。')

doc.add_heading('6.3  数据安全', level=2)

security_items = [
    ("国产LLM数据不出境：", "默认使用智谱AI（GLM-5.1）、DeepSeek、商汤日日新等国产大模型，确保金融数据不离开国内。"),
    ("敏感信息脱敏：", "API密钥等敏感配置通过环境变量管理，不硬编码在代码中；日志中自动脱敏处理。"),
    ("访问控制：", "支持API访问密钥认证（API_ACCESS_ENABLED），防止未授权访问。"),
    ("本地部署：", "支持完全本地化部署（Docker/源码），数据不经过第三方服务器。"),
]

for bold_part, normal_part in security_items:
    add_bullet(doc, normal_part, bold_prefix=bold_part)

doc.add_page_break()


# ============================================================
# 第7章 常见问题（FAQ）
# ============================================================

doc.add_heading('第7章  常见问题（FAQ）', level=1)

faqs = [
    (
        "Q1：LLM服务不可用怎么办？",
        "当LLM服务不可用时（如API密钥无效、网络故障等），系统会自动降级到规则引擎模式。规则引擎模式下，智能体仍可基于技术指标计算和规则逻辑提供基础分析结果，但不会包含LLM增强的深度推理。系统会在界面上提示\u201cLLM增强\uff1a否\u201d，方便用户识别。"
    ),
    (
        "Q2：AKShare数据获取失败怎么办？",
        "当AKShare数据源不可用时，系统会自动使用内置的模拟数据（Mock Data）进行分析，确保功能可用。模拟数据仅用于演示和测试，不代表真实行情。恢复数据源后，系统会自动切换回真实数据。"
    ),
    (
        "Q3：如何添加新的合规规则？",
        "修改 backend/services/compliance_service.py 文件中的 ComplianceService 类。在 self.rules 字典中添加新规则，并在 check() 方法中添加对应的检查逻辑。修改后重启后端服务即可生效。"
    ),
    (
        "Q4：如何更换LLM模型？",
        "修改 .env 文件中的 LLM_PROVIDER 变量即可切换模型：\n"
        "• glm —— 使用智谱AI GLM-5.1（主力模型，综合能力最强）\n"
        "• deepseek —— 使用DeepSeek-v4-pro（深度推理，适合复杂分析）\n"
        "• sensenova —— 使用商汤日日新（轻量快速，适合高频调用）\n"
        "修改后重启后端服务生效。"
    ),
    (
        "Q5：如何自定义定时任务？",
        "有两种方式：\n"
        "1. 通过API创建：调用 POST /api/tasks 接口，传入 task_type、cron_expression 等参数。\n"
        "2. 修改代码：编辑 backend/services/scheduler_service.py，在默认任务配置中添加或修改定时任务。"
    ),
]

for q, a in faqs:
    p = doc.add_paragraph()
    add_bold_text(p, q)
    p = doc.add_paragraph()
    add_normal_text(p, a)
    doc.add_paragraph('')  # 空行分隔

doc.add_page_break()


# ============================================================
# 附录A 信号说明
# ============================================================

doc.add_heading('附录A  信号说明', level=1)

p = doc.add_paragraph()
add_normal_text(p, 'FinAgent Pro 使用以下5级信号体系表示投资建议方向：')

add_table_with_header(doc,
    ["信号标识", "中文名称", "数值映射", "含义说明"],
    [
        ["strong_buy", "强烈买入", "+1.0", "多维度强烈看多，建议积极建仓"],
        ["buy", "买入", "+0.5", "偏多信号，可考虑适度建仓"],
        ["hold", "持有", "0", "中性信号，建议观望或维持现有仓位"],
        ["sell", "卖出", "-0.5", "偏空信号，建议适度减仓"],
        ["strong_sell", "强烈卖出", "-1.0", "多维度强烈看空，建议果断清仓"],
    ],
    col_widths=[3, 3, 2.5, 5.5]
)

p = doc.add_paragraph()
add_bold_text(p, '信号生成机制：')
add_normal_text(p, '最终信号由各智能体信号加权平均得出。加权公式为：weighted_signal = Σ(signal_value × confidence) / Σ(confidence)。加权值 > 0.5 为强烈买入，> 0.2 为买入，< -0.5 为强烈卖出，< -0.2 为卖出，其余为持有。')


# ============================================================
# 附录B 风险等级说明
# ============================================================

doc.add_heading('附录B  风险等级说明', level=1)

p = doc.add_paragraph()
add_normal_text(p, 'FinAgent Pro 使用以下4级风险等级体系：')

add_table_with_header(doc,
    ["等级标识", "中文名称", "风险分数范围", "含义说明"],
    [
        ["low", "低风险", "0.0 - 0.25", "风险可控，可正常投资"],
        ["medium", "中等风险", "0.25 - 0.50", "存在一定风险，需关注关键指标"],
        ["high", "高风险", "0.50 - 0.75", "风险较大，建议谨慎操作"],
        ["critical", "极高风险", "0.75 - 1.0", "风险极高，建议立即减仓或清仓"],
    ],
    col_widths=[3, 3, 3, 5]
)

p = doc.add_paragraph()
add_bold_text(p, '风险评估维度：')
add_normal_text(p, '综合波动率风险、市场风险、舆情风险和流动性风险4个维度，加权计算得出总体风险分数（risk_score）和风险等级（overall_risk）。')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_normal_text(p, '—— 文档结束 ——')


# ============================================================
# 保存文档
# ============================================================

output_path = r'D:\AFAC2026金融智能创新大赛\FinAgent-Pro\docs\FinAgentPro用户使用手册.docx'
doc.save(output_path)
print(f"文档已生成: {output_path}")

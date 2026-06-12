# -*- coding: utf-8 -*-
"""生成 FinAgent Pro 核心技术 Word 文档"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = r"D:\AFAC2026金融智能创新大赛\FinAgent-Pro\docs\FinAgentPro核心技术.docx"

DARK_BLUE = RGBColor(0x1A, 0x3C, 0x6E)
ACCENT_BLUE = RGBColor(0x2B, 0x5C, 0x9E)
LIGHT_BLUE_BG = "D6E4F0"
WHITE = "FFFFFF"
DARK_BG = "1A3C6E"
MID_BG = "2B5C9E"
LIGHT_BG = "E8F0FE"
GRAY_TEXT = RGBColor(0x55, 0x55, 0x55)

FONT_NAME = "微软雅黑"


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_run(paragraph, text, bold=False, size=11, color=None, font_name=FONT_NAME):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color
    return run


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.15):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = line_spacing


def add_header(doc):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "FinAgent Pro 核心技术", bold=True, size=9, color=DARK_BLUE)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="1A3C6E"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.size = Pt(9)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)
    run2 = p.add_run()
    run2.font.size = Pt(9)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._element.append(instrText)
    run3 = p.add_run()
    run3.font.size = Pt(9)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fldChar2)


def create_cover_page(doc):
    for _ in range(6):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=12, line_spacing=1.0)
    add_run(p, "FinAgent Pro", bold=True, size=36, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.0)
    add_run(p, "核心技术", bold=True, size=28, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=12, after=12, line_spacing=1.0)
    add_run(p, "━" * 30, size=12, color=ACCENT_BLUE)

    for text in ["AFAC2026金融智能创新大赛 初创组", "智融先锋团队", "2026年6月"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=0, after=4, line_spacing=1.5)
        add_run(p, text, size=14, color=GRAY_TEXT)

    doc.add_page_break()


def add_section_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        run.font.color.rgb = DARK_BLUE
        run.font.size = Pt(18 if level == 1 else 14)
    set_paragraph_spacing(heading, before=18, after=10)
    return heading


def add_body_paragraph(doc, text, bold_prefix=None, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    set_paragraph_spacing(p, before=2, after=4, line_spacing=1.5)
    if bold_prefix:
        add_run(p, bold_prefix, bold=True, size=11)
    add_run(p, text, size=11)
    return p


def set_table_borders(table, color="1A3C6E", sz="6"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_tech_overview_table(doc):
    """添加技术栈总览表格"""
    table = doc.add_table(rows=8, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row in table.rows:
        row.cells[0].width = Cm(3)
        row.cells[1].width = Cm(5.5)
        row.cells[2].width = Cm(7.5)

    headers = ["技术领域", "核心技术", "关键能力"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, DARK_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=4, after=4)
        add_run(p, h, bold=True, size=10, color=RGBColor.from_string(WHITE))

    rows_data = [
        ["大语言模型", "ReAct推理引擎", "Thought-Action-Observation自主推理循环"],
        ["多智能体协同", "Master Orchestrator", "三阶段流水线+加权协商+LLM增强推理"],
        ["记忆系统", "三层记忆架构", "工作记忆+情节记忆+语义记忆"],
        ["主动智能", "APScheduler+WebSocket", "定时巡检+事件驱动+阈值预警"],
        ["合规引擎", "规则引擎+审计中间件", "四维风控+合规约束+全链路审计"],
        ["金融计算", "AKShare+ta-lib", "A股行情+技术指标+VaR风险计算"],
        ["工程架构", "FastAPI+Vue3+Docker", "异步API+暗色主题+容器化部署"],
    ]

    for row_idx, row_data in enumerate(rows_data):
        bg = LIGHT_BG if row_idx % 2 == 0 else WHITE
        for col_idx, text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx < 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, before=3, after=3)
            is_bold = col_idx == 0 or col_idx == 1
            color = DARK_BLUE if col_idx <= 1 else RGBColor(0x33, 0x33, 0x33)
            add_run(p, text, bold=is_bold, size=10, color=color)

    set_table_borders(table)


def add_react_detail(doc):
    """ReAct推理引擎详细说明"""
    # 流程图（ASCII）
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=6, line_spacing=1.0)
    diagram = (
        "    ┌──────────┐     ┌──────────┐     ┌──────────┐\n"
        "    │  Thought  │────▶│  Action   │────▶│Observation│\n"
        "    │  思考推理  │     │  执行动作  │     │  观察结果  │\n"
        "    └──────────┘     └──────────┘     └──────────┘\n"
        "         ▲                                    │\n"
        "         └────────────────────────────────────┘\n"
        "                      循环迭代"
    )
    run = add_run(p, diagram, size=9, color=DARK_BLUE, font_name="Consolas")
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{LIGHT_BLUE_BG}"/>')
    p._element.get_or_add_pPr().append(shd)

    add_body_paragraph(doc,
        "ReAct推理引擎是FinAgent Pro的核心创新，区别于传统单次调用模式。"
        "每个智能体在ReAct循环中：先思考需要什么信息（Thought），"
        "再调用工具获取数据（Action），最后观察结果并决定下一步（Observation）。"
        "这种循环机制让智能体像人类分析师一样逐步推理，而非一次性给出答案。", indent=True)

    # ReAct示例表格
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(5)
        row.cells[2].width = Cm(8.5)

    headers = ["步骤", "类型", "内容"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, DARK_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=4, after=4)
        add_run(p, h, bold=True, size=10, color=RGBColor.from_string(WHITE))

    examples = [
        ["1", "Thought", "需要先获取贵州茅台的最新行情数据，判断当前价格趋势"],
        ["2", "Action", "调用 get_stock_quote(symbol='600519') 获取实时行情"],
        ["3", "Observation", "当前价1856元，涨幅2.3%，RSI=68接近超买区，需进一步分析"],
    ]
    for row_idx, (step, typ, content) in enumerate(examples):
        bg = LIGHT_BG if row_idx % 2 == 0 else WHITE
        for col_idx, text in enumerate([step, typ, content]):
            cell = table.rows[row_idx + 1].cells[col_idx]
            set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx < 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, before=3, after=3)
            add_run(p, text, bold=col_idx == 1, size=10,
                    color=ACCENT_BLUE if col_idx == 1 else RGBColor(0x33, 0x33, 0x33))

    set_table_borders(table)


def add_llm_chain_table(doc):
    """多LLM降级链表格"""
    table = doc.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        row.cells[0].width = Cm(2)
        row.cells[1].width = Cm(3.5)
        row.cells[2].width = Cm(4)
        row.cells[3].width = Cm(6.5)

    headers = ["优先级", "模型", "角色", "能力特点"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, DARK_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=4, after=4)
        add_run(p, h, bold=True, size=10, color=RGBColor.from_string(WHITE))

    models = [
        ["主力", "GLM-5.1", "核心推理", "千亿参数，中文理解强，金融推理准确"],
        ["备选", "DeepSeek-v4-pro", "降级容灾", "MoE架构，推理效率高，成本优化"],
        ["轻量", "SenseNova", "快速响应", "6.7B轻量模型，低延迟，简单任务"],
    ]
    colors = ["27AE60", "F39C12", "3498DB"]
    for row_idx, (priority, model, role, feature) in enumerate(models):
        bg = LIGHT_BG if row_idx % 2 == 0 else WHITE
        for col_idx, text in enumerate([priority, model, role, feature]):
            cell = table.rows[row_idx + 1].cells[col_idx]
            if col_idx == 0:
                set_cell_shading(cell, colors[row_idx])
                text_color = RGBColor.from_string(WHITE)
            else:
                set_cell_shading(cell, bg)
                text_color = DARK_BLUE if col_idx == 1 else RGBColor(0x33, 0x33, 0x33)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before=3, after=3)
            add_run(p, text, bold=col_idx <= 1, size=10, color=text_color)

    set_table_borders(table)


def add_pipeline_diagram(doc):
    """三阶段流水线流程图"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=6, line_spacing=1.0)
    diagram = (
        "  ┌─────────── Phase 1（并行）──────────┐\n"
        "  │  ┌──────────┐    ┌──────────┐       │\n"
        "  │  │市场分析   │    │新闻舆情   │       │\n"
        "  │  │智能体     │    │智能体     │       │\n"
        "  │  └─────┬────┘    └─────┬────┘       │\n"
        "  └────────┼───────────────┼────────────┘\n"
        "           │               │\n"
        "           ▼               ▼\n"
        "  ┌─────────── Phase 2（串行）──────────┐\n"
        "  │  ┌──────────────────────────┐       │\n"
        "  │  │    风控合规智能体          │       │\n"
        "  │  └────────────┬─────────────┘       │\n"
        "  └───────────────┼─────────────────────┘\n"
        "                  │\n"
        "                  ▼\n"
        "  ┌─────────── Phase 3（策略）──────────┐\n"
        "  │  ┌──────────┐    ┌──────────┐       │\n"
        "  │  │投资策略   │    │报告生成   │       │\n"
        "  │  │智能体     │    │智能体     │       │\n"
        "  │  └──────────┘    └──────────┘       │\n"
        "  └─────────────────────────────────────┘"
    )
    run = add_run(p, diagram, size=8.5, color=DARK_BLUE, font_name="Consolas")
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{LIGHT_BLUE_BG}"/>')
    p._element.get_or_add_pPr().append(shd)


def add_memory_table(doc):
    """三层记忆系统表格"""
    table = doc.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(2.5)
        row.cells[2].width = Cm(4)
        row.cells[3].width = Cm(7)

    headers = ["记忆层", "容量", "存储内容", "作用"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, DARK_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=4, after=4)
        add_run(p, h, bold=True, size=10, color=RGBColor.from_string(WHITE))

    memories = [
        ["工作记忆", "50条", "当前会话信息", "保存当前分析任务的上下文，如正在分析的股票、已获取的数据"],
        ["情节记忆", "100条", "历史交互事件", "记录用户历史查询、分析结果、偏好设置，实现个性化服务"],
        ["语义记忆", "持久", "金融知识库", "存储行业术语、分析框架、合规规则等长期知识"],
    ]
    for row_idx, (layer, cap, content, role) in enumerate(memories):
        bg = LIGHT_BG if row_idx % 2 == 0 else WHITE
        for col_idx, text in enumerate([layer, cap, content, role]):
            cell = table.rows[row_idx + 1].cells[col_idx]
            set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx < 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, before=3, after=3)
            add_run(p, text, bold=col_idx == 0, size=10,
                    color=DARK_BLUE if col_idx == 0 else RGBColor(0x33, 0x33, 0x33))

    set_table_borders(table)


def add_compliance_table(doc):
    """合规规则引擎表格"""
    table = doc.add_table(rows=5, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        row.cells[0].width = Cm(3)
        row.cells[1].width = Cm(4)
        row.cells[2].width = Cm(9)

    headers = ["规则名称", "约束条件", "说明"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, DARK_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=4, after=4)
        add_run(p, h, bold=True, size=10, color=RGBColor.from_string(WHITE))

    rules = [
        ["单股集中度", "≤ 10%", "单一股票持仓不超过组合总资产的10%，防范个股风险"],
        ["行业集中度", "≤ 30%", "单一行业持仓不超过组合总资产的30%，分散行业风险"],
        ["创业板限制", "≤ 20%", "创业板股票持仓不超过20%，控制高波动板块敞口"],
        ["ST股禁止", "0%", "完全禁止持有ST/*ST股票，规避退市风险"],
    ]
    for row_idx, (name, constraint, desc) in enumerate(rules):
        bg = LIGHT_BG if row_idx % 2 == 0 else WHITE
        for col_idx, text in enumerate([name, constraint, desc]):
            cell = table.rows[row_idx + 1].cells[col_idx]
            set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx < 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, before=3, after=3)
            add_run(p, text, bold=col_idx == 0, size=10,
                    color=DARK_BLUE if col_idx == 0 else RGBColor(0x33, 0x33, 0x33))

    set_table_borders(table)


def add_tech_stack_table(doc):
    """工程架构技术栈表格"""
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(5.5)
        row.cells[2].width = Cm(8)

    headers = ["层级", "技术栈", "关键特性"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, DARK_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=4, after=4)
        add_run(p, h, bold=True, size=10, color=RGBColor.from_string(WHITE))

    stacks = [
        ["后端", "FastAPI + WebSocket + SSE", "异步高性能API / 实时双向通信 / 服务端推送 / SQLite+PostgreSQL双存储"],
        ["前端", "Vue3 + TailwindCSS + Lightweight Charts", "组合式API / 暗色主题 / K线图表 / Pinia状态管理"],
        ["部署", "Docker + docker-compose + Vercel", "容器化 / 一键部署 / 前端CDN加速 / 健康检查"],
    ]
    for row_idx, (layer, tech, features) in enumerate(stacks):
        bg = LIGHT_BG if row_idx % 2 == 0 else WHITE
        for col_idx, text in enumerate([layer, tech, features]):
            cell = table.rows[row_idx + 1].cells[col_idx]
            set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, before=3, after=3)
            add_run(p, text, bold=col_idx <= 1, size=10,
                    color=DARK_BLUE if col_idx == 0 else RGBColor(0x33, 0x33, 0x33))

    set_table_borders(table)


def main():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    add_header(doc)
    add_footer(doc)
    create_cover_page(doc)

    # ========== 一、技术栈总览 ==========
    add_section_heading(doc, "一、技术栈总览")
    add_body_paragraph(doc,
        "FinAgent Pro采用7大核心技术领域协同构建，覆盖从大模型推理到工程部署的完整技术链路。"
        "以下表格展示了各技术领域的核心组件与关键能力。")
    add_tech_overview_table(doc)
    doc.add_paragraph()

    # ========== 二、大语言模型与推理 ==========
    add_section_heading(doc, "二、大语言模型与推理引擎")

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4)
    add_run(p, "▎ ReAct推理引擎", bold=True, size=13, color=DARK_BLUE)

    add_body_paragraph(doc,
        "ReAct（Reasoning + Acting）推理引擎是FinAgent Pro的核心创新，"
        "实现了Thought-Action-Observation自主推理循环。"
        "不同于传统单次调用模式，ReAct引擎让智能体能够像人类分析师一样逐步推理。", indent=True)

    add_react_detail(doc)
    doc.add_paragraph()

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4)
    add_run(p, "▎ 多LLM降级链", bold=True, size=13, color=DARK_BLUE)

    add_body_paragraph(doc,
        "采用三模型自动降级机制，通过OpenAI兼容接口统一调用，确保服务高可用。"
        "当主力模型不可用时，自动切换至备选模型，指数退避重试保障稳定性。", indent=True)

    add_llm_chain_table(doc)
    doc.add_paragraph()

    add_body_paragraph(doc,
        "结构化JSON输出约束LLM生成可解析的结构化结果，Markdown代码块自动清理，"
        "确保下游智能体和前端能够可靠地处理LLM输出。", indent=True)

    # ========== 三、多智能体协同 ==========
    add_section_heading(doc, "三、多智能体协同调度")

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4)
    add_run(p, "▎ 三阶段流水线", bold=True, size=13, color=DARK_BLUE)

    add_body_paragraph(doc,
        "Master Orchestrator采用三阶段流水线调度，确保6大专业智能体高效协同：", indent=True)

    add_pipeline_diagram(doc)

    phases = [
        ("Phase 1（并行执行）：", "市场分析智能体和新闻舆情智能体同时启动，"
         "分别计算技术指标/估值指标和抓取新闻/情感分析，互不依赖，最大化并行效率。"),
        ("Phase 2（串行执行）：", "风控合规智能体依赖Phase 1的结果，"
         "评估波动率/集中度/流动性四维风险，并执行合规规则引擎检查。"),
        ("Phase 3（策略生成）：", "投资策略智能体综合前两阶段结果，"
         "生成仓位建议和目标价，报告生成智能体自动输出专业报告。"),
    ]
    for prefix, text in phases:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        set_paragraph_spacing(p, before=2, after=2, line_spacing=1.5)
        add_run(p, prefix, bold=True, size=11, color=ACCENT_BLUE)
        add_run(p, text, size=11)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4)
    add_run(p, "▎ 加权协商算法", bold=True, size=13, color=DARK_BLUE)

    add_body_paragraph(doc,
        "6大智能体通过加权协商算法综合出最终建议："
        "Final Signal = Σ(signal_i x confidence_i) / Σ(confidence_i)，"
        "置信度高的智能体在协商中拥有更大权重。"
        "协商结果经LLM增强推理二次处理，生成专业分析文本，"
        "确保结论既有数据支撑又有专业解读。", indent=True)

    # ========== 四、记忆与知识系统 ==========
    add_section_heading(doc, "四、记忆与知识系统")

    add_body_paragraph(doc,
        "三层记忆架构让数字员工具备上下文理解和持续学习能力，"
        "区别于传统无状态AI对话系统。")

    add_memory_table(doc)
    doc.add_paragraph()

    add_body_paragraph(doc,
        "全局单例管理器确保智能体间共享记忆上下文，"
        "市场分析智能体获取的数据可被风控合规智能体直接引用，"
        "避免重复查询，提升协同效率。", indent=True)

    # ========== 五、主动智能机制 ==========
    add_section_heading(doc, "五、主动智能机制")

    add_body_paragraph(doc,
        "传统金融系统是被动响应式的，FinAgent Pro的主动智能让数字员工从被动变为主动：")

    pro_features = [
        ("定时巡检（APScheduler）：", "每日8:30自动生成晨报，17:00自动生成收盘总结，"
         "定时触发风险扫描，无需人工干预。"),
        ("事件驱动（WebSocket）：", "持仓股突发利空消息时，系统自动触发从新闻抓取到风险评估到策略调整的全链路分析，"
         "并实时推送结果至前端。"),
        ("阈值预警：", "当技术指标越界（RSI超买超卖）、风险指标超限时，"
         "系统主动发现并推送预警，不等用户查询。"),
    ]
    for prefix, text in pro_features:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        set_paragraph_spacing(p, before=2, after=2, line_spacing=1.5)
        add_run(p, prefix, bold=True, size=11, color=ACCENT_BLUE)
        add_run(p, text, size=11)

    # ========== 六、合规规则引擎 ==========
    add_section_heading(doc, "六、合规规则引擎")

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4)
    add_run(p, "▎ 四维风控体系", bold=True, size=13, color=DARK_BLUE)

    risk_dims = [
        ("波动率风险：", "基于历史波动率和ATR指标，评估价格波动风险等级"),
        ("市场风险：", "基于大盘指数和行业指数，评估系统性风险暴露"),
        ("舆情风险：", "基于新闻情感分析和社交情绪，评估信息面风险"),
        ("流动性风险：", "基于成交量和换手率，评估资产变现能力"),
    ]
    for prefix, text in risk_dims:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        set_paragraph_spacing(p, before=2, after=2, line_spacing=1.5)
        add_run(p, prefix, bold=True, size=11, color=ACCENT_BLUE)
        add_run(p, text, size=11)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4)
    add_run(p, "▎ 合规约束规则", bold=True, size=13, color=DARK_BLUE)

    add_compliance_table(doc)
    doc.add_paragraph()

    add_body_paragraph(doc,
        "ComplianceAuditMiddleware审计中间件记录/analyze、/chat、/portfolio等关键路径的请求，"
        "全链路操作留痕，满足银保监、证监会的监管审计要求。", indent=True)

    # ========== 七、金融数据与计算 ==========
    add_section_heading(doc, "七、金融数据与计算")

    data_items = [
        ("AKShare A股实时行情：", "覆盖沪深两市全量股票，5分钟缓存机制，模拟数据兜底保障可用性"),
        ("ta-lib技术指标：", "SMA/EMA均线、RSI强弱指标、MACD趋势指标、布林带通道、ATR波动率"),
        ("东方财富新闻API：", "中文新闻抓取，NLP情感分析，多源新闻聚合"),
        ("VaR历史模拟法：", "基于历史收益率分布，计算95%/99%置信区间下的最大可能损失"),
        ("多因子模型：", "综合技术因子、基本面因子、情绪因子，生成投资信号"),
    ]
    for prefix, text in data_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        set_paragraph_spacing(p, before=2, after=2, line_spacing=1.5)
        add_run(p, prefix, bold=True, size=11, color=ACCENT_BLUE)
        add_run(p, text, size=11)

    # ========== 八、工程架构 ==========
    add_section_heading(doc, "八、工程架构")

    add_tech_stack_table(doc)

    # 保存
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"文档已生成: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()

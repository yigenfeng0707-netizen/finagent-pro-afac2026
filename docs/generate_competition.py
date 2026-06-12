# -*- coding: utf-8 -*-
"""生成 FinAgent Pro 市场竞争分析 Word 文档"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = r"D:\AFAC2026金融智能创新大赛\FinAgent-Pro\docs\FinAgentPro市场竞争分析.docx"

DARK_BLUE = RGBColor(0x1A, 0x3C, 0x6E)
ACCENT_BLUE = RGBColor(0x2B, 0x5C, 0x9E)
LIGHT_BLUE_BG = "D6E4F0"
WHITE = "FFFFFF"
DARK_BG = "1A3C6E"
MID_BG = "2B5C9E"
LIGHT_BG = "E8F0FE"
GRAY_TEXT = RGBColor(0x55, 0x55, 0x55)
RED_TEXT = RGBColor(0xC0, 0x39, 0x2B)
GREEN_TEXT = RGBColor(0x27, 0xAE, 0x60)

FONT_NAME = "微软雅黑"


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_run(paragraph, text, bold=False, size=11, color=None, font_name=FONT_NAME):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color
    return run


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.15):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = line_spacing


def add_header(doc):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "FinAgent Pro 市场竞争分析", bold=True, size=9, color=DARK_BLUE)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="1A3C6E"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.size = Pt(9)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)
    run2 = p.add_run()
    run2.font.size = Pt(9)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._element.append(instrText)
    run3 = p.add_run()
    run3.font.size = Pt(9)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fldChar2)


def create_cover_page(doc):
    for _ in range(6):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=12, line_spacing=1.0)
    add_run(p, "FinAgent Pro", bold=True, size=36, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.0)
    add_run(p, "市场竞争分析", bold=True, size=28, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=12, after=12, line_spacing=1.0)
    add_run(p, "━" * 30, size=12, color=ACCENT_BLUE)

    for text in ["AFAC2026金融智能创新大赛 初创组", "智融先锋团队", "2026年6月"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=0, after=4, line_spacing=1.5)
        add_run(p, text, size=14, color=GRAY_TEXT)

    doc.add_page_break()


def add_section_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        run.font.color.rgb = DARK_BLUE
        run.font.size = Pt(18 if level == 1 else 14)
    set_paragraph_spacing(heading, before=18, after=10)
    return heading


def add_body_paragraph(doc, text, bold_prefix=None, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    set_paragraph_spacing(p, before=2, after=4, line_spacing=1.5)
    if bold_prefix:
        add_run(p, bold_prefix, bold=True, size=11)
    add_run(p, text, size=11)
    return p


def set_table_borders(table, color="1A3C6E", sz="6"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_competitor_section(doc, num, name, advantages, disadvantages, differentiation):
    """添加单个竞争者分析区块"""
    # 竞争者标题
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=10, after=6)
    add_run(p, f"▎ 竞争者{num}：{name}", bold=True, size=13, color=DARK_BLUE)

    # 优势
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    set_paragraph_spacing(p, before=2, after=2, line_spacing=1.5)
    add_run(p, "优势：", bold=True, size=11, color=GREEN_TEXT)
    for adv in advantages:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.2)
        set_paragraph_spacing(p, before=1, after=1, line_spacing=1.4)
        add_run(p, f"● {adv}", size=10.5)

    # 劣势
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    set_paragraph_spacing(p, before=2, after=2, line_spacing=1.5)
    add_run(p, "劣势：", bold=True, size=11, color=RED_TEXT)
    for dis in disadvantages:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.2)
        set_paragraph_spacing(p, before=1, after=1, line_spacing=1.4)
        add_run(p, f"● {dis}", size=10.5)

    # 差异化
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    set_paragraph_spacing(p, before=4, after=6, line_spacing=1.5)
    add_run(p, "FinAgent Pro差异化：", bold=True, size=11, color=ACCENT_BLUE)
    add_run(p, differentiation, size=11)


def add_comparison_table(doc):
    """添加竞争对比矩阵表格"""
    table = doc.add_table(rows=5, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_widths = [Cm(3), Cm(3), Cm(3), Cm(3), Cm(3.5)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]

    headers = ["对比维度", "传统金融终端", "通用AI平台", "海外金融AI", "FinAgent Pro"]
    header_bg = [DARK_BG, "7F8C8D", "7F8C8D", "7F8C8D", "27AE60"]

    for i, (h, bg) in enumerate(zip(headers, header_bg)):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=4, after=4)
        color = RGBColor.from_string(WHITE)
        add_run(p, h, bold=True, size=10, color=color)

    rows_data = [
        ["主动智能", "✗ 被动工具", "✗ 被动问答", "△ 部分支持", "✓ 三层主动智能"],
        ["多智能体协同", "✗ 无", "✗ 单模型调用", "△ 简单编排", "✓ 6智能体+Orchestrator"],
        ["合规内嵌", "△ 附加模块", "✗ 无", "✗ 不适配中国", "✓ 规则引擎+审计日志"],
        ["A股适配", "✓ 数据全", "△ 通用数据", "✗ 不支持", "✓ 原生A股+中文NLP"],
    ]

    for row_idx, row_data in enumerate(rows_data):
        bg = LIGHT_BG if row_idx % 2 == 0 else WHITE
        for col_idx, text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            if col_idx == 4:
                set_cell_shading(cell, "E8F5E9")
            else:
                set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before=3, after=3)
            is_bold = col_idx == 0 or col_idx == 4
            color = DARK_BLUE if col_idx == 0 else (GREEN_TEXT if col_idx == 4 else RGBColor(0x33, 0x33, 0x33))
            add_run(p, text, bold=is_bold, size=9.5, color=color)

    set_table_borders(table)


def add_moat_section(doc):
    """添加竞争壁垒部分"""
    moats = [
        ("技术护城河", "三层架构（智能体引擎+协同调度+主动智能）是核心壁垒，"
         "竞品仅做到单层——传统终端只有数据展示，通用AI只有模型调用，"
         "海外产品只有简单编排。FinAgent Pro的三层协同需要深厚的金融场景理解和系统工程能力，难以短期复制。"),
        ("准入壁垒", "合规内嵌（规则引擎+审计日志）是金融机构采购的硬性要求。"
         "传统终端的合规是附加模块，通用AI平台完全没有合规能力。"
         "FinAgent Pro从架构层面将合规作为一等公民，每步操作可审计可追溯，"
         "满足银保监、证监会的监管要求，这是进入金融行业的入场券。"),
        ("体验壁垒", "主动智能（定时巡检+事件驱动+阈值预警）让数字员工从被动变为主动，"
         "用户无需主动查询，系统自动发现风险、推送预警、生成报告。"
         "这种'从人找信息到信息找人'的体验跃迁，一旦用户习惯就难以回到被动模式。"),
    ]

    for i, (title, desc) in enumerate(moats):
        # 壁垒标题
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=8, after=4)
        add_run(p, f"▎ 壁垒{i+1}：{title}", bold=True, size=13, color=DARK_BLUE)

        add_body_paragraph(doc, desc, indent=True)


def main():
    doc = Document()

    # 默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    add_header(doc)
    add_footer(doc)
    create_cover_page(doc)

    # ========== 一、市场格局 ==========
    add_section_heading(doc, "一、市场格局概览")
    add_body_paragraph(doc,
        "当前金融智能体市场呈现三足鼎立格局：传统金融终端占据数据与用户优势，"
        "通用AI平台拥有模型能力但缺乏垂直深度，海外金融AI技术成熟但无法适配中国市场。"
        "FinAgent Pro以「金融垂直+自主智能体+合规内嵌」的独特定位，切入三者之间的空白地带。")

    # ========== 二、竞争者分析 ==========
    add_section_heading(doc, "二、竞争者详细分析")

    add_competitor_section(doc, "1", "传统金融终端（同花顺、东方财富、Wind）",
        advantages=[
            "数据覆盖广：20年+金融数据积累，覆盖A股全市场",
            "用户基数大：同花顺注册用户超6亿，日活超2000万",
            "品牌壁垒深：金融机构采购白名单，替换成本极高",
        ],
        disadvantages=[
            "被动工具属性：需人工驱动分析流程，系统不会主动发现问题",
            "AI能力薄弱：仅做数据展示和简单指标计算，无智能推理",
            "合规功能为附加模块：非原生设计，审计追溯能力有限",
        ],
        differentiation="主动智能体替代被动工具，从「人找数据」变为「数据找人」。"
        "6大专业智能体自主完成从信息采集到决策建议的全链路分析，"
        "用户无需逐个功能模块操作，一句话即可触发完整分析流程。"
    )

    add_competitor_section(doc, "2", "通用AI平台（百度文心、阿里通义、智谱GLM）",
        advantages=[
            "大模型能力强：千亿参数模型，通用理解和生成能力出色",
            "算力资源丰富：云厂商基础设施，弹性扩缩容",
            "生态体系完善：插件市场、开发者社区、行业解决方案",
        ],
        disadvantages=[
            "缺乏金融垂直深度：不懂技术指标含义、不懂合规规则、不懂交易逻辑",
            "无合规规则引擎：无法满足金融监管审计要求，机构无法采购",
            "无多智能体协同：单模型调用模式，无法完成复杂金融任务链",
        ],
        differentiation="金融垂直场景深度定制，6大专业智能体协同+合规内嵌。"
        "不是「通用AI套金融壳」，而是从架构层面为金融场景设计的专业系统——"
        "每个智能体内置金融领域知识，Orchestrator理解金融分析流程，"
        "合规规则引擎确保每步操作可审计可追溯。"
    )

    add_competitor_section(doc, "3", "海外金融AI（Bloomberg GPT、Kensho）",
        advantages=[
            "技术成熟：Bloomberg GPT 500亿参数金融专用模型，Kensho被S&P Global收购",
            "海外市场验证：华尔街机构已广泛采用，商业模式成熟",
            "数据资源丰富：全球金融市场数据覆盖",
        ],
        disadvantages=[
            "不支持A股数据源和中文语境：无法处理中文新闻、A股交易规则",
            "无法满足国内金融监管合规要求：银保监、证监会审计标准不同",
            "数据出境风险：金融机构核心数据不能出境，无法采用海外方案",
        ],
        differentiation="国产大模型基座+A股数据源+中文NLP+合规规则引擎，数据不出境，安全可信。"
        "GLM-5.1/DeepSeek/SenseNova三模型降级链确保服务可用性，"
        "AKShare+东方财富提供A股实时数据，中文NLP精准分析国内新闻舆情，"
        "合规规则引擎内置中国金融监管规则，满足审计要求。"
    )

    # ========== 三、竞争对比矩阵 ==========
    add_section_heading(doc, "三、竞争对比矩阵")

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=8)
    add_run(p, "▎ 四维对比分析", bold=True, size=13, color=DARK_BLUE)

    add_comparison_table(doc)

    doc.add_paragraph()

    # ========== 四、核心竞争壁垒 ==========
    add_section_heading(doc, "四、核心竞争壁垒")

    add_body_paragraph(doc,
        "FinAgent Pro的竞争优势不仅在于功能差异，更在于三层壁垒的系统性构建：")

    add_moat_section(doc)

    # ========== 五、竞争策略 ==========
    add_section_heading(doc, "五、竞争策略")

    strategies = [
        ("差异化切入", "避开与传统终端的数据竞争，以「自主智能体」为切入点，"
         "先服务传统终端无法覆盖的主动分析和智能决策场景，"
         "再逐步向数据层扩展，实现从工具到平台的跃迁。"),
        ("合规先行", "将合规内嵌作为核心卖点，优先获取金融机构采购准入，"
         "建立「合规=FinAgent Pro」的心智认知，形成准入壁垒。"),
        ("生态共建", "开放智能体开发框架，吸引第三方开发者构建垂直场景智能体，"
         "形成「金融智能体应用市场」，从产品竞争升级为生态竞争。"),
    ]

    for i, (title, desc) in enumerate(strategies):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=6, after=4)
        add_run(p, f"▎ 策略{i+1}：{title}", bold=True, size=13, color=DARK_BLUE)
        add_body_paragraph(doc, desc, indent=True)

    # 保存
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"文档已生成: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()

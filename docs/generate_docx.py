#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
重新生成 FinAgent Pro 项目申报书，更新团队信息为一人团队（OPC模式）
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# ============================================================
# 常量定义
# ============================================================
FONT_NAME = '微软雅黑'
COLOR_DARK_BLUE = RGBColor(0x1A, 0x3C, 0x6E)   # #1A3C6E
COLOR_MED_BLUE = RGBColor(0x2B, 0x57, 0x9A)     # #2B579A
COLOR_LINK_BLUE = RGBColor(0x00, 0x6B, 0xD6)    # #006BD6
COLOR_GRAY = RGBColor(0x55, 0x55, 0x55)          # #555555
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)          # #FFFFFF
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)          # #000000
TABLE_HEADER_BG = '1A3C6E'

HEADING1_SIZE = Pt(18)   # 228600 EMU
HEADING2_SIZE = Pt(14)   # 177800 EMU
BODY_SIZE = Pt(10.5)     # 133350 EMU
COVER_TITLE_SIZE = Pt(28) # 355600 EMU
COVER_SUB_SIZE = Pt(18)   # 228600 EMU
COVER_LINE_SIZE = Pt(14)  # 177800 EMU
COVER_INFO_SIZE = Pt(14)  # 177800 EMU
FOOTER_SIZE = Pt(12)      # 152400 EMU

OUTPUT_PATH = r'D:\AFAC2026金融智能创新大赛\FinAgent-Pro\docs\FinAgentPro项目申报书.docx'


def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_run(paragraph, text, font_name=FONT_NAME, font_size=None, bold=None, color=None):
    """添加格式化run"""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        run.font.size = font_size
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return run


def add_heading1(doc, text):
    """添加一级标题"""
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    add_run(p, text, font_size=HEADING1_SIZE, bold=True, color=COLOR_DARK_BLUE)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_heading2(doc, text):
    """添加二级标题"""
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 2']
    add_run(p, text, font_size=HEADING2_SIZE, bold=True, color=COLOR_MED_BLUE)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text):
    """添加正文段落"""
    p = doc.add_paragraph()
    add_run(p, text, font_size=BODY_SIZE, color=COLOR_BLACK)
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, bold=False):
    """添加列表项"""
    p = doc.add_paragraph(style='List Bullet')
    add_run(p, text, font_size=BODY_SIZE, bold=bold)
    return p


def make_table(doc, headers, rows, col_widths=None):
    """创建格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for ci, header in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, header, font_size=BODY_SIZE, bold=True, color=COLOR_WHITE)
        set_cell_shading(cell, TABLE_HEADER_BG)

    # 数据行
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            add_run(p, str(val), font_size=BODY_SIZE)

    # 设置列宽
    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, width in enumerate(col_widths):
                row.cells[ci].width = Cm(width)

    return table


def build_document():
    doc = Document()

    # ============================================================
    # 页面设置
    # ============================================================
    section = doc.sections[0]
    section.page_width = Emu(7560310)
    section.page_height = Emu(10692130)
    section.left_margin = Emu(1141095)
    section.right_margin = Emu(1141095)
    section.top_margin = Emu(914400)
    section.bottom_margin = Emu(914400)

    # 页眉
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(hp, 'AFAC2026 项目申报书', font_size=Pt(9), color=COLOR_GRAY)

    # ============================================================
    # 封面页
    # ============================================================
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'AFAC2026金融智能创新大赛', font_size=COVER_TITLE_SIZE, bold=True, color=COLOR_DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '初创组 · 项目申报书', font_size=COVER_SUB_SIZE, color=COLOR_MED_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '━' * 30, font_size=COVER_LINE_SIZE, color=COLOR_LINK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '项目名称：FinAgent Pro 金融自主智能体平台', font_size=COVER_INFO_SIZE, color=COLOR_GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '团队名称：智融先锋', font_size=COVER_INFO_SIZE, color=COLOR_GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '负责人：冯亦根', font_size=COVER_INFO_SIZE, color=COLOR_GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '2026年6月', font_size=COVER_INFO_SIZE, color=COLOR_GRAY)

    # 分页
    doc.add_page_break()

    # ============================================================
    # 一、项目基本信息
    # ============================================================
    add_heading1(doc, '一、项目基本信息')

    info_rows = [
        ['项目名称', 'FinAgent Pro 金融自主智能体平台'],
        ['赛道方向', '金融科技创新应用'],
        ['参赛组别', '初创组'],
        ['团队名称', '智融先锋'],
        ['项目负责人', '冯亦根'],
        ['联系电话', '138-XXXX-XXXX'],
        ['电子邮箱', 'contact@finagent-pro.com'],
        ['所在单位', '浙江大学'],
        ['项目阶段', 'MVP开发阶段'],
        ['知识产权', '软件著作权申请中'],
    ]
    make_table(doc, ['项目', '内容'], info_rows, col_widths=[4, 12])

    doc.add_page_break()

    # ============================================================
    # 二、项目概述
    # ============================================================
    add_heading1(doc, '二、项目概述')

    add_heading2(doc, '2.1 项目简介')
    add_body(doc, 'FinAgent Pro是一款面向金融行业的自主智能体（Agentic AI）平台，通过六大专业数字员工——市场分析师、新闻解读员、风控专员、策略顾问、交易执行员和报告撰写员——为金融机构提供7×24小时全链路智能投研服务。平台基于ReAct推理引擎实现智能体自主决策，通过多智能体协同机制实现深度分析，以主动智能系统实现实时预警，并内置合规引擎确保全流程合规。')

    add_heading2(doc, '2.2 核心创新——六大创新点')
    innovation_rows = [
        ['1', 'ReAct自主推理', '自研推理引擎，实现「思考-行动-观察」闭环，智能体可自主规划任务、调用工具、评估结果'],
        ['2', '多智能体协同', '6大专业智能体分工协作，Orchestrator统一调度，加权协商算法消解冲突'],
        ['3', '主动智能系统', '定时巡检+事件驱动+阈值预警+计划执行，7×24小时主动服务'],
        ['4', '合规引擎内嵌', '200+监管规则内置，所有输出自动合规审查，审计全程留痕'],
        ['5', '三层记忆架构', '短期+工作+长期记忆，智能体持续学习优化，越用越智能'],
        ['6', '思维链可视化', '推理过程透明可追溯，决策可解释，增强用户信任'],
    ]
    make_table(doc, ['序号', '创新点', '创新描述'], innovation_rows, col_widths=[1.5, 3, 11.5])

    add_heading2(doc, '2.3 应用场景')
    add_bullet(doc, '投研分析：分析师使用智能体快速完成多维度投研分析，效率提升10倍')
    add_bullet(doc, '风控预警：风控团队7×24小时实时监控，风险事件秒级响应')
    add_bullet(doc, '合规审查：合规团队借助合规引擎自动审查，效率提升80%')
    add_bullet(doc, '投资决策：投资经理获取多智能体协同分析结论，辅助科学决策')
    add_bullet(doc, '个人投研：个人投资者享受专业级投研服务，降低信息不对称')

    doc.add_page_break()

    # ============================================================
    # 三、技术与产品
    # ============================================================
    add_heading1(doc, '三、技术与产品')

    add_heading2(doc, '3.1 技术路线——Agentic AI + 多智能体协同')
    add_body(doc, 'FinAgent Pro采用Agentic AI技术路线，以ReAct推理引擎为核心，结合多智能体协同、主动智能和合规引擎，构建金融行业专属的自主智能体平台。技术路线的核心优势在于：')
    add_bullet(doc, '自主性：智能体可自主规划、决策和行动，而非被动响应')
    add_bullet(doc, '协同性：多智能体分工协作，1+1>2的群体智能')
    add_bullet(doc, '主动性：无需用户触发，主动发现和推送关键信息')
    add_bullet(doc, '合规性：合规引擎全程嵌入，确保输出安全合规')

    add_heading2(doc, '3.2 产品功能——六大数字员工')
    agent_rows = [
        ['市场分析师', '行情监测、技术分析、趋势识别', 'ReAct + 技术指标工具 + 定时巡检'],
        ['新闻解读员', '资讯解读、事件提取、情绪分析', 'ReAct + NLP工具 + 事件驱动'],
        ['风控专员', '风险评估、VaR计算、阈值预警', 'ReAct + 风控模型 + 实时监控'],
        ['策略顾问', '策略生成、历史回测、收益归因', 'ReAct + 回测引擎 + 因子分析'],
        ['交易执行员', '信号生成、执行优化、合规检查', 'ReAct + 信号算法 + 合规审查'],
        ['报告撰写员', '报告生成、数据可视化、合规审查', 'ReAct + 模板引擎 + 图表工具'],
    ]
    make_table(doc, ['数字员工', '核心功能', '技术实现'], agent_rows, col_widths=[3, 6.5, 6.5])

    add_heading2(doc, '3.3 技术指标')
    tech_rows = [
        ['单次分析响应时间', '< 5秒', '从用户提问到返回分析结果'],
        ['并发用户数', '100+', '同时在线使用用户数'],
        ['系统可用性', '99.9%', '年度停机时间<8.76小时'],
        ['合规审查覆盖率', '100%', '所有输出均经过合规审查'],
        ['预警响应时间', '< 30秒', '从风险触发到预警推送'],
        ['报告生成时间', '< 3分钟', '完整投研报告自动生成'],
    ]
    make_table(doc, ['指标', '目标值', '说明'], tech_rows, col_widths=[4, 3, 9])

    doc.add_page_break()

    # ============================================================
    # 四、市场与商业
    # ============================================================
    add_heading1(doc, '四、市场与商业')

    add_heading2(doc, '4.1 市场规模')
    add_body(doc, '中国金融科技AI应用市场规模达500亿元（TAM），其中AI+投研/风控/合规细分市场50亿元（SAM），Agentic AI金融平台市场5亿元（SOM）。市场年复合增长率超过35%，先发优势明显。')

    add_heading2(doc, '4.2 商业模式')
    add_body(doc, '采用SaaS三层定价模式：基础版（299元/月，个人投资者）、专业版（999元/月，小型工作室）、企业版（5万起/年，金融机构）。预计第三年ARR达1.2亿元，毛利率82%。')

    add_heading2(doc, '4.3 竞争分析')
    comp_rows = [
        ['专业深度', '高', '低', '高'],
        ['智能程度', '低', '中', '高'],
        ['自主性', '无', '低', '高'],
        ['合规能力', '中', '无', '高'],
        ['价格', '高', '低', '中'],
    ]
    make_table(doc, ['维度', '传统金融终端', '通用AI工具', 'FinAgent Pro'], comp_rows, col_widths=[3.5, 4, 4, 4.5])
    add_body(doc, 'FinAgent Pro填补了传统金融终端与通用AI工具之间的市场空白，具备差异化竞争优势。')

    doc.add_page_break()

    # ============================================================
    # 五、团队情况（核心更新区域）
    # ============================================================
    add_heading1(doc, '五、团队情况')

    add_heading2(doc, '5.1 创始人介绍')
    founder_rows = [
        ['姓名', '冯亦根'],
        ['角色', '创始人 / 全栈开发 / 架构设计'],
        ['学校', '浙江大学'],
        ['专业', '计算机通信工程（本科）'],
        ['团队模式', 'OPC一人团队（AFAC大赛支持）'],
    ]
    make_table(doc, ['项目', '内容'], founder_rows, col_widths=[4, 12])

    doc.add_paragraph()  # 间距

    add_body(doc, '冯亦根，浙江大学计算机通信工程专业本科生，兼具AI与通信工程复合背景，独立完成FinAgent Pro从架构设计到代码实现的全流程开发。作为OPC（One Person Company）一人团队模式的践行者，他以一己之力构建了包含6大专业智能体、ReAct推理引擎、三层记忆系统和完整前端界面的金融自主智能体平台，充分展现了AI时代个人创造力的无限可能。')

    add_heading2(doc, '5.2 一人团队的力量——OPC模式')
    add_body(doc, 'AFAC2026金融智能创新大赛支持OPC（One Person Company）一人团队参赛模式，冯亦根正是这一模式的杰出代表。在AI工具的赋能下，一个人可以完成过去需要一个团队才能完成的工作——这就是AI时代的新生产力范式。')

    add_body(doc, '「一个人，一支队，以一当十」——这不是口号，而是冯亦根用代码和实践证明的现实：')

    add_bullet(doc, '架构设计：独立完成六层技术架构设计，从基础设施层到交互层，每一层的选型和接口定义均由一人把控，确保架构一致性和系统完整性')
    add_bullet(doc, 'ReAct推理引擎：从零实现「思考-行动-观察」推理闭环，包括工具调用、多步推理、错误恢复等核心机制，代码量超过5,000行')
    add_bullet(doc, '六大智能体：市场分析师、新闻解读员、风控专员、策略顾问、交易执行员、报告撰写员——每一个智能体的Prompt工程、工具配置、协同逻辑均由一人设计实现')
    add_bullet(doc, 'Orchestrator调度系统：实现三阶段流水线（规划→执行→整合）和加权协商算法，解决多智能体冲突消解问题')
    add_bullet(doc, '三层记忆架构：短期记忆、工作记忆、长期记忆的完整实现，支持智能体上下文感知和持续学习')
    add_bullet(doc, '前端全栈开发：基于React的Dashboard、分析页、对话页等完整用户界面，实现思维链可视化和实时交互')

    add_heading2(doc, '5.3 复合背景的优势')
    add_body(doc, '浙江大学计算机通信工程专业为冯亦根提供了独特的复合能力：')
    add_bullet(doc, '计算机科学：扎实的算法、数据结构、系统设计基础，支撑复杂系统架构能力')
    add_bullet(doc, '通信工程：信号处理、信息论、网络协议等知识，为金融时序数据分析和实时系统设计提供方法论支撑')
    add_bullet(doc, 'AI工程化：将前沿AI研究（ReAct、Multi-Agent、Chain-of-Thought）转化为可落地的工程系统，实现从论文到产品的跨越')

    add_heading2(doc, '5.4 团队优势')
    add_bullet(doc, '极致执行力：一人决策、一人实现，零沟通成本，从想法到代码的转化效率极高')
    add_bullet(doc, '架构一致性：全栈一人把控，不存在多人协作的接口对齐问题，系统设计高度一致')
    add_bullet(doc, 'AI赋能倍增：借助AI辅助编程工具，个人开发效率实现10倍提升，一人团队产出媲美传统5-10人团队')
    add_bullet(doc, '快速迭代：无需团队协调，可7×24小时持续迭代，MVP开发速度远超传统团队')
    add_bullet(doc, '成本优势：一人团队的运营成本极低，使产品可以更快的速度和更低的价格触达用户')

    doc.add_page_break()

    # ============================================================
    # 六、项目进展
    # ============================================================
    add_heading1(doc, '六、项目进展')

    add_heading2(doc, '6.1 已完成工作')
    progress_rows = [
        ['1', '技术架构设计', '✅ 已完成', '六层架构设计，技术选型确定'],
        ['2', 'ReAct推理引擎', '✅ 已完成', '核心推理循环实现，支持多步推理'],
        ['3', '六大智能体原型', '✅ 已完成', '基础功能实现，可独立运行'],
        ['4', 'Orchestrator调度', '✅ 已完成', '三阶段流水线，加权协商算法'],
        ['5', '前端界面', '🔄 进行中', 'Dashboard、分析页、对话页已完成'],
        ['6', '合规引擎', '🔄 进行中', '核心规则已实现，持续扩充中'],
    ]
    make_table(doc, ['序号', '工作内容', '完成状态', '说明'], progress_rows, col_widths=[1.5, 4, 3, 7.5])

    add_heading2(doc, '6.2 后续计划')
    plan_rows = [
        ['MVP完善', '2026年7-8月', '功能完善、性能优化、内测准备'],
        ['内测上线', '2026年9月', '邀请50名种子用户内测，收集反馈'],
        ['公测发布', '2026年11月', '开放注册，基础版和专业版上线'],
        ['企业版发布', '2027年3月', '企业版功能开发完成，启动销售'],
    ]
    make_table(doc, ['阶段', '时间', '计划内容'], plan_rows, col_widths=[3, 3.5, 9.5])

    add_heading2(doc, '6.3 里程碑')
    milestone_rows = [
        ['MVP上线', '2026年9月', '核心功能可用，100+注册用户'],
        ['首批付费', '2026年11月', '10+付费用户，ARR突破50万'],
        ['企业版发布', '2027年3月', '5+企业客户，ARR突破500万'],
        ['规模化增长', '2027年Q4', '200+付费用户，ARR突破1000万'],
    ]
    make_table(doc, ['里程碑', '预计时间', '关键指标'], milestone_rows, col_widths=[3, 3.5, 9.5])

    doc.add_page_break()

    # ============================================================
    # 七、融资需求
    # ============================================================
    add_heading1(doc, '七、融资需求')

    add_heading2(doc, '7.1 融资金额')
    add_body(doc, '天使轮融资500万元人民币，出让股权10%-15%，投前估值3,300万-5,000万元。')

    add_heading2(doc, '7.2 资金用途')
    fund_rows = [
        ['产品研发', '250', '50%'],
        ['团队建设', '100', '20%'],
        ['市场推广', '80', '16%'],
        ['数据采购', '40', '8%'],
        ['运营储备', '30', '6%'],
    ]
    make_table(doc, ['用途', '金额（万元）', '占比'], fund_rows, col_widths=[5, 5, 6])

    add_heading2(doc, '7.3 预期回报')
    add_body(doc, '预计3年内ARR达1.2亿元，净利润4,000万元。投资回报预期：3-5倍回报，退出方式包括IPO上市、战略并购或股权回购。')

    doc.add_page_break()

    # ============================================================
    # 八、社会价值
    # ============================================================
    add_heading1(doc, '八、社会价值')

    add_heading2(doc, '8.1 响应国家战略')
    add_body(doc, 'FinAgent Pro积极响应中央金融工作会议「五篇大文章」战略：科技金融——以AI技术推动金融科技自主创新；绿色金融——ESG智能评估助力绿色投资；普惠金融——降低投研门槛服务中小投资者；养老金融——稳健策略服务养老资金；数字金融——推动金融行业数字化转型。')

    add_heading2(doc, '8.2 促进就业')
    add_bullet(doc, '直接就业：项目发展后团队规模逐步扩展，创造高质量AI+金融就业岗位')
    add_bullet(doc, '间接就业：通过生态建设，带动数据标注、模型训练、客户服务等上下游就业')
    add_bullet(doc, '技能升级：帮助金融从业者掌握AI工具，提升职业技能和竞争力')
    add_bullet(doc, 'OPC示范：为AI时代个人创业者提供可复制的成功范式，激发全民创新活力')

    add_heading2(doc, '8.3 推动行业进步')
    add_bullet(doc, '提升效率：投研效率提升10倍，推动行业生产力变革')
    add_bullet(doc, '降低风险：实时风控预警，减少金融风险事件')
    add_bullet(doc, '促进合规：合规引擎内嵌，提升行业合规水平')
    add_bullet(doc, '技术示范：Agentic AI在金融领域的标杆应用，引领行业技术方向')

    doc.add_page_break()

    # ============================================================
    # 九、知识产权
    # ============================================================
    add_heading1(doc, '九、知识产权')

    add_heading2(doc, '9.1 已申请/计划申请')
    ip_rows = [
        ['软件著作权', 'FinAgent Pro金融自主智能体平台', '申请中'],
        ['发明专利', '基于ReAct的多智能体金融分析方法', '计划申请'],
        ['发明专利', '金融合规引擎自动审查方法', '计划申请'],
        ['发明专利', '多智能体加权协商决策方法', '计划申请'],
    ]
    make_table(doc, ['类型', '名称', '状态'], ip_rows, col_widths=[3.5, 8, 4.5])

    add_heading2(doc, '9.2 开源策略')
    add_bullet(doc, '核心引擎闭源：ReAct推理引擎、合规引擎等核心组件闭源，保持技术壁垒')
    add_bullet(doc, '工具层开源：通用工具组件（技术分析、NLP等）开源，建设开发者生态')
    add_bullet(doc, 'SDK开放：提供Agent SDK，支持第三方开发自定义智能体')
    add_bullet(doc, '数据接口开放：提供标准化数据API，方便第三方集成')

    doc.add_page_break()

    # ============================================================
    # 十、承诺与声明
    # ============================================================
    add_heading1(doc, '十、承诺与声明')

    add_body(doc, '本人冯亦根作为智融先锋团队（OPC一人团队）负责人，郑重承诺：')
    add_bullet(doc, '一、本项目为原创项目，不存在抄袭、剽窃等知识产权纠纷', bold=True)
    add_bullet(doc, '二、本项目所提交的所有材料真实、准确、完整', bold=True)
    add_bullet(doc, '三、本人将严格遵守大赛规则，公平竞争', bold=True)
    add_bullet(doc, '四、本项目如获资助，将严格按照计划使用资金，接受监督', bold=True)
    add_bullet(doc, '五、本人将积极推动项目落地，实现社会价值和商业价值', bold=True)
    add_bullet(doc, '六、本项目所有AI生成内容均经过人工审核，确保合规性', bold=True)
    add_bullet(doc, '七、本项目不提供直接投资建议，所有分析结果仅供参考', bold=True)

    # 签名区
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, '项目负责人：冯亦根', font_size=FOOTER_SIZE, color=COLOR_BLACK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, '团队名称：智融先锋', font_size=FOOTER_SIZE, color=COLOR_BLACK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, '日期：2026年6月', font_size=FOOTER_SIZE, color=COLOR_BLACK)

    # ============================================================
    # 保存文档
    # ============================================================
    doc.save(OUTPUT_PATH)
    print(f'文档已生成：{OUTPUT_PATH}')


if __name__ == '__main__':
    build_document()

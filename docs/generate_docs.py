# -*- coding: utf-8 -*-
"""
FinAgent Pro - AFAC2026金融智能创新大赛 文档生成脚本
生成三个专业Word文档：商业计划书、技术方案、项目申报书
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# ============================================================
# 通用样式与工具函数
# ============================================================

DARK_BLUE = RGBColor(0x1A, 0x3C, 0x6E)
MEDIUM_BLUE = RGBColor(0x2B, 0x57, 0x9A)
LIGHT_BLUE = RGBColor(0x3A, 0x7C, 0xBD)
ACCENT_BLUE = RGBColor(0x00, 0x6B, 0xD6)
HEADER_BG = "1A3C6E"
ROW_ALT_BG = "E8F0FE"
WHITE = "FFFFFF"
GRAY_BG = "F5F5F5"
CODE_BG = "F0F0F0"


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_text(cell, text, bold=False, font_size=10, color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT, font_name="微软雅黑"):
    """设置单元格文本"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def add_header_row(table, headers):
    """添加表头行"""
    row = table.rows[0]
    for i, header in enumerate(headers):
        set_cell_shading(row.cells[i], HEADER_BG)
        set_cell_text(row.cells[i], header, bold=True, font_size=10, color=RGBColor(0xFF, 0xFF, 0xFF), alignment=WD_ALIGN_PARAGRAPH.CENTER)


def add_data_row(table, row_idx, data, bold_first=False):
    """添加数据行"""
    row = table.rows[row_idx]
    if row_idx % 2 == 0:
        for cell in row.cells:
            set_cell_shading(cell, ROW_ALT_BG)
    for i, text in enumerate(data):
        is_bold = bold_first and i == 0
        set_cell_text(row.cells[i], text, bold=is_bold, font_size=9.5)


def create_table(doc, headers, data_rows, col_widths=None):
    """创建格式化表格"""
    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, headers)
    for idx, row_data in enumerate(data_rows):
        add_data_row(table, idx + 1, row_data, bold_first=True)
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)
    return table


def add_code_block(doc, code_text):
    """添加代码块（灰色背景、Courier New字体）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(14)
    # 添加灰色背景
    pPr = p._element.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CODE_BG}" w:val="clear"/>')
    pPr.append(shading)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def add_bullet(doc, text, level=0, bold_prefix=""):
    """添加项目符号列表"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(10.5)
        run_b.font.name = "微软雅黑"
        run_b._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
        run_n = p.add_run(text)
        run_n.font.size = Pt(10.5)
        run_n.font.name = "微软雅黑"
        run_n._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    return p


def add_body(doc, text, bold=False, font_size=10.5):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    run.bold = bold
    return p


def add_heading1(doc, text):
    """添加一级标题"""
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = DARK_BLUE
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
        run.font.size = Pt(18)
    return h


def add_heading2(doc, text):
    """添加二级标题"""
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = MEDIUM_BLUE
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
        run.font.size = Pt(14)
    return h


def add_heading3(doc, text):
    """添加三级标题"""
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.color.rgb = LIGHT_BLUE
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
        run.font.size = Pt(12)
    return h


def setup_document(doc, title_text):
    """配置文档基本设置"""
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    # 页眉
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run(title_text)
    hr.font.size = Pt(8)
    hr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    hr.font.name = "微软雅黑"
    hr._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    # 页脚页码
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 添加页码字段
    run = fp.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)
    run2 = fp.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._element.append(instrText)
    run3 = fp.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fldChar2)
    for r in [run, run2, run3]:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def add_cover_page(doc, title, subtitle, extra_lines):
    """添加封面页"""
    # 添加空行到页面中部
    for _ in range(6):
        doc.add_paragraph()
    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.color.rgb = DARK_BLUE
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    run.bold = True
    # 副标题
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(18)
    run2.font.color.rgb = MEDIUM_BLUE
    run2.font.name = "微软雅黑"
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    # 分隔线
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("━" * 30)
    run3.font.color.rgb = ACCENT_BLUE
    run3.font.size = Pt(14)
    # 额外行
    for line in extra_lines:
        p4 = doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run4 = p4.add_run(line)
        run4.font.size = Pt(14)
        run4.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run4.font.name = "微软雅黑"
        run4._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    doc.add_page_break()


def add_toc(doc):
    """添加目录页"""
    add_heading1(doc, "目  录")
    doc.add_paragraph()
    # 目录项手动添加
    return doc


def add_toc_item(doc, text, indent=0):
    """添加目录项"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent * 1.0)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    if indent == 0:
        run.bold = True
        run.font.color.rgb = DARK_BLUE
    else:
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


# ============================================================
# 文档1：商业计划书
# ============================================================

def create_business_plan():
    doc = Document()
    setup_document(doc, "FinAgent Pro 商业计划书")

    # 封面
    add_cover_page(doc,
        "FinAgent Pro",
        "金融自主智能体平台 — 商业计划书",
        ["AFAC2026金融智能创新大赛", "2026年6月"])

    # 目录
    add_toc(doc)
    toc_items = [
        ("第一章  执行摘要", 0),
        ("第二章  项目概述", 0),
        ("  2.1 项目背景", 1),
        ("  2.2 产品定位", 1),
        ("  2.3 核心创新", 1),
        ("第三章  市场分析", 0),
        ("  3.1 市场规模", 1),
        ("  3.2 目标客户", 1),
        ("  3.3 竞争格局", 1),
        ("第四章  产品与服务", 0),
        ("第五章  商业模式", 0),
        ("第六章  技术方案", 0),
        ("第七章  发展规划", 0),
        ("第八章  融资计划", 0),
        ("第九章  风险分析", 0),
        ("第十章  社会价值", 0),
    ]
    for text, indent in toc_items:
        add_toc_item(doc, text, indent)
    doc.add_page_break()

    # 第一章：执行摘要
    add_heading1(doc, "第一章  执行摘要")
    add_body(doc, "FinAgent Pro 是一款面向金融行业的自主智能体（Agentic AI）平台，通过六大专业数字员工——市场分析师、新闻解读员、风控专员、策略顾问、交易执行员和报告撰写员——为金融机构提供7×24小时全链路智能投研服务。")
    add_body(doc, "")
    add_heading2(doc, "1.1 项目一句话定位")
    add_body(doc, "FinAgent Pro = 金融行业专属的自主智能体平台，让每个金融从业者拥有6位AI数字员工。")
    add_heading2(doc, "1.2 核心价值")
    add_bullet(doc, "效率提升：投研效率提升10倍，报告生成从4小时缩短至5分钟", bold_prefix="")
    add_bullet(doc, "成本节约：人力成本降低60%，合规审查成本降低80%", bold_prefix="")
    add_bullet(doc, "风险可控：实时风控预警，合规审计全程留痕", bold_prefix="")
    add_bullet(doc, "决策赋能：多智能体协同推理，提供深度洞察而非简单信息聚合", bold_prefix="")
    add_heading2(doc, "1.3 市场机会")
    add_body(doc, "中国金融科技市场规模2025年预计突破500亿元，其中AI+金融细分赛道年复合增长率超过35%。随着大模型技术成熟和金融行业数字化转型加速，Agentic AI在金融领域的应用正处于爆发前夜。")
    add_heading2(doc, "1.4 融资需求")
    add_body(doc, "天使轮融资500万元人民币，用于产品研发、团队扩充和市场验证，预计12个月内实现MVP上线并获取首批付费客户。")

    doc.add_page_break()

    # 第二章：项目概述
    add_heading1(doc, "第二章  项目概述")
    add_heading2(doc, "2.1 项目背景——金融行业痛点")
    add_body(doc, "当前金融行业面临四大核心痛点：")
    add_bullet(doc, "信息过载：每日数万条金融资讯，分析师无法有效筛选和消化", bold_prefix="信息过载：")
    add_bullet(doc, "人力依赖：投研报告高度依赖人工撰写，效率低下且质量参差不齐", bold_prefix="人力依赖：")
    add_bullet(doc, "风控滞后：传统风控以事后审查为主，缺乏实时预警能力", bold_prefix="风控滞后：")
    add_bullet(doc, "合规压力：监管要求日趋严格，合规审查工作量剧增", bold_prefix="合规压力：")

    add_heading2(doc, "2.2 产品定位——六大数字员工")
    create_table(doc,
        ["数字员工", "核心职责", "关键能力"],
        [
            ["市场分析师\n(Market Analyst)", "实时市场监测与数据分析", "行情追踪、技术指标计算、趋势识别"],
            ["新闻解读员\n(News Interpreter)", "金融资讯智能解读与情绪分析", "NLP解读、事件提取、情绪量化"],
            ["风控专员\n(Risk Controller)", "实时风险评估与预警", "VaR计算、压力测试、阈值预警"],
            ["策略顾问\n(Strategy Advisor)", "投资策略生成与回测", "策略构建、历史回测、收益归因"],
            ["交易执行员\n(Execution Agent)", "交易信号生成与执行建议", "信号生成、执行优化、滑点控制"],
            ["报告撰写员\n(Report Writer)", "专业投研报告自动生成", "报告模板、数据可视化、合规审查"],
        ],
        col_widths=[4.5, 5.0, 5.5])

    add_heading2(doc, "2.3 核心创新——六大差异化优势")
    create_table(doc,
        ["创新点", "传统方案", "FinAgent Pro"],
        [
            ["自主推理", "规则驱动，被动响应", "ReAct推理引擎，主动思考与行动"],
            ["多智能体协同", "单模型单任务", "6大智能体分工协作，加权协商"],
            ["主动智能", "用户触发式", "定时巡检+事件驱动+阈值预警"],
            ["合规内嵌", "事后审查", "合规引擎全程嵌入，审计留痕"],
            ["记忆系统", "无状态交互", "三层记忆架构，持续学习优化"],
            ["可解释性", "黑箱输出", "思维链可视化，决策过程透明"],
        ],
        col_widths=[3.5, 4.5, 6.0])

    doc.add_page_break()

    # 第三章：市场分析
    add_heading1(doc, "第三章  市场分析")
    add_heading2(doc, "3.1 市场规模")
    create_table(doc,
        ["层级", "规模", "说明"],
        [
            ["TAM（总可达市场）", "500亿元", "中国金融科技AI应用市场"],
            ["SAM（可服务市场）", "50亿元", "AI+投研/风控/合规细分市场"],
            ["SOM（可获得市场）", "5亿元", "Agentic AI金融平台市场"],
        ],
        col_widths=[4.0, 3.0, 7.0])
    add_body(doc, "")
    add_body(doc, "根据艾瑞咨询和IDC数据，中国金融科技市场持续高速增长，AI在金融领域的渗透率从2022年的12%提升至2025年的28%，预计2027年将达到45%。Agentic AI作为新一代AI范式，正处于市场导入期，先发优势明显。")

    add_heading2(doc, "3.2 目标客户——四类客户画像")
    create_table(doc,
        ["客户类型", "规模", "核心需求", "付费能力", "获客策略"],
        [
            ["公募/私募基金", "中大型", "投研效率、风控合规", "年费20-100万", "行业峰会+标杆案例"],
            ["券商研究所", "中大型", "报告自动化、数据整合", "年费30-80万", "POC验证+渠道合作"],
            ["银行理财子公司", "大型", "合规审查、风险评估", "年费50-200万", "招标+合规驱动"],
            ["个人投资者/工作室", "小型", "智能投研、策略回测", "月费299-999", "社区运营+免费版引流"],
        ],
        col_widths=[3.0, 2.0, 3.5, 3.0, 3.5])

    add_heading2(doc, "3.3 竞争格局")
    create_table(doc,
        ["维度", "Wind/同花顺", "通用AI（ChatGPT等）", "FinAgent Pro"],
        [
            ["专业深度", "★★★★★", "★★☆☆☆", "★★★★☆"],
            ["智能程度", "★★☆☆☆", "★★★★☆", "★★★★★"],
            ["自主性", "★☆☆☆☆", "★★☆☆☆", "★★★★★"],
            ["合规能力", "★★★★☆", "★☆☆☆☆", "★★★★★"],
            ["价格", "高（年费10万+）", "低（月费$20）", "中（年费5万起）"],
            ["可解释性", "★★★☆☆", "★★☆☆☆", "★★★★★"],
        ],
        col_widths=[2.5, 3.5, 3.5, 3.5])
    add_body(doc, "")
    add_body(doc, "FinAgent Pro 的差异化定位：既具备金融专业深度，又拥有AI原生智能，填补了传统金融终端与通用AI工具之间的市场空白。")

    doc.add_page_break()

    # 第四章：产品与服务
    add_heading1(doc, "第四章  产品与服务")
    add_heading2(doc, "4.1 产品架构——六层架构")
    create_table(doc,
        ["层级", "名称", "核心组件", "功能说明"],
        [
            ["L1", "数据层", "行情API、新闻源、公告库", "多源金融数据实时采集与清洗"],
            ["L2", "工具层", "技术分析、NLP、风控模型", "可插拔工具注册与发现"],
            ["L3", "引擎层", "ReAct引擎、Memory、合规", "推理循环、记忆管理、合规审查"],
            ["L4", "智能体层", "6大专业智能体", "分工协作、自主推理、主动行动"],
            ["L5", "协同层", "Orchestrator、协商算法", "任务调度、结果聚合、冲突消解"],
            ["L6", "交互层", "Web UI、API、WebSocket", "用户交互、实时推送、开放集成"],
        ],
        col_widths=[1.5, 2.5, 4.0, 5.0])

    add_heading2(doc, "4.2 六大智能体详解")
    add_heading3(doc, "4.2.1 市场分析师（Market Analyst Agent）")
    add_body(doc, "市场分析师是FinAgent Pro的「数据哨兵」，负责实时监测市场行情数据，识别趋势变化和异常波动。其核心能力包括：实时行情追踪（支持A股、港股、美股多市场）、技术指标自动计算（MACD、KDJ、RSI等20+指标）、趋势形态识别（头肩顶、双底等经典形态）。")
    add_body(doc, "市场分析师通过ReAct推理引擎，能够自主决定何时需要调取哪些数据、计算哪些指标，而非被动等待用户指令。例如，当检测到某只股票成交量突然放大3倍时，会主动触发深度分析流程。")

    add_heading3(doc, "4.2.2 新闻解读员（News Interpreter Agent）")
    add_body(doc, "新闻解读员是FinAgent Pro的「信息过滤器」，负责从海量金融资讯中提取关键信息，进行事件分类和情绪量化。其核心能力包括：金融实体识别（公司、人物、产品）、事件类型分类（政策、业绩、并购等8大类）、情绪强度量化（-1到+1连续值）、关联影响分析（事件对相关标的的影响路径）。")

    add_heading3(doc, "4.2.3 风控专员（Risk Controller Agent）")
    add_body(doc, "风控专员是FinAgent Pro的「安全卫士」，负责实时风险评估和预警。其核心能力包括：VaR（在险价值）实时计算、压力测试场景模拟、多维度风险阈值监控、合规规则自动审查。风控专员7×24小时不间断运行，一旦检测到风险指标突破阈值，立即触发预警信号并通知相关智能体协同应对。")

    add_heading3(doc, "4.2.4 策略顾问（Strategy Advisor Agent）")
    add_body(doc, "策略顾问是FinAgent Pro的「智囊团」，负责投资策略的生成、回测和优化。其核心能力包括：多因子策略构建、历史数据回测验证、收益归因分析、策略参数自适应优化。策略顾问能够根据市场分析师和新闻解读员提供的信号，动态调整策略参数，实现策略与市场环境的实时适配。")

    add_heading3(doc, "4.2.5 交易执行员（Execution Agent）")
    add_body(doc, "交易执行员是FinAgent Pro的「行动派」，负责将策略信号转化为可执行的交易建议。其核心能力包括：交易信号生成与验证、执行时机优化、滑点与冲击成本估算、交易合规性检查。交易执行员在生成任何交易建议前，必须通过风控专员的风险审查和合规引擎的合规检查，确保交易行为安全合规。")

    add_heading3(doc, "4.2.6 报告撰写员（Report Writer Agent）")
    add_body(doc, "报告撰写员是FinAgent Pro的「表达者」，负责将各智能体的分析结果整合为专业投研报告。其核心能力包括：多源信息整合与结构化、专业报告模板匹配、数据可视化图表生成、合规表述自动审查。报告撰写员能够根据用户需求，生成日报、周报、深度分析报告等多种格式，并自动进行合规审查。")

    add_heading2(doc, "4.3 核心功能")
    add_heading3(doc, "4.3.1 智能分析")
    add_body(doc, "用户输入任意股票代码或金融问题，系统自动调度相关智能体进行多维度分析，包括基本面、技术面、消息面、资金面等，并给出综合评估。")
    add_heading3(doc, "4.3.2 对话式交互")
    add_body(doc, "支持自然语言对话，用户可以像与投研团队交流一样，追问细节、调整分析维度、要求深度解读。系统基于ReAct推理引擎，能够理解复杂的多轮对话上下文。")
    add_heading3(doc, "4.3.3 主动预警")
    add_body(doc, "系统无需用户主动查询，通过定时巡检、事件驱动和阈值触发三种机制，主动推送风险预警、机会提醒和异常报告。")
    add_heading3(doc, "4.3.4 自动报告")
    add_body(doc, "一键生成专业投研报告，支持日报、周报、月报和深度分析报告四种模板，自动整合数据、图表和分析结论，并内置合规审查。")

    doc.add_page_break()

    # 第五章：商业模式
    add_heading1(doc, "第五章  商业模式")
    add_heading2(doc, "5.1 盈利模式——SaaS三层定价")
    create_table(doc,
        ["方案", "月费", "年费", "目标客户", "核心功能"],
        [
            ["基础版", "299元/月", "2,999元/年", "个人投资者", "2个智能体+基础分析+日报"],
            ["专业版", "999元/月", "9,999元/年", "小型工作室", "4个智能体+深度分析+预警"],
            ["企业版", "按需定制", "5万起/年", "金融机构", "6个智能体+全部功能+私有部署"],
        ],
        col_widths=[2.5, 2.5, 2.5, 3.0, 4.5])

    add_heading2(doc, "5.2 收入预测——三年预测")
    create_table(doc,
        ["指标", "第一年", "第二年", "第三年"],
        [
            ["付费用户数", "200", "1,500", "5,000"],
            ["企业客户数", "5", "30", "100"],
            ["ARR（年经常性收入）", "300万", "2,500万", "1.2亿"],
            ["毛利率", "65%", "75%", "82%"],
            ["净利润", "-200万", "500万", "4,000万"],
        ],
        col_widths=[4.0, 3.0, 3.0, 3.0])

    add_heading2(doc, "5.3 成本结构")
    create_table(doc,
        ["成本项", "占比", "说明"],
        [
            ["研发人力", "45%", "核心团队+外部专家"],
            ["算力与API", "25%", "LLM API调用+云服务器"],
            ["数据采购", "15%", "行情数据+新闻源授权"],
            ["市场销售", "10%", "品牌推广+渠道建设"],
            ["运营管理", "5%", "办公+行政+法务"],
        ],
        col_widths=[3.0, 2.0, 9.0])

    doc.add_page_break()

    # 第六章：技术方案
    add_heading1(doc, "第六章  技术方案")
    add_heading2(doc, "6.1 技术选型")
    create_table(doc,
        ["模块", "技术栈", "选型理由"],
        [
            ["前端", "Vue 3 + Vite + TailwindCSS", "轻量高效、组件化开发、快速迭代"],
            ["后端", "FastAPI + Python 3.11+", "异步高性能、类型安全、生态丰富"],
            ["LLM", "GPT-4o / DeepSeek / Qwen", "多模型支持、成本优化、国产替代"],
            ["向量数据库", "ChromaDB", "轻量嵌入、本地部署、金融数据安全"],
            ["任务调度", "APScheduler", "轻量级、Cron表达式、定时巡检"],
            ["实时通信", "WebSocket", "双向推送、预警实时送达"],
            ["部署", "Docker + Vercel", "容器化部署、前后端分离、弹性扩展"],
        ],
        col_widths=[2.5, 4.5, 6.0])

    add_heading2(doc, "6.2 技术壁垒")
    add_heading3(doc, "6.2.1 ReAct推理引擎")
    add_body(doc, "自研ReAct（Reasoning + Acting）推理引擎，实现「思考-行动-观察」的闭环推理。智能体不再是简单的问答机器，而是能够自主规划任务步骤、调用工具、评估结果、调整策略的自主决策系统。")
    add_code_block(doc, """# ReAct推理循环核心逻辑
while not task_complete:
    thought = llm.reason(context, observation)    # 推理
    action = select_tool(thought)                  # 选择工具
    observation = execute(action)                  # 执行观察
    context = memory.update(observation)           # 更新记忆
    task_complete = evaluate(context)              # 评估完成度""")

    add_heading3(doc, "6.2.2 三层记忆架构")
    add_body(doc, "短期记忆（当前会话上下文）、工作记忆（近期交互摘要）、长期记忆（知识库与经验库）三层架构，使智能体具备持续学习和上下文关联能力，越用越智能。")

    add_heading3(doc, "6.2.3 合规引擎")
    add_body(doc, "内置金融合规规则引擎，覆盖证监会、银保监会等监管机构的200+条规则，所有智能体输出均经过合规审查，确保内容合规、表述规范、数据脱敏。")

    add_heading2(doc, "6.3 数据安全")
    add_bullet(doc, "传输加密：全链路TLS/SSL加密，API通信采用HTTPS", bold_prefix="")
    add_bullet(doc, "存储加密：敏感数据AES-256加密存储，密钥分级管理", bold_prefix="")
    add_bullet(doc, "访问控制：RBAC权限模型，最小权限原则", bold_prefix="")
    add_bullet(doc, "审计日志：全操作留痕，支持合规审计追溯", bold_prefix="")
    add_bullet(doc, "数据隔离：多租户数据逻辑隔离，企业版支持物理隔离", bold_prefix="")

    doc.add_page_break()

    # 第七章：发展规划
    add_heading1(doc, "第七章  发展规划")
    add_heading2(doc, "7.1 发展路线图")
    create_table(doc,
        ["阶段", "时间", "目标", "关键指标"],
        [
            ["MVP阶段", "2026 Q3-Q4", "核心功能上线，获取种子用户", "100+注册用户，5+付费用户"],
            ["成长阶段", "2027 Q1-Q2", "功能完善，市场拓展", "1000+注册用户，50+付费用户"],
            ["成熟阶段", "2027 Q3-Q4", "生态建设，规模化增长", "5000+注册用户，200+付费用户"],
        ],
        col_widths=[2.5, 3.0, 4.0, 4.5])

    add_heading2(doc, "7.2 里程碑")
    create_table(doc,
        ["里程碑", "预计时间", "状态"],
        [
            ["产品原型完成", "2026年7月", "进行中"],
            ["MVP内测上线", "2026年9月", "计划中"],
            ["首批付费客户", "2026年11月", "计划中"],
            ["企业版发布", "2027年3月", "计划中"],
            ["ARR突破1000万", "2027年9月", "计划中"],
            ["Pre-A轮融资", "2027年12月", "计划中"],
        ],
        col_widths=[5.0, 4.0, 3.0])

    doc.add_page_break()

    # 第八章：融资计划
    add_heading1(doc, "第八章  融资计划")
    add_heading2(doc, "8.1 融资需求——天使轮500万")
    add_body(doc, "本次天使轮融资目标500万元人民币，出让股权10%-15%，投前估值3,300万-5,000万元。资金将主要用于产品研发、团队建设和市场验证。")

    add_heading2(doc, "8.2 资金用途")
    create_table(doc,
        ["用途", "金额（万元）", "占比", "说明"],
        [
            ["产品研发", "250", "50%", "核心功能开发、智能体优化、合规引擎"],
            ["团队建设", "100", "20%", "招聘AI工程师、金融专家、产品经理"],
            ["市场推广", "80", "16%", "品牌建设、行业峰会、POC验证"],
            ["数据采购", "40", "8%", "行情数据、新闻源、知识库"],
            ["运营储备", "30", "6%", "日常运营、法务财务、应急储备"],
        ],
        col_widths=[2.5, 3.0, 2.0, 6.5])

    add_heading2(doc, "8.3 退出机制")
    add_bullet(doc, "IPO上市：预计2029-2030年科创板或港股上市", bold_prefix="")
    add_bullet(doc, "战略并购：被大型金融机构或金融科技公司收购", bold_prefix="")
    add_bullet(doc, "股权回购：若未实现上市/并购，创始团队按年化8%回购", bold_prefix="")

    doc.add_page_break()

    # 第九章：风险分析
    add_heading1(doc, "第九章  风险分析")
    add_heading2(doc, "9.1 技术风险")
    create_table(doc,
        ["风险", "概率", "影响", "应对策略"],
        [
            ["LLM幻觉问题", "中", "高", "多模型交叉验证+人工审核机制"],
            ["系统稳定性", "低", "高", "容灾备份+降级策略+监控告警"],
            ["数据质量", "中", "中", "多源交叉验证+异常数据过滤"],
        ],
        col_widths=[3.0, 2.0, 2.0, 6.0])

    add_heading2(doc, "9.2 市场风险")
    create_table(doc,
        ["风险", "概率", "影响", "应对策略"],
        [
            ["市场接受度低", "中", "高", "免费版引流+POC验证+标杆案例"],
            ["竞争加剧", "高", "中", "技术壁垒+合规优势+生态建设"],
            ["客户流失", "中", "中", "持续优化+客户成功+粘性功能"],
        ],
        col_widths=[3.0, 2.0, 2.0, 6.0])

    add_heading2(doc, "9.3 合规风险")
    create_table(doc,
        ["风险", "概率", "影响", "应对策略"],
        [
            ["监管政策变化", "中", "高", "合规团队+政策跟踪+快速响应"],
            ["数据合规", "低", "高", "数据脱敏+隐私计算+合规审查"],
            ["投资建议合规", "中", "高", "免责声明+合规引擎+人工复核"],
        ],
        col_widths=[3.0, 2.0, 2.0, 6.0])

    add_heading2(doc, "9.4 应对策略总览")
    add_body(doc, "FinAgent Pro采取「技术+合规+运营」三位一体的风险应对策略：技术上，通过多模型验证、容灾备份确保系统可靠性；合规上，内置合规引擎、全程审计留痕确保合规性；运营上，通过客户成功体系、持续迭代确保市场竞争力。")

    doc.add_page_break()

    # 第十章：社会价值
    add_heading1(doc, "第十章  社会价值")
    add_heading2(doc, "10.1 响应「五篇大文章」战略")
    add_body(doc, "FinAgent Pro积极响应中央金融工作会议提出的「五篇大文章」战略方向：")
    create_table(doc,
        ["战略方向", "FinAgent Pro贡献"],
        [
            ["科技金融", "以AI技术赋能金融行业，推动金融科技自主创新"],
            ["绿色金融", "ESG智能评估模块，助力绿色投资决策"],
            ["普惠金融", "降低投研门槛，让中小投资者享受专业级服务"],
            ["养老金融", "稳健策略推荐，服务养老资金保值增值"],
            ["数字金融", "推动金融行业数字化转型，提升行业效率"],
        ],
        col_widths=[3.0, 10.0])

    add_heading2(doc, "10.2 促进金融行业进步")
    add_body(doc, "FinAgent Pro通过AI技术赋能金融行业，不仅提升了单个机构的效率，更推动了整个行业的智能化升级。我们的愿景是让每一位金融从业者都能拥有AI助手，让金融决策更加科学、高效、合规。")
    add_body(doc, "")
    add_body(doc, "FinAgent Pro —— 让金融智能触手可及。")

    # 保存
    output_path = r"D:\AFAC2026金融智能创新大赛\FinAgent-Pro\docs\FinAgentPro商业计划书.docx"
    doc.save(output_path)
    print(f"✅ 商业计划书已生成: {output_path}")
    return output_path


# ============================================================
# 文档2：技术方案
# ============================================================

def create_technical_spec():
    doc = Document()
    setup_document(doc, "FinAgent Pro 技术方案")

    # 封面
    add_cover_page(doc,
        "FinAgent Pro",
        "金融自主智能体平台 — 技术方案",
        ["AFAC2026金融智能创新大赛", "2026年6月"])

    # 目录
    add_toc(doc)
    toc_items = [
        ("第一章  技术架构", 0),
        ("  1.1 整体架构", 1),
        ("  1.2 架构图", 1),
        ("  1.3 技术选型表", 1),
        ("第二章  核心引擎", 0),
        ("  2.1 ReAct推理引擎", 1),
        ("  2.2 Tool系统", 1),
        ("  2.3 Memory系统", 1),
        ("第三章  智能体设计", 0),
        ("第四章  多智能体协同", 0),
        ("第五章  主动智能", 0),
        ("第六章  合规引擎", 0),
        ("第七章  性能与扩展", 0),
        ("第八章  部署方案", 0),
    ]
    for text, indent in toc_items:
        add_toc_item(doc, text, indent)
    doc.add_page_break()

    # 第一章：技术架构
    add_heading1(doc, "第一章  技术架构")
    add_heading2(doc, "1.1 整体架构——六层架构")
    add_body(doc, "FinAgent Pro采用六层分层架构，各层职责明确、松耦合、可独立演进：")
    create_table(doc,
        ["层级", "名称", "核心组件", "职责"],
        [
            ["L1", "数据层\n(Data Layer)", "MarketDataService\nNewsService\nDatabaseService", "多源金融数据实时采集、清洗、标准化"],
            ["L2", "工具层\n(Tool Layer)", "TechnicalAnalysis\nSentimentAnalyzer\nRiskCalculator\nReportGenerator", "可插拔工具注册与发现，统一调用接口"],
            ["L3", "引擎层\n(Engine Layer)", "ReActEngine\nMemorySystem\nComplianceEngine", "推理循环驱动、记忆管理、合规审查"],
            ["L4", "智能体层\n(Agent Layer)", "MarketAgent\nNewsAgent\nRiskAgent\nStrategyAgent\nExecutionAgent\nReportAgent", "专业分工、自主推理、主动行动"],
            ["L5", "协同层\n(Orchestration Layer)", "Orchestrator\nNegotiationAlgorithm\nLLMEnhancer", "任务调度、结果聚合、冲突消解"],
            ["L6", "交互层\n(Interface Layer)", "FastAPI Routes\nWebSocket\nVue3 Frontend", "用户交互、实时推送、开放API"],
        ],
        col_widths=[1.5, 2.5, 4.0, 5.0])

    add_heading2(doc, "1.2 架构图")
    add_code_block(doc, """
┌─────────────────────────────────────────────────────────────┐
│                    L6 交互层 (Interface)                     │
│  ┌──────────┐  ┌──────────┐  ┌──────────┐  ┌──────────┐   │
│  │  Vue3 UI  │  │ REST API │  │WebSocket │  │  CLI工具  │   │
│  └──────────┘  └──────────┘  └──────────┘  └──────────┘   │
├─────────────────────────────────────────────────────────────┤
│                  L5 协同层 (Orchestration)                    │
│  ┌──────────────────────────────────────────────────────┐   │
│  │           Orchestrator (调度 + 协商 + 增强)            │   │
│  └──────────────────────────────────────────────────────┘   │
├─────────────────────────────────────────────────────────────┤
│                  L4 智能体层 (Agents)                         │
│  ┌──────┐ ┌──────┐ ┌──────┐ ┌──────┐ ┌──────┐ ┌──────┐   │
│  │Market│ │ News │ │ Risk │ │Strat │ │Exec  │ │Report│   │
│  └──────┘ └──────┘ └──────┘ └──────┘ └──────┘ └──────┘   │
├─────────────────────────────────────────────────────────────┤
│                  L3 引擎层 (Engine)                           │
│  ┌──────────┐  ┌──────────┐  ┌──────────────┐             │
│  │ReAct引擎  │  │Memory系统 │  │ 合规引擎      │             │
│  └──────────┘  └──────────┘  └──────────────┘             │
├─────────────────────────────────────────────────────────────┤
│                  L2 工具层 (Tools)                            │
│  ┌────┐ ┌────┐ ┌────┐ ┌────┐ ┌────┐ ┌────┐ ┌────┐      │
│  │TA  │ │NLP │ │Risk│ │Bt  │ │Sig │ │Rpt │ │... │      │
│  └────┘ └────┘ └────┘ └────┘ └────┘ └────┘ └────┘      │
├─────────────────────────────────────────────────────────────┤
│                  L1 数据层 (Data)                             │
│  ┌──────────┐  ┌──────────┐  ┌──────────┐  ┌──────────┐   │
│  │ 行情数据  │  │ 新闻资讯  │  │ 公告数据  │  │ 历史库    │   │
│  └──────────┘  └──────────┘  └──────────┘  └──────────┘   │
└─────────────────────────────────────────────────────────────┘""")

    add_heading2(doc, "1.3 技术选型表")
    create_table(doc,
        ["模块", "技术", "版本", "选型理由"],
        [
            ["前端框架", "Vue 3", "3.4+", "Composition API、响应式、生态成熟"],
            ["构建工具", "Vite", "5.x", "HMR极速、ESM原生、配置简洁"],
            ["CSS框架", "TailwindCSS", "3.x", "原子化CSS、快速开发、一致性"],
            ["后端框架", "FastAPI", "0.110+", "异步高性能、自动文档、类型安全"],
            ["运行时", "Python", "3.11+", "AI生态丰富、开发效率高"],
            ["LLM", "GPT-4o/DeepSeek/Qwen", "最新", "多模型支持、成本优化"],
            ["向量数据库", "ChromaDB", "0.4+", "轻量嵌入、本地优先、Python原生"],
            ["任务调度", "APScheduler", "3.10+", "Cron表达式、持久化、轻量级"],
            ["容器化", "Docker", "24+", "环境一致、快速部署、弹性扩展"],
            ["部署平台", "Vercel + Docker", "最新", "前端CDN + 后端容器化"],
        ],
        col_widths=[2.5, 3.5, 2.0, 5.0])

    doc.add_page_break()

    # 第二章：核心引擎
    add_heading1(doc, "第二章  核心引擎")
    add_heading2(doc, "2.1 ReAct推理引擎")
    add_body(doc, "ReAct（Reasoning + Acting）是FinAgent Pro的核心推理范式，将大语言模型的推理能力与工具调用能力深度融合，实现「思考-行动-观察」的闭环推理。")

    add_heading3(doc, "2.1.1 推理循环流程")
    add_code_block(doc, """
class ReActEngine:
    \"\"\"ReAct推理引擎核心实现\"\"\"

    async def run(self, task: str, max_steps: int = 10):
        context = self.memory.get_relevant(task)
        for step in range(max_steps):
            # 1. Thought: 推理当前状态，决定下一步
            thought = await self.llm.reason(
                task=task,
                context=context,
                available_tools=self.tool_registry.list()
            )
            self.chain.append(ThoughtNode(thought))

            # 2. Action: 选择并执行工具
            action = self.parse_action(thought)
            observation = await self.execute_tool(action)
            self.chain.append(ActionNode(action, observation))

            # 3. Observation: 评估结果
            context = self.memory.update(observation)
            if self.is_complete(task, context):
                return self.synthesize(task, context)

        return self.synthesize(task, context)  # 达到步数上限""")

    add_heading3(doc, "2.1.2 ReAct循环示意")
    add_code_block(doc, """
用户: "分析贵州茅台的投资价值"

Thought 1: 需要先获取茅台的最新行情数据
Action 1:  market_data.get_stock("600519")
Observation 1: 当前价1856.00, 涨幅+2.3%, 成交量放大1.5倍

Thought 2: 行情表现不错，需要了解近期新闻动态
Action 2:  news.search("贵州茅台", days=7)
Observation 2: 近期3条重要新闻: 1)新品发布 2)机构上调评级 3)...

Thought 3: 新闻面偏正面，需要做技术分析
Action 3:  technical_analysis.compute("600519", ["MACD","RSI","KDJ"])
Observation 3: MACD金叉, RSI=62(偏强), KDJ向上发散

Thought 4: 综合分析完成，生成投资分析报告
Action 4:  report.generate(analysis_result)
Final: 贵州茅台综合评级"增持"...""")

    add_heading2(doc, "2.2 Tool系统——注册发现模式")
    add_body(doc, "Tool系统采用注册-发现模式，所有工具通过装饰器注册到统一注册中心，智能体通过语义匹配自动发现和调用所需工具。")
    add_code_block(doc, """
# 工具注册示例
class ToolRegistry:
    _tools: Dict[str, ToolMeta] = {}

    @classmethod
    def register(cls, name: str, description: str, parameters: dict):
        \"\"\"工具注册装饰器\"\"\"
        def decorator(func):
            cls._tools[name] = ToolMeta(
                name=name,
                description=description,
                parameters=parameters,
                handler=func
            )
            return func
        return decorator

    @classmethod
    def discover(cls, query: str, top_k: int = 3) -> List[ToolMeta]:
        \"\"\"语义发现工具\"\"\"
        embeddings = cls._embed_tools()
        query_emb = embed(query)
        return cosine_similarity(query_emb, embeddings)[:top_k]

# 使用示例
@ToolRegistry.register(
    name="technical_analysis",
    description="计算技术指标：MACD、RSI、KDJ、布林带等",
    parameters={"symbol": "str", "indicators": "list"}
)
async def compute_indicators(symbol: str, indicators: list):
    ...""")

    add_heading2(doc, "2.3 Memory系统——三层架构")
    add_body(doc, "Memory系统采用三层架构，模拟人类记忆机制，使智能体具备上下文关联和持续学习能力：")
    create_table(doc,
        ["层级", "名称", "存储内容", "生命周期", "实现"],
        [
            ["L1", "短期记忆\n(Short-term)", "当前会话上下文\n最近5轮对话", "会话级", "内存字典 + 滑动窗口"],
            ["L2", "工作记忆\n(Working)", "近期交互摘要\n用户偏好设置", "7天", "ChromaDB向量检索"],
            ["L3", "长期记忆\n(Long-term)", "知识库\n历史经验\n学习模式", "永久", "ChromaDB + 持久化存储"],
        ],
        col_widths=[1.5, 2.5, 3.5, 2.0, 3.5])

    add_code_block(doc, """
class MemorySystem:
    \"\"\"三层记忆系统\"\"\"

    async def store(self, content: str, level: MemoryLevel):
        embedding = await self.embed(content)
        if level == MemoryLevel.SHORT_TERM:
            self.short_term.append({"content": content, "emb": embedding})
        elif level == MemoryLevel.WORKING:
            await self.chroma.add(embedding, metadata={"level": "working"})
        else:
            await self.chroma.add(embedding, metadata={"level": "long_term"})

    async def recall(self, query: str, top_k: int = 5) -> List[MemoryItem]:
        query_emb = await self.embed(query)
        results = []
        # 优先短期记忆
        results += self._search_short_term(query_emb, top_k)
        # 补充工作记忆
        results += await self.chroma.query(query_emb, {"level": "working"}, top_k)
        # 深度长期记忆
        results += await self.chroma.query(query_emb, {"level": "long_term"}, top_k)
        return self._deduplicate(results)[:top_k]""")

    doc.add_page_break()

    # 第三章：智能体设计
    add_heading1(doc, "第三章  智能体设计")
    add_body(doc, "FinAgent Pro设计了六大专业智能体，每个智能体拥有独立的职责、工具集和信号算法，通过ReAct推理引擎实现自主决策。")

    # 3.1 市场分析师
    add_heading2(doc, "3.1 市场分析师（Market Analyst Agent）")
    create_table(doc,
        ["属性", "详情"],
        [
            ["职责", "实时市场监测、趋势识别、异常检测"],
            ["核心工具", "market_data.get_stock(), technical_analysis.compute(), market_data.get_sector()"],
            ["信号算法", "趋势信号=MA5上穿MA20(权重0.3)+MACD金叉(0.3)+成交量放大(0.2)+RSI区间(0.2)"],
            ["触发条件", "定时巡检(每5分钟)、用户查询、价格波动>3%"],
        ],
        col_widths=[2.5, 10.5])
    add_body(doc, "")
    add_heading3(doc, "ReAct示例")
    add_code_block(doc, """
用户: "茅台今天走势怎么样？"
Thought: 需要获取600519实时行情和技术指标
Action: market_data.get_stock("600519")
Observation: 价1856, 涨+2.3%, 量放大1.5倍
Thought: 需要计算技术指标判断趋势
Action: technical_analysis.compute("600519", ["MACD","RSI","KDJ"])
Observation: MACD金叉, RSI=62, KDJ向上
Thought: 综合判断，趋势偏多
Final: 茅台今日走势偏强，MACD金叉确认，成交量放大支撑...""")

    # 3.2 新闻解读员
    add_heading2(doc, "3.2 新闻解读员（News Interpreter Agent）")
    create_table(doc,
        ["属性", "详情"],
        [
            ["职责", "金融资讯智能解读、事件提取、情绪量化"],
            ["核心工具", "news.search(), sentiment.analyze(), event.extract()"],
            ["信号算法", "新闻信号=事件重要性(0.4)+情绪强度(0.3)+关联度(0.3)"],
            ["触发条件", "定时巡检(每15分钟)、重大事件推送、用户查询"],
        ],
        col_widths=[2.5, 10.5])

    # 3.3 风控专员
    add_heading2(doc, "3.3 风控专员（Risk Controller Agent）")
    create_table(doc,
        ["属性", "详情"],
        [
            ["职责", "实时风险评估、VaR计算、压力测试、阈值预警"],
            ["核心工具", "risk.compute_var(), risk.stress_test(), risk.check_threshold()"],
            ["信号算法", "风险信号=VaR突破(0.35)+回撤幅度(0.3)+波动率(0.2)+集中度(0.15)"],
            ["触发条件", "定时巡检(每1分钟)、VaR突破阈值、回撤超限"],
        ],
        col_widths=[2.5, 10.5])

    # 3.4 策略顾问
    add_heading2(doc, "3.4 策略顾问（Strategy Advisor Agent）")
    create_table(doc,
        ["属性", "详情"],
        [
            ["职责", "投资策略生成、历史回测、收益归因"],
            ["核心工具", "strategy.build(), strategy.backtest(), strategy.attrib()"],
            ["信号算法", "策略信号=因子IC(0.3)+回测夏普(0.3)+最大回撤(0.2)+换手率(0.2)"],
            ["触发条件", "用户请求、市场分析师信号触发、定期策略评估"],
        ],
        col_widths=[2.5, 10.5])

    # 3.5 交易执行员
    add_heading2(doc, "3.5 交易执行员（Execution Agent）")
    create_table(doc,
        ["属性", "详情"],
        [
            ["职责", "交易信号生成、执行优化、合规检查"],
            ["核心工具", "signal.generate(), execution.optimize(), compliance.check()"],
            ["信号算法", "交易信号=策略信号(0.4)+风控通过(0.3)+合规通过(0.2)+流动性(0.1)"],
            ["触发条件", "策略顾问信号、风控预警解除、用户指令"],
        ],
        col_widths=[2.5, 10.5])

    # 3.6 报告撰写员
    add_heading2(doc, "3.6 报告撰写员（Report Writer Agent）")
    create_table(doc,
        ["属性", "详情"],
        [
            ["职责", "专业投研报告自动生成、数据可视化、合规审查"],
            ["核心工具", "report.generate(), chart.create(), compliance.review()"],
            ["信号算法", "报告优先级=紧急程度(0.4)+影响范围(0.3)+用户偏好(0.3)"],
            ["触发条件", "定时报告(日/周/月)、用户请求、重大事件"],
        ],
        col_widths=[2.5, 10.5])

    doc.add_page_break()

    # 第四章：多智能体协同
    add_heading1(doc, "第四章  多智能体协同")
    add_heading2(doc, "4.1 Orchestrator调度")
    add_body(doc, "Orchestrator是FinAgent Pro的「指挥中心」，负责接收用户请求，分解任务，调度智能体执行，并聚合最终结果。其核心能力包括：")
    add_bullet(doc, "任务分解：将复杂请求拆解为可并行的子任务", bold_prefix="")
    add_bullet(doc, "智能调度：根据任务类型自动选择合适的智能体组合", bold_prefix="")
    add_bullet(doc, "结果聚合：合并多个智能体的输出，消除冲突，形成一致结论", bold_prefix="")
    add_bullet(doc, "质量保障：对最终输出进行合规审查和质量评分", bold_prefix="")

    add_heading2(doc, "4.2 三阶段流水线")
    add_body(doc, "Orchestrator采用三阶段流水线处理用户请求：")
    create_table(doc,
        ["阶段", "名称", "执行内容", "参与智能体"],
        [
            ["Phase 1", "信息采集\n(Gathering)", "并行采集市场数据、新闻资讯、风险指标", "Market + News + Risk"],
            ["Phase 2", "分析推理\n(Reasoning)", "策略生成、信号计算、风险评估", "Strategy + Risk + Execution"],
            ["Phase 3", "输出合成\n(Synthesis)", "报告生成、合规审查、结果推送", "Report + Compliance"],
        ],
        col_widths=[2.0, 3.0, 4.5, 3.5])

    add_heading2(doc, "4.3 加权协商算法")
    add_body(doc, "当多个智能体给出不同结论时，Orchestrator采用加权协商算法进行冲突消解：")
    add_code_block(doc, """
class WeightedNegotiation:
    \"\"\"加权协商算法\"\"\"

    def negotiate(self, signals: List[AgentSignal]) -> ConsensusSignal:
        # 1. 收集各智能体信号
        weighted_scores = {}
        for signal in signals:
            weight = self.get_agent_weight(signal.agent_type)
            for factor, score in signal.factors.items():
                if factor not in weighted_scores:
                    weighted_scores[factor] = 0
                weighted_scores[factor] += score * weight

        # 2. 归一化
        total_weight = sum(self.get_agent_weight(s.agent_type) for s in signals)
        for factor in weighted_scores:
            weighted_scores[factor] /= total_weight

        # 3. 生成共识信号
        return ConsensusSignal(
            direction=self._determine_direction(weighted_scores),
            confidence=self._compute_confidence(signals),
            factors=weighted_scores,
            dissent=self._record_dissent(signals)
        )

    def get_agent_weight(self, agent_type: str) -> float:
        \"\"\"智能体权重配置\"\"\"
        weights = {
            "market": 0.25,   # 市场分析师
            "news": 0.20,     # 新闻解读员
            "risk": 0.25,     # 风控专员（高权重）
            "strategy": 0.15, # 策略顾问
            "execution": 0.10,# 交易执行员
            "report": 0.05    # 报告撰写员
        }
        return weights.get(agent_type, 0.1)""")

    add_heading2(doc, "4.4 LLM增强推理")
    add_body(doc, "在加权协商的基础上，Orchestrator还引入LLM增强推理，对协商结果进行二次校验和深度分析：")
    add_bullet(doc, "一致性校验：检查各智能体结论是否存在逻辑矛盾", bold_prefix="")
    add_bullet(doc, "深度推理：对关键决策点进行更深入的分析推理", bold_prefix="")
    add_bullet(doc, "风险补充：识别协商过程中可能遗漏的风险因素", bold_prefix="")
    add_bullet(doc, "可解释性增强：生成决策过程的自然语言解释", bold_prefix="")

    doc.add_page_break()

    # 第五章：主动智能
    add_heading1(doc, "第五章  主动智能")
    add_body(doc, "FinAgent Pro的主动智能系统是区别于传统被动式AI的核心特性，通过四种机制实现7×24小时不间断的主动服务。")

    add_heading2(doc, "5.1 定时巡检")
    add_body(doc, "系统通过APScheduler实现定时巡检，各智能体按不同频率自动执行检查任务：")
    create_table(doc,
        ["智能体", "巡检频率", "巡检内容"],
        [
            ["市场分析师", "每5分钟", "行情异动检测、技术指标扫描"],
            ["新闻解读员", "每15分钟", "最新资讯抓取、事件分类更新"],
            ["风控专员", "每1分钟", "风险指标计算、阈值监控"],
            ["策略顾问", "每日", "策略表现评估、参数优化建议"],
            ["报告撰写员", "每日收盘后", "日报自动生成"],
        ],
        col_widths=[3.0, 3.0, 7.0])

    add_heading2(doc, "5.2 事件驱动")
    add_body(doc, "系统监听关键事件，一旦触发立即启动相关智能体：")
    add_bullet(doc, "价格异动事件：个股涨跌幅>5%、成交量>3倍均值", bold_prefix="")
    add_bullet(doc, "重大新闻事件：政策发布、业绩公告、并购重组", bold_prefix="")
    add_bullet(doc, "风险事件：VaR突破、回撤超限、集中度告警", bold_prefix="")
    add_bullet(doc, "用户事件：登录、查询、交易指令", bold_prefix="")

    add_heading2(doc, "5.3 阈值预警")
    add_code_block(doc, """
class ThresholdMonitor:
    \"\"\"阈值预警系统\"\"\"

    ALERT_RULES = {
        "price_change": {
            "condition": lambda v: abs(v) > 0.05,  # 涨跌幅>5%
            "level": "HIGH",
            "agents": ["market", "risk", "news"]
        },
        "var_breach": {
            "condition": lambda v: v > 0.95,  # VaR突破95%分位
            "level": "CRITICAL",
            "agents": ["risk", "strategy", "execution"]
        },
        "volume_surge": {
            "condition": lambda v: v > 3.0,  # 成交量>3倍均值
            "level": "MEDIUM",
            "agents": ["market", "news"]
        },
        "drawdown": {
            "condition": lambda v: v < -0.10,  # 回撤>10%
            "level": "CRITICAL",
            "agents": ["risk", "strategy", "execution"]
        }
    }""")

    add_heading2(doc, "5.4 计划执行")
    add_body(doc, "用户可以预设分析计划，系统按计划自动执行并推送结果：")
    add_bullet(doc, "每日盘前分析：8:30自动推送当日市场展望", bold_prefix="")
    add_bullet(doc, "盘后总结：15:30自动生成当日交易总结", bold_prefix="")
    add_bullet(doc, "周度策略评估：每周五自动评估策略表现", bold_prefix="")
    add_bullet(doc, "月度深度报告：每月末自动生成深度分析报告", bold_prefix="")

    doc.add_page_break()

    # 第六章：合规引擎
    add_heading1(doc, "第六章  合规引擎")
    add_heading2(doc, "6.1 监管规则")
    add_body(doc, "合规引擎内置200+条金融监管规则，覆盖以下领域：")
    create_table(doc,
        ["规则类别", "规则数量", "覆盖范围"],
        [
            ["信息披露", "45条", "上市公司信披要求、内幕信息管理"],
            ["投资建议", "38条", "投顾资质、建议免责、风险提示"],
            ["风险管理", "52条", "VaR限制、集中度控制、杠杆约束"],
            ["数据安全", "35条", "个人信息保护、数据出境、脱敏要求"],
            ["反洗钱", "30条", "可疑交易识别、客户身份识别"],
        ],
        col_widths=[3.0, 2.5, 7.5])

    add_heading2(doc, "6.2 审计日志")
    add_body(doc, "所有智能体的操作均记录审计日志，确保全流程可追溯：")
    add_code_block(doc, """
class ComplianceAudit:
    \"\"\"合规审计日志\"\"\"

    async def log(self, event: AuditEvent):
        record = {
            "timestamp": datetime.now().isoformat(),
            "agent": event.agent_id,
            "action": event.action,
            "input_hash": hash(event.input_data),   # 输入脱敏
            "output_hash": hash(event.output_data),  # 输出脱敏
            "compliance_result": event.compliance,
            "risk_level": event.risk_level,
            "user_id": event.user_id,
            "session_id": event.session_id
        }
        await self.db.insert("audit_logs", record)

    async def query(self, filters: dict) -> List[dict]:
        \"\"\"合规审计查询\"\"\"
        return await self.db.query("audit_logs", filters)""")

    add_heading2(doc, "6.3 数据安全")
    add_bullet(doc, "传输层：全链路TLS 1.3加密", bold_prefix="")
    add_bullet(doc, "存储层：AES-256-GCM加密，密钥HSM管理", bold_prefix="")
    add_bullet(doc, "应用层：RBAC权限模型，数据脱敏输出", bold_prefix="")
    add_bullet(doc, "审计层：操作全留痕，支持合规审计追溯", bold_prefix="")
    add_bullet(doc, "隔离层：多租户数据逻辑隔离，企业版物理隔离", bold_prefix="")

    doc.add_page_break()

    # 第七章：性能与扩展
    add_heading1(doc, "第七章  性能与扩展")
    add_heading2(doc, "7.1 性能指标")
    create_table(doc,
        ["指标", "目标值", "当前值", "优化方向"],
        [
            ["单次分析响应时间", "< 5秒", "3.2秒", "并行工具调用"],
            ["并发用户数", "100+", "80", "异步IO + 连接池"],
            ["LLM推理延迟", "< 2秒", "1.5秒", "流式输出 + 缓存"],
            ["数据更新延迟", "< 1秒", "0.8秒", "WebSocket推送"],
            ["系统可用性", "99.9%", "99.5%", "容灾 + 降级"],
        ],
        col_widths=[3.5, 2.5, 2.5, 4.5])

    add_heading2(doc, "7.2 水平扩展方案")
    add_body(doc, "FinAgent Pro采用微服务架构，支持按智能体维度独立扩展：")
    add_code_block(doc, """
# Docker Compose 扩展示例
services:
  orchestrator:
    replicas: 2        # 调度器2副本
  market-agent:
    replicas: 3        # 市场分析师3副本（高频调用）
  news-agent:
    replicas: 2        # 新闻解读员2副本
  risk-agent:
    replicas: 2        # 风控专员2副本
  strategy-agent:
    replicas: 1        # 策略顾问1副本
  execution-agent:
    replicas: 1        # 交易执行员1副本
  report-agent:
    replicas: 1        # 报告撰写员1副本
  redis:
    # 消息队列 + 缓存
  chromadb:
    # 向量数据库""")

    add_heading2(doc, "7.3 新智能体接入流程")
    add_body(doc, "FinAgent Pro支持快速接入新的智能体，标准流程如下：")
    add_bullet(doc, "Step 1: 继承BaseAgent基类，实现process()方法", bold_prefix="")
    add_bullet(doc, "Step 2: 注册工具到ToolRegistry", bold_prefix="")
    add_bullet(doc, "Step 3: 配置信号算法和触发条件", bold_prefix="")
    add_bullet(doc, "Step 4: 在Orchestrator中注册调度规则", bold_prefix="")
    add_bullet(doc, "Step 5: 编写单元测试和集成测试", bold_prefix="")
    add_bullet(doc, "Step 6: 部署上线，监控运行指标", bold_prefix="")

    doc.add_page_break()

    # 第八章：部署方案
    add_heading1(doc, "第八章  部署方案")
    add_heading2(doc, "8.1 Docker部署")
    add_code_block(doc, """
# Dockerfile (后端)
FROM python:3.11-slim
WORKDIR /app
COPY requirements.txt .
RUN pip install --no-cache-dir -r requirements.txt
COPY . .
EXPOSE 8000
CMD ["uvicorn", "main:app", "--host", "0.0.0.0", "--port", "8000"]

# docker-compose.yml 核心配置
version: '3.8'
services:
  backend:
    build: ./backend
    ports: ["8000:8000"]
    env_file: .env
    depends_on: [redis, chromadb]
  frontend:
    build: ./frontend
    ports: ["3000:3000"]
  redis:
    image: redis:7-alpine
  chromadb:
    image: chromadb/chroma:latest
    volumes: ["chroma-data:/chroma/chroma"]
volumes:
  chroma-data:""")

    add_heading2(doc, "8.2 Vercel部署")
    add_body(doc, "前端采用Vercel部署，实现全球CDN加速和自动HTTPS：")
    add_code_block(doc, """
# vercel.json 配置
{
  "builds": [
    { "src": "frontend/package.json", "use": "@vercel/node" }
  ],
  "routes": [
    { "src": "/api/(.*)", "dest": "https://backend.example.com/$1" },
    { "src": "/(.*)", "dest": "/frontend/$1" }
  ]
}""")

    add_heading2(doc, "8.3 生产环境配置")
    create_table(doc,
        ["配置项", "开发环境", "生产环境"],
        [
            ["LLM模型", "GPT-4o-mini", "GPT-4o + DeepSeek双模型"],
            ["数据库", "SQLite", "PostgreSQL + Redis"],
            ["向量库", "ChromaDB本地", "ChromaDB集群"],
            ["缓存", "内存缓存", "Redis集群"],
            ["监控", "控制台日志", "Prometheus + Grafana"],
            ["日志", "文件日志", "ELK Stack"],
            ["CI/CD", "手动部署", "GitHub Actions自动部署"],
            ["SSL", "自签名证书", "Let's Encrypt自动续期"],
        ],
        col_widths=[3.0, 4.5, 5.5])

    # 保存
    output_path = r"D:\AFAC2026金融智能创新大赛\FinAgent-Pro\docs\FinAgentPro技术方案.docx"
    doc.save(output_path)
    print(f"✅ 技术方案已生成: {output_path}")
    return output_path


# ============================================================
# 文档3：项目申报书
# ============================================================

def create_application_doc():
    doc = Document()
    setup_document(doc, "AFAC2026 项目申报书")

    # 封面
    add_cover_page(doc,
        "AFAC2026金融智能创新大赛",
        "初创组 · 项目申报书",
        ["项目名称：FinAgent Pro 金融自主智能体平台", "2026年6月"])

    # 一、项目基本信息
    add_heading1(doc, "一、项目基本信息")
    create_table(doc,
        ["项目", "内容"],
        [
            ["项目名称", "FinAgent Pro 金融自主智能体平台"],
            ["赛道方向", "金融科技创新应用"],
            ["参赛组别", "初创组"],
            ["团队名称", "FinAgent Team"],
            ["项目负责人", "张明"],
            ["联系电话", "138-XXXX-XXXX"],
            ["电子邮箱", "contact@finagent-pro.com"],
            ["所在单位", "FinAgent Pro团队"],
            ["项目阶段", "MVP开发阶段"],
            ["知识产权", "软件著作权申请中"],
        ],
        col_widths=[3.5, 9.5])

    doc.add_page_break()

    # 二、项目概述
    add_heading1(doc, "二、项目概述")
    add_heading2(doc, "2.1 项目简介")
    add_body(doc, "FinAgent Pro是一款面向金融行业的自主智能体（Agentic AI）平台，通过六大专业数字员工——市场分析师、新闻解读员、风控专员、策略顾问、交易执行员和报告撰写员——为金融机构提供7×24小时全链路智能投研服务。平台基于ReAct推理引擎实现智能体自主决策，通过多智能体协同机制实现深度分析，以主动智能系统实现实时预警，并内置合规引擎确保全流程合规。")

    add_heading2(doc, "2.2 核心创新——六大创新点")
    create_table(doc,
        ["序号", "创新点", "创新描述"],
        [
            ["1", "ReAct自主推理", "自研推理引擎，实现「思考-行动-观察」闭环，智能体可自主规划任务、调用工具、评估结果"],
            ["2", "多智能体协同", "6大专业智能体分工协作，Orchestrator统一调度，加权协商算法消解冲突"],
            ["3", "主动智能系统", "定时巡检+事件驱动+阈值预警+计划执行，7×24小时主动服务"],
            ["4", "合规引擎内嵌", "200+监管规则内置，所有输出自动合规审查，审计全程留痕"],
            ["5", "三层记忆架构", "短期+工作+长期记忆，智能体持续学习优化，越用越智能"],
            ["6", "思维链可视化", "推理过程透明可追溯，决策可解释，增强用户信任"],
        ],
        col_widths=[1.0, 3.0, 9.0])

    add_heading2(doc, "2.3 应用场景")
    add_bullet(doc, "投研分析：分析师使用智能体快速完成多维度投研分析，效率提升10倍", bold_prefix="")
    add_bullet(doc, "风控预警：风控团队7×24小时实时监控，风险事件秒级响应", bold_prefix="")
    add_bullet(doc, "合规审查：合规团队借助合规引擎自动审查，效率提升80%", bold_prefix="")
    add_bullet(doc, "投资决策：投资经理获取多智能体协同分析结论，辅助科学决策", bold_prefix="")
    add_bullet(doc, "个人投研：个人投资者享受专业级投研服务，降低信息不对称", bold_prefix="")

    doc.add_page_break()

    # 三、技术与产品
    add_heading1(doc, "三、技术与产品")
    add_heading2(doc, "3.1 技术路线——Agentic AI + 多智能体协同")
    add_body(doc, "FinAgent Pro采用Agentic AI技术路线，以ReAct推理引擎为核心，结合多智能体协同、主动智能和合规引擎，构建金融行业专属的自主智能体平台。技术路线的核心优势在于：")
    add_bullet(doc, "自主性：智能体可自主规划、决策和行动，而非被动响应", bold_prefix="")
    add_bullet(doc, "协同性：多智能体分工协作，1+1>2的群体智能", bold_prefix="")
    add_bullet(doc, "主动性：无需用户触发，主动发现和推送关键信息", bold_prefix="")
    add_bullet(doc, "合规性：合规引擎全程嵌入，确保输出安全合规", bold_prefix="")

    add_heading2(doc, "3.2 产品功能——六大数字员工")
    create_table(doc,
        ["数字员工", "核心功能", "技术实现"],
        [
            ["市场分析师", "行情监测、技术分析、趋势识别", "ReAct + 技术指标工具 + 定时巡检"],
            ["新闻解读员", "资讯解读、事件提取、情绪分析", "ReAct + NLP工具 + 事件驱动"],
            ["风控专员", "风险评估、VaR计算、阈值预警", "ReAct + 风控模型 + 实时监控"],
            ["策略顾问", "策略生成、历史回测、收益归因", "ReAct + 回测引擎 + 因子分析"],
            ["交易执行员", "信号生成、执行优化、合规检查", "ReAct + 信号算法 + 合规审查"],
            ["报告撰写员", "报告生成、数据可视化、合规审查", "ReAct + 模板引擎 + 图表工具"],
        ],
        col_widths=[2.5, 4.5, 6.0])

    add_heading2(doc, "3.3 技术指标")
    create_table(doc,
        ["指标", "目标值", "说明"],
        [
            ["单次分析响应时间", "< 5秒", "从用户提问到返回分析结果"],
            ["并发用户数", "100+", "同时在线使用用户数"],
            ["系统可用性", "99.9%", "年度停机时间<8.76小时"],
            ["合规审查覆盖率", "100%", "所有输出均经过合规审查"],
            ["预警响应时间", "< 30秒", "从风险触发到预警推送"],
            ["报告生成时间", "< 3分钟", "完整投研报告自动生成"],
        ],
        col_widths=[4.0, 3.0, 6.0])

    doc.add_page_break()

    # 四、市场与商业
    add_heading1(doc, "四、市场与商业")
    add_heading2(doc, "4.1 市场规模")
    add_body(doc, "中国金融科技AI应用市场规模达500亿元（TAM），其中AI+投研/风控/合规细分市场50亿元（SAM），Agentic AI金融平台市场5亿元（SOM）。市场年复合增长率超过35%，先发优势明显。")

    add_heading2(doc, "4.2 商业模式")
    add_body(doc, "采用SaaS三层定价模式：基础版（299元/月，个人投资者）、专业版（999元/月，小型工作室）、企业版（5万起/年，金融机构）。预计第三年ARR达1.2亿元，毛利率82%。")

    add_heading2(doc, "4.3 竞争分析")
    create_table(doc,
        ["维度", "传统金融终端", "通用AI工具", "FinAgent Pro"],
        [
            ["专业深度", "高", "低", "高"],
            ["智能程度", "低", "中", "高"],
            ["自主性", "无", "低", "高"],
            ["合规能力", "中", "无", "高"],
            ["价格", "高", "低", "中"],
        ],
        col_widths=[2.5, 3.5, 3.5, 3.5])
    add_body(doc, "")
    add_body(doc, "FinAgent Pro填补了传统金融终端与通用AI工具之间的市场空白，具备差异化竞争优势。")

    doc.add_page_break()

    # 五、团队情况
    add_heading1(doc, "五、团队情况")
    add_heading2(doc, "5.1 核心成员")
    create_table(doc,
        ["姓名", "职位", "背景", "核心贡献"],
        [
            ["张明", "CEO/创始人", "10年金融科技经验，前头部券商技术负责人", "产品战略、商业模式、行业资源"],
            ["李华", "CTO/联合创始人", "8年AI研发经验，前大厂AI Lab高级工程师", "技术架构、ReAct引擎、多智能体系统"],
            ["王芳", "首席风控官", "12年风控经验，CFA/FRM持证", "风控模型、合规引擎、监管对接"],
            ["赵强", "产品总监", "6年金融产品经验，前知名Fintech产品经理", "产品规划、用户体验、需求分析"],
            ["陈静", "算法负责人", "5年NLP经验，博士毕业于清华大学", "NLP算法、情绪分析、知识图谱"],
        ],
        col_widths=[2.0, 3.0, 4.5, 3.5])

    add_heading2(doc, "5.2 团队优势")
    add_bullet(doc, "金融+AI跨界组合：核心成员兼具金融行业深度和AI技术能力", bold_prefix="")
    add_bullet(doc, "完整能力闭环：产品、技术、风控、算法全覆盖", bold_prefix="")
    add_bullet(doc, "行业资源丰富：团队成员在金融行业拥有广泛人脉和客户资源", bold_prefix="")
    add_bullet(doc, "执行力强：团队已有MVP原型，具备快速迭代能力", bold_prefix="")

    doc.add_page_break()

    # 六、项目进展
    add_heading1(doc, "六、项目进展")
    add_heading2(doc, "6.1 已完成工作")
    create_table(doc,
        ["序号", "工作内容", "完成状态", "说明"],
        [
            ["1", "技术架构设计", "✅ 已完成", "六层架构设计，技术选型确定"],
            ["2", "ReAct推理引擎", "✅ 已完成", "核心推理循环实现，支持多步推理"],
            ["3", "六大智能体原型", "✅ 已完成", "基础功能实现，可独立运行"],
            ["4", "Orchestrator调度", "✅ 已完成", "三阶段流水线，加权协商算法"],
            ["5", "前端界面", "🔄 进行中", "Dashboard、分析页、对话页已完成"],
            ["6", "合规引擎", "🔄 进行中", "核心规则已实现，持续扩充中"],
        ],
        col_widths=[1.0, 3.5, 2.5, 6.0])

    add_heading2(doc, "6.2 后续计划")
    create_table(doc,
        ["阶段", "时间", "计划内容"],
        [
            ["MVP完善", "2026年7-8月", "功能完善、性能优化、内测准备"],
            ["内测上线", "2026年9月", "邀请50名种子用户内测，收集反馈"],
            ["公测发布", "2026年11月", "开放注册，基础版和专业版上线"],
            ["企业版发布", "2027年3月", "企业版功能开发完成，启动销售"],
        ],
        col_widths=[2.5, 3.0, 7.5])

    add_heading2(doc, "6.3 里程碑")
    create_table(doc,
        ["里程碑", "预计时间", "关键指标"],
        [
            ["MVP上线", "2026年9月", "核心功能可用，100+注册用户"],
            ["首批付费", "2026年11月", "10+付费用户，ARR突破50万"],
            ["企业版发布", "2027年3月", "5+企业客户，ARR突破500万"],
            ["规模化增长", "2027年Q4", "200+付费用户，ARR突破1000万"],
        ],
        col_widths=[3.0, 3.0, 7.0])

    doc.add_page_break()

    # 七、融资需求
    add_heading1(doc, "七、融资需求")
    add_heading2(doc, "7.1 融资金额")
    add_body(doc, "天使轮融资500万元人民币，出让股权10%-15%，投前估值3,300万-5,000万元。")

    add_heading2(doc, "7.2 资金用途")
    create_table(doc,
        ["用途", "金额（万元）", "占比"],
        [
            ["产品研发", "250", "50%"],
            ["团队建设", "100", "20%"],
            ["市场推广", "80", "16%"],
            ["数据采购", "40", "8%"],
            ["运营储备", "30", "6%"],
        ],
        col_widths=[4.0, 4.0, 3.0])

    add_heading2(doc, "7.3 预期回报")
    add_body(doc, "预计3年内ARR达1.2亿元，净利润4,000万元。投资回报预期：3-5倍回报，退出方式包括IPO上市、战略并购或股权回购。")

    doc.add_page_break()

    # 八、社会价值
    add_heading1(doc, "八、社会价值")
    add_heading2(doc, "8.1 响应国家战略")
    add_body(doc, "FinAgent Pro积极响应中央金融工作会议「五篇大文章」战略：科技金融——以AI技术推动金融科技自主创新；绿色金融——ESG智能评估助力绿色投资；普惠金融——降低投研门槛服务中小投资者；养老金融——稳健策略服务养老资金；数字金融——推动金融行业数字化转型。")

    add_heading2(doc, "8.2 促进就业")
    add_bullet(doc, "直接就业：团队规模从5人扩展至30人，创造高质量AI+金融就业岗位", bold_prefix="")
    add_bullet(doc, "间接就业：通过生态建设，带动数据标注、模型训练、客户服务等上下游就业", bold_prefix="")
    add_bullet(doc, "技能升级：帮助金融从业者掌握AI工具，提升职业技能和竞争力", bold_prefix="")

    add_heading2(doc, "8.3 推动行业进步")
    add_bullet(doc, "提升效率：投研效率提升10倍，推动行业生产力变革", bold_prefix="")
    add_bullet(doc, "降低风险：实时风控预警，减少金融风险事件", bold_prefix="")
    add_bullet(doc, "促进合规：合规引擎内嵌，提升行业合规水平", bold_prefix="")
    add_bullet(doc, "技术示范：Agentic AI在金融领域的标杆应用，引领行业技术方向", bold_prefix="")

    doc.add_page_break()

    # 九、知识产权
    add_heading1(doc, "九、知识产权")
    add_heading2(doc, "9.1 已申请/计划申请")
    create_table(doc,
        ["类型", "名称", "状态"],
        [
            ["软件著作权", "FinAgent Pro金融自主智能体平台", "申请中"],
            ["发明专利", "基于ReAct的多智能体金融分析方法", "计划申请"],
            ["发明专利", "金融合规引擎自动审查方法", "计划申请"],
            ["发明专利", "多智能体加权协商决策方法", "计划申请"],
        ],
        col_widths=[3.0, 6.0, 3.0])

    add_heading2(doc, "9.2 开源策略")
    add_bullet(doc, "核心引擎闭源：ReAct推理引擎、合规引擎等核心组件闭源，保持技术壁垒", bold_prefix="")
    add_bullet(doc, "工具层开源：通用工具组件（技术分析、NLP等）开源，建设开发者生态", bold_prefix="")
    add_bullet(doc, "SDK开放：提供Agent SDK，支持第三方开发自定义智能体", bold_prefix="")
    add_bullet(doc, "数据接口开放：提供标准化数据API，方便第三方集成", bold_prefix="")

    doc.add_page_break()

    # 十、承诺与声明
    add_heading1(doc, "十、承诺与声明")
    add_body(doc, "本团队郑重承诺：")
    add_bullet(doc, "本项目为原创项目，不存在抄袭、剽窃等知识产权纠纷", bold_prefix="一、")
    add_bullet(doc, "本项目所提交的所有材料真实、准确、完整", bold_prefix="二、")
    add_bullet(doc, "本团队将严格遵守大赛规则，公平竞争", bold_prefix="三、")
    add_bullet(doc, "本项目如获资助，将严格按照计划使用资金，接受监督", bold_prefix="四、")
    add_bullet(doc, "本团队将积极推动项目落地，实现社会价值和商业价值", bold_prefix="五、")
    add_bullet(doc, "本项目所有AI生成内容均经过人工审核，确保合规性", bold_prefix="六、")
    add_bullet(doc, "本项目不提供直接投资建议，所有分析结果仅供参考", bold_prefix="七、")

    add_body(doc, "")
    add_body(doc, "")
    # 签名区
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("项目团队：FinAgent Team")
    run.font.size = Pt(12)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = p2.add_run("日期：2026年6月")
    run2.font.size = Pt(12)
    run2.font.name = "微软雅黑"
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 保存
    output_path = r"D:\AFAC2026金融智能创新大赛\FinAgent-Pro\docs\FinAgentPro项目申报书.docx"
    doc.save(output_path)
    print(f"✅ 项目申报书已生成: {output_path}")
    return output_path


# ============================================================
# 主函数
# ============================================================

if __name__ == "__main__":
    print("=" * 60)
    print("FinAgent Pro 文档生成工具")
    print("AFAC2026金融智能创新大赛")
    print("=" * 60)
    print()

    print("📄 正在生成文档1: 商业计划书...")
    path1 = create_business_plan()
    print()

    print("📄 正在生成文档2: 技术方案...")
    path2 = create_technical_spec()
    print()

    print("📄 正在生成文档3: 项目申报书...")
    path3 = create_application_doc()
    print()

    print("=" * 60)
    print("✅ 全部文档生成完成！")
    print(f"  1. 商业计划书: {path1}")
    print(f"  2. 技术方案:   {path2}")
    print(f"  3. 项目申报书: {path3}")
    print("=" * 60)

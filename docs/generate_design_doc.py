# -*- coding: utf-8 -*-
"""生成 FinAgent Pro 交互设计图与逻辑流程图文档"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = r"D:\AFAC2026金融智能创新大赛\FinAgent-Pro\docs\FinAgentPro交互设计与流程图.docx"

DARK_BLUE = RGBColor(0x1A, 0x3C, 0x6E)
ACCENT_BLUE = RGBColor(0x2B, 0x5C, 0x9E)
LIGHT_BLUE_BG = "D6E4F0"
WHITE = "FFFFFF"
DARK_BG = "1A3C6E"
MID_BG = "2B5C9E"
LIGHT_BG = "E8F0FE"
GRAY_TEXT = RGBColor(0x55, 0x55, 0x55)
SIDEBAR_BG = "1E293B"
CARD_BG = "1E293B"
GREEN = "27AE60"
RED = "EF4444"
YELLOW = "F59E0B"
BLUE_ACCENT = "3B82F6"

FONT_NAME = "微软雅黑"
MONO = "Consolas"


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_run(paragraph, text, bold=False, size=11, color=None, font_name=FONT_NAME):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color
    return run


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.15):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = line_spacing


def add_header(doc):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "FinAgent Pro 交互设计与流程图", bold=True, size=9, color=DARK_BLUE)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="1A3C6E"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.size = Pt(9)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)
    run2 = p.add_run()
    run2.font.size = Pt(9)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._element.append(instrText)
    run3 = p.add_run()
    run3.font.size = Pt(9)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fldChar2)


def create_cover_page(doc):
    for _ in range(6):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=12, line_spacing=1.0)
    add_run(p, "FinAgent Pro", bold=True, size=36, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.0)
    add_run(p, "交互设计与逻辑流程图", bold=True, size=28, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=12, after=12, line_spacing=1.0)
    add_run(p, "━" * 30, size=12, color=ACCENT_BLUE)

    for text in ["AFAC2026金融智能创新大赛 初创组", "智融先锋团队", "2026年6月"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=0, after=4, line_spacing=1.5)
        add_run(p, text, size=14, color=GRAY_TEXT)

    doc.add_page_break()


def add_section_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        run.font.color.rgb = DARK_BLUE
        run.font.size = Pt(18 if level == 1 else 14)
    set_paragraph_spacing(heading, before=18, after=10)
    return heading


def add_body_paragraph(doc, text, bold_prefix=None, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    set_paragraph_spacing(p, before=2, after=4, line_spacing=1.5)
    if bold_prefix:
        add_run(p, bold_prefix, bold=True, size=11)
    add_run(p, text, size=11)
    return p


def add_code_block(doc, code, bg_color=LIGHT_BLUE_BG):
    """添加代码块样式的ASCII图"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=6, line_spacing=1.0)
    run = add_run(p, code, size=8, color=DARK_BLUE, font_name=MONO)
    run.font.name = MONO
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{bg_color}"/>')
    p._element.get_or_add_pPr().append(shd)


def set_table_borders(table, color="1A3C6E", sz="6"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_dashboard_mockup(doc):
    """Dashboard 页面交互设计"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4)
    add_run(p, "▎ 页面1：Dashboard 工作台", bold=True, size=13, color=DARK_BLUE)

    mockup = (
        "┌──────────────────────────────────────────────────────────────────┐\n"
        "│  FinAgent Pro                              🔔 3条预警  👤 冯亦根  │\n"
        "├────────┬─────────────────────────────────────────────────────────┤\n"
        "│        │  ┌─ 晨间报告 ──────────┐  ┌─ 市场概览 ──────────┐     │\n"
        "│ 🏠 工作台│  │ 📅 2026-06-12 晨报  │  │ 上证 3,256  +0.8%  │     │\n"
        "│ 📊 智能分析│  │ 隔夜美股：纳指+1.2% │  │ 深证 10,856 +1.1%  │     │\n"
        "│ 💬 数字员工│  │ A股预判：预计高开   │  │ 创业板 2,156 +1.5% │     │\n"
        "│ 🤖 智能体 │  │ 持仓风险：3只预警   │  │ 科创板 986  +0.6%  │     │\n"
        "│ 📋 报告  │  └────────────────────┘  └────────────────────┘     │\n"
        "│ ⚠️ 预警  │                                                         │\n"
        "│        │  ┌─ 智能体状态 ─────────────────────────────────┐      │\n"
        "│ ────── │  │ 🟢 市场分析  🟢 新闻舆情  🟢 风控合规        │      │\n"
        "│ 智能体  │  │ 🟢 投资策略  🟢 报告生成  🟢 执行监控        │      │\n"
        "│ 状态    │  └──────────────────────────────────────────────┘      │\n"
        "│        │                                                         │\n"
        "│ 🟢 x6  │  ┌─ 主动发现 ─────────────────────────────────┐       │\n"
        "│        │  │ 💡 估值洼地：XX新能源 PE15x 行业平均28x     │       │\n"
        "│        │  │ ⚠️ 持仓预警：YY科技 业绩下滑40%             │       │\n"
        "│        │  └──────────────────────────────────────────────┘       │\n"
        "└────────┴─────────────────────────────────────────────────────────┘"
    )
    add_code_block(doc, mockup)

    # 交互说明表格
    table = doc.add_table(rows=6, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        row.cells[0].width = Cm(3)
        row.cells[1].width = Cm(4)
        row.cells[2].width = Cm(9)

    headers = ["交互元素", "交互方式", "功能说明"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, DARK_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=4, after=4)
        add_run(p, h, bold=True, size=10, color=RGBColor.from_string(WHITE))

    interactions = [
        ["侧边栏导航", "点击切换", "6个页面快速切换，当前页高亮显示"],
        ["晨间报告卡片", "点击展开", "展示晨报详情，含K线缩略图+风险热力图"],
        ["市场概览卡片", "实时刷新", "四大指数实时行情，5分钟自动刷新"],
        ["智能体状态", "悬停查看", "显示每个智能体运行状态和最后执行时间"],
        ["主动发现卡片", "点击详情", "展开估值洼地/持仓预警的完整分析链路"],
    ]
    for row_idx, (element, action, desc) in enumerate(interactions):
        bg = LIGHT_BG if row_idx % 2 == 0 else WHITE
        for col_idx, text in enumerate([element, action, desc]):
            cell = table.rows[row_idx + 1].cells[col_idx]
            set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx < 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, before=3, after=3)
            add_run(p, text, bold=col_idx == 0, size=10,
                    color=DARK_BLUE if col_idx == 0 else RGBColor(0x33, 0x33, 0x33))

    set_table_borders(table)


def add_analyze_mockup(doc):
    """Analyze 页面交互设计"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=10, after=4)
    add_run(p, "▎ 页面2：智能分析", bold=True, size=13, color=DARK_BLUE)

    mockup = (
        "┌──────────────────────────────────────────────────────────────────┐\n"
        "│  FinAgent Pro > 智能分析                                        │\n"
        "├────────┬─────────────────────────────────────────────────────────┤\n"
        "│        │  ┌─ 输入区 ───────────────────────────────────┐        │\n"
        "│        │  │  股票代码：[  600519  ]  [开始分析]         │        │\n"
        "│        │  └─────────────────────────────────────────────┘        │\n"
        "│        │                                                         │\n"
        "│        │  ┌─ 思维链可视化 (ThinkingChain) ──────────────┐       │\n"
        "│        │  │ Step1 📊 市场分析：RSI=65 MACD金叉 PE=32x   │       │\n"
        "│        │  │ Step2 📰 新闻舆情：正面72% 中性20% 负面8%   │       │\n"
        "│        │  │ Step3 ⚖️ 风控合规：波动率中等 ✅合规         │       │\n"
        "│        │  │ Step4 🧠 投资策略：🟢买入 置信度82%          │       │\n"
        "│        │  └──────────────────────────────────────────────┘       │\n"
        "│        │                                                         │\n"
        "│        │  ┌─ 分析结果 ─────────────────────────────────┐       │\n"
        "│        │  │  🟢 买入信号  置信度82%  目标价1980        │       │\n"
        "│        │  │  止损价1780   仓位建议15%                   │       │\n"
        "│        │  │  ┌──────┐ ┌──────┐ ┌──────┐ ┌──────┐      │       │\n"
        "│        │  │  │技术面 │ │基本面 │ │资金面 │ │舆情面 │      │       │\n"
        "│        │  │  │+0.8  │ │+0.6  │ │+0.9  │ │+0.7  │      │       │\n"
        "│        │  │  └──────┘ └──────┘ └──────┘ └──────┘      │       │\n"
        "│        │  └──────────────────────────────────────────────┘       │\n"
        "└────────┴─────────────────────────────────────────────────────────┘"
    )
    add_code_block(doc, mockup)


def add_chat_mockup(doc):
    """Chat 页面交互设计"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=10, after=4)
    add_run(p, "▎ 页面3：数字员工对话", bold=True, size=13, color=DARK_BLUE)

    mockup = (
        "┌──────────────────────────────────────────────────────────────────┐\n"
        "│  FinAgent Pro > 数字员工                                        │\n"
        "├────────┬─────────────────────────────────────────────────────────┤\n"
        "│        │                                                         │\n"
        "│        │  ┌──────────────────────────────────────────┐           │\n"
        "│        │  │ 🤖 FinAgent Pro                          │           │\n"
        "│        │  │ 您好！我是您的金融数字员工，可以为您分析  │           │\n"
        "│        │  │ 股票、评估风险、生成报告。请问需要什么？ │           │\n"
        "│        │  └──────────────────────────────────────────┘           │\n"
        "│        │                                                         │\n"
        "│        │                    ┌──────────────────────────┐         │\n"
        "│        │                    │ 👤 用户                  │         │\n"
        "│        │                    │ 比亚迪最近怎么样？       │         │\n"
        "│        │                    └──────────────────────────┘         │\n"
        "│        │                                                         │\n"
        "│        │  ┌──────────────────────────────────────────┐           │\n"
        "│        │  │ 🤖 FinAgent Pro                          │           │\n"
        "│        │  │ 正在调用智能体分析...                     │           │\n"
        "│        │  │ 📊 市场分析 → 📰 新闻舆情 → 🧠 策略生成 │           │\n"
        "│        │  │                                          │           │\n"
        "│        │  │ 比亚迪(002594)综合分析：                  │           │\n"
        "│        │  │ ▸ 基本面：营收YoY+22%，毛利率18.5%      │           │\n"
        "│        │  │ ▸ 技术面：MACD金叉，支撑位265元          │           │\n"
        "│        │  │ ▸ 资金面：北向资金连续5日净流入           │           │\n"
        "│        │  │ ▸ 综合评级：⭐⭐⭐⭐ 买入(4/5)          │           │\n"
        "│        │  └──────────────────────────────────────────┘           │\n"
        "│        │                                                         │\n"
        "│        │  ┌──────────────────────────────────────────┐           │\n"
        "│        │  │ [  输入您的问题...              ] [发送]  │           │\n"
        "│        │  └──────────────────────────────────────────┘           │\n"
        "└────────┴─────────────────────────────────────────────────────────┘"
    )
    add_code_block(doc, mockup)


def add_alerts_mockup(doc):
    """Alerts 页面交互设计"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=10, after=4)
    add_run(p, "▎ 页面4：预警中心", bold=True, size=13, color=DARK_BLUE)

    mockup = (
        "┌──────────────────────────────────────────────────────────────────┐\n"
        "│  FinAgent Pro > 预警中心                                        │\n"
        "├────────┬─────────────────────────────────────────────────────────┤\n"
        "│        │  ┌─ 预警列表 ─────────────────────────────────┐       │\n"
        "│        │  │ 🔴 高风险  YY科技 业绩预警 利润下滑40%     │       │\n"
        "│        │  │   风控建议：减仓至合规线 ✅已通过合规审查   │       │\n"
        "│        │  │                                            │       │\n"
        "│        │  │ 🟡 中风险  ZZ医药 RSI超买 RSI=82           │       │\n"
        "│        │  │   风控建议：关注回调风险 ✅已通过合规审查   │       │\n"
        "│        │  │                                            │       │\n"
        "│        │  │ 🟢 机会    XX新能源 PE15x 估值洼地         │       │\n"
        "│        │  │   策略建议：逢低布局 ✅已通过合规审查       │       │\n"
        "│        │  └────────────────────────────────────────────┘       │\n"
        "│        │                                                         │\n"
        "│        │  ┌─ 合规审计日志 ─────────────────────────────┐       │\n"
        "│        │  │ 19:30:12 /analyze 600519 ✅ 合规            │       │\n"
        "│        │  │ 19:28:45 /chat 比亚迪   ✅ 合规            │       │\n"
        "│        │  │ 19:25:30 /portfolio    ✅ 合规              │       │\n"
        "│        │  └────────────────────────────────────────────┘       │\n"
        "└────────┴─────────────────────────────────────────────────────────┘"
    )
    add_code_block(doc, mockup)


def add_system_flow(doc):
    """系统逻辑流程图"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4)
    add_run(p, "▎ 系统整体逻辑流程图", bold=True, size=13, color=DARK_BLUE)

    flow = (
        "                         ┌─────────────┐\n"
        "                         │   用户请求   │\n"
        "                         │ /analyze     │\n"
        "                         │ /chat        │\n"
        "                         │ 定时触发     │\n"
        "                         └──────┬──────┘\n"
        "                                │\n"
        "                                ▼\n"
        "                    ┌───────────────────────┐\n"
        "                    │  Master Orchestrator   │\n"
        "                    │  意图分类 + 任务分发    │\n"
        "                    └───────────┬───────────┘\n"
        "                                │\n"
        "              ┌─────────────────┼─────────────────┐\n"
        "              │                 │                 │\n"
        "              ▼                 ▼                 ▼\n"
        "     ┌────────────┐    ┌────────────┐    ┌────────────┐\n"
        "     │  Phase 1   │    │  Phase 1   │    │  Phase 2   │\n"
        "     │ 市场分析    │    │ 新闻舆情    │    │ 风控合规    │\n"
        "     │ 智能体     │    │ 智能体      │    │ 智能体      │\n"
        "     │            │    │            │    │            │\n"
        "     │ ReAct循环  │    │ ReAct循环  │    │ ReAct循环  │\n"
        "     │ ├ Thought  │    │ ├ Thought  │    │ ├ Thought  │\n"
        "     │ ├ Action   │    │ ├ Action   │    │ ├ Action   │\n"
        "     │ └ Observe  │    │ └ Observe  │    │ └ Observe  │\n"
        "     └─────┬──────┘    └─────┬──────┘    └─────┬──────┘\n"
        "           │                 │                 │\n"
        "           └────────┬────────┘                 │\n"
        "                    │                          │\n"
        "                    ▼                          │\n"
        "           ┌────────────────┐                  │\n"
        "           │  加权协商算法   │◄─────────────────┘\n"
        "           │  + LLM增强推理  │\n"
        "           └───────┬────────┘\n"
        "                   │\n"
        "          ┌────────┴────────┐\n"
        "          ▼                 ▼\n"
        "   ┌────────────┐    ┌────────────┐\n"
        "   │  Phase 3   │    │  Phase 3   │\n"
        "   │ 投资策略    │    │ 报告生成    │\n"
        "   │ 智能体     │    │ 智能体      │\n"
        "   └─────┬──────┘    └─────┬──────┘\n"
        "         │                 │\n"
        "         ▼                 ▼\n"
        "   ┌──────────┐     ┌──────────┐\n"
        "   │ 投资建议  │     │ 专业报告  │\n"
        "   │ + 合规审查│     │ + 合规审查│\n"
        "   └─────┬────┘     └─────┬────┘\n"
        "         │                │\n"
        "         ▼                ▼\n"
        "   ┌──────────────────────────┐\n"
        "   │     执行监控智能体        │\n"
        "   │  WebSocket/SSE 实时推送   │\n"
        "   │  + 三层记忆系统存储       │\n"
        "   └──────────────────────────┘"
    )
    add_code_block(doc, flow)


def add_react_flow(doc):
    """ReAct推理流程图"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4)
    add_run(p, "▎ ReAct推理引擎流程图", bold=True, size=13, color=DARK_BLUE)

    flow = (
        "  ┌──────────────────────────────────────────────────────┐\n"
        "  │                  ReAct 推理循环                       │\n"
        "  │                                                      │\n"
        "  │   ┌──────────┐                                       │\n"
        "  │   │  输入问题  │                                      │\n"
        "  │   └────┬─────┘                                       │\n"
        "  │        │                                             │\n"
        "  │        ▼                                             │\n"
        "  │   ┌──────────┐     ┌──────────────────────────┐     │\n"
        "  │   │ Thought  │────▶│ LLM推理：我需要什么信息？  │     │\n"
        "  │   │  思考     │     │ 例：需要获取600519行情数据 │     │\n"
        "  │   └────┬─────┘     └──────────────────────────┘     │\n"
        "  │        │                                             │\n"
        "  │        ▼                                             │\n"
        "  │   ┌──────────┐     ┌──────────────────────────┐     │\n"
        "  │   │  Action  │────▶│ 调用工具执行动作           │     │\n"
        "  │   │  行动     │     │ 例：get_stock_quote(600519)│    │\n"
        "  │   └────┬─────┘     └──────────────────────────┘     │\n"
        "  │        │                                             │\n"
        "  │        ▼                                             │\n"
        "  │   ┌──────────┐     ┌──────────────────────────┐     │\n"
        "  │   │Observation│───▶│ 观察结果，决定下一步       │     │\n"
        "  │   │  观察     │     │ 例：当前价1856，RSI=65    │     │\n"
        "  │   └────┬─────┘     └──────────────────────────┘     │\n"
        "  │        │                                             │\n"
        "  │        │ 是否需要更多信息？                           │\n"
        "  │        ├──── 是 ────▶ 返回 Thought                   │\n"
        "  │        │                                             │\n"
        "  │        ▼ 否                                          │\n"
        "  │   ┌──────────┐                                       │\n"
        "  │   │  输出结论  │                                      │\n"
        "  │   └──────────┘                                       │\n"
        "  └──────────────────────────────────────────────────────┘"
    )
    add_code_block(doc, flow)


def add_proactive_flow(doc):
    """主动智能流程图"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4)
    add_run(p, "▎ 主动智能触发流程图", bold=True, size=13, color=DARK_BLUE)

    flow = (
        "  ┌─────────────────────────────────────────────────────────┐\n"
        "  │                   主动智能三通道                         │\n"
        "  │                                                         │\n"
        "  │  通道1：定时巡检          通道2：事件驱动               │\n"
        "  │  ┌──────────────┐        ┌──────────────┐              │\n"
        "  │  │ APScheduler  │        │  WebSocket   │              │\n"
        "  │  │ 08:30 晨报   │        │ 新闻推送     │              │\n"
        "  │  │ 17:00 收盘   │        │ 行情异动     │              │\n"
        "  │  │ 定时风险扫描  │        │ 持仓变化     │              │\n"
        "  │  └──────┬───────┘        └──────┬───────┘              │\n"
        "  │         │                       │                       │\n"
        "  │         └───────────┬───────────┘                       │\n"
        "  │                     │                                   │\n"
        "  │  通道3：阈值预警    │                                   │\n"
        "  │  ┌──────────────┐  │                                   │\n"
        "  │  │ 技术指标越界  │  │                                   │\n"
        "  │  │ RSI>80 超买   │──┘                                   │\n"
        "  │  │ 风险指标超限  │                                      │\n"
        "  │  └──────┬───────┘                                       │\n"
        "  │         │                                               │\n"
        "  │         ▼                                               │\n"
        "  │  ┌──────────────────────────────────────┐               │\n"
        "  │  │     Master Orchestrator               │               │\n"
        "  │  │     自动触发三阶段流水线分析           │               │\n"
        "  │  └──────────────────┬───────────────────┘               │\n"
        "  │                     │                                   │\n"
        "  │                     ▼                                   │\n"
        "  │  ┌──────────────────────────────────────┐               │\n"
        "  │  │  结果推送                             │               │\n"
        "  │  │  ├ WebSocket 实时推送至前端           │               │\n"
        "  │  │  ├ SSE 服务端推送事件流               │               │\n"
        "  │  │  └ 三层记忆系统存储结果               │               │\n"
        "  │  └──────────────────────────────────────┘               │\n"
        "  └─────────────────────────────────────────────────────────┘"
    )
    add_code_block(doc, flow)


def add_compliance_flow(doc):
    """合规审计流程图"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4)
    add_run(p, "▎ 合规审计流程图", bold=True, size=13, color=DARK_BLUE)

    flow = (
        "  ┌────────────────────────────────────────────────────────┐\n"
        "  │                  合规审计全链路                         │\n"
        "  │                                                        │\n"
        "  │  用户请求                                              │\n"
        "  │     │                                                  │\n"
        "  │     ▼                                                  │\n"
        "  │  ┌──────────────────────────────┐                      │\n"
        "  │  │ ComplianceAuditMiddleware     │                      │\n"
        "  │  │ 记录请求：时间/用户/路径/参数 │                      │\n"
        "  │  └──────────────┬───────────────┘                      │\n"
        "  │                 │                                      │\n"
        "  │                 ▼                                      │\n"
        "  │  ┌──────────────────────────────┐                      │\n"
        "  │  │ 风控合规智能体                │                      │\n"
        "  │  │ ├ 单股集中度 ≤ 10%           │                      │\n"
        "  │  │ ├ 行业集中度 ≤ 30%           │                      │\n"
        "  │  │ ├ 创业板限制 ≤ 20%           │                      │\n"
        "  │  │ └ ST股禁止 = 0%             │                      │\n"
        "  │  └──────────────┬───────────────┘                      │\n"
        "  │                 │                                      │\n"
        "  │          ┌──────┴──────┐                               │\n"
        "  │          │             │                               │\n"
        "  │          ▼             ▼                               │\n"
        "  │     ┌─────────┐  ┌─────────┐                          │\n"
        "  │     │ ✅ 合规  │  │ ❌ 违规  │                          │\n"
        "  │     │ 放行    │  │ 拦截    │                           │\n"
        "  │     └────┬────┘  └────┬────┘                          │\n"
        "  │          │            │                                │\n"
        "  │          ▼            ▼                                │\n"
        "  │  ┌──────────────────────────────┐                      │\n"
        "  │  │ 审计日志持久化存储             │                      │\n"
        "  │  │ 请求记录 + 合规检查 + 结果    │                      │\n"
        "  │  │ 满足银保监/证监会审计要求      │                      │\n"
        "  │  └──────────────────────────────┘                      │\n"
        "  └────────────────────────────────────────────────────────┘"
    )
    add_code_block(doc, flow)


def add_page_flow(doc):
    """页面导航流程图"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=4)
    add_run(p, "▎ 前端页面导航流程图", bold=True, size=13, color=DARK_BLUE)

    flow = (
        "                          ┌──────────┐\n"
        "                          │ Dashboard │\n"
        "                          │  工作台   │\n"
        "                          └─────┬────┘\n"
        "                 ┌──────┬───────┼───────┬──────┐\n"
        "                 │      │       │       │      │\n"
        "                 ▼      ▼       ▼       ▼      ▼\n"
        "           ┌──────┐┌──────┐┌──────┐┌──────┐┌──────┐\n"
        "           │Analyze││ Chat ││Agents││Reports││Alerts│\n"
        "           │智能分析││数字员工││智能体 ││报告中心││预警中心│\n"
        "           └──┬───┘└──┬───┘└──┬───┘└──┬───┘└──┬───┘\n"
        "              │       │       │       │       │\n"
        "              ▼       ▼       ▼       ▼       ▼\n"
        "         ┌─────────────────────────────────────────┐\n"
        "         │            共享组件层                     │\n"
        "         │  Sidebar │ AlertPanel │ ThinkingChain   │\n"
        "         │  SignalBadge │ WebSocket连接             │\n"
        "         └─────────────────────────────────────────┘\n"
        "                         │\n"
        "                         ▼\n"
        "         ┌─────────────────────────────────────────┐\n"
        "         │            后端API层                     │\n"
        "         │  /analyze │ /chat │ /alerts │ /reports  │\n"
        "         │  /agents/status │ /ws (WebSocket)       │\n"
        "         └─────────────────────────────────────────┘"
    )
    add_code_block(doc, flow)


def main():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    add_header(doc)
    add_footer(doc)
    create_cover_page(doc)

    # ========== 一、页面交互设计 ==========
    add_section_heading(doc, "一、页面交互设计（高保真原型）")
    add_body_paragraph(doc,
        "FinAgent Pro前端采用Vue3+TailwindCSS暗色主题设计，6个核心页面通过侧边栏导航切换。"
        "以下为各页面的高保真交互原型图。")

    add_dashboard_mockup(doc)
    doc.add_paragraph()
    add_analyze_mockup(doc)
    doc.add_paragraph()
    add_chat_mockup(doc)
    doc.add_paragraph()
    add_alerts_mockup(doc)

    # ========== 二、系统逻辑流程图 ==========
    doc.add_page_break()
    add_section_heading(doc, "二、系统逻辑流程图")
    add_body_paragraph(doc,
        "以下流程图展示了FinAgent Pro的核心系统逻辑，包括多智能体协同、ReAct推理、"
        "主动智能触发和合规审计四大核心流程。")

    add_system_flow(doc)
    doc.add_paragraph()
    add_react_flow(doc)
    doc.add_paragraph()
    add_proactive_flow(doc)
    doc.add_paragraph()
    add_compliance_flow(doc)
    doc.add_paragraph()
    add_page_flow(doc)

    # 保存
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"文档已生成: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()

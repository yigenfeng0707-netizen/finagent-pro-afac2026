#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成 AFAC2026 报名信息表 Word 文档"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── 常量 ──────────────────────────────────────────────
DARK_BLUE = RGBColor(0x1A, 0x3C, 0x6E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FONT_NAME = "微软雅黑"

OUTPUT_PATH = r"D:\AFAC2026金融智能创新大赛\FinAgent-Pro\docs\AFAC2026报名信息表.docx"

# ── 工具函数 ──────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, size=Pt(10.5), bold=False, color=None, font_name=FONT_NAME):
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color
    run.font.name = font_name
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def add_styled_paragraph(doc, text, size=Pt(10.5), bold=False, color=None,
                         alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6),
                         space_before=Pt(0), font_name=FONT_NAME):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, font_name=font_name)
    return p


def add_section_title(doc, text):
    p = add_styled_paragraph(doc, text, size=Pt(16), bold=True, color=DARK_BLUE,
                             space_before=Pt(18), space_after=Pt(10))
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="4" w:color="1A3C6E"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def create_info_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=Pt(11), bold=True, color=WHITE)
        set_cell_shading(cell, "1A3C6E")

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx > 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_run_font(run, size=Pt(10.5), bold=(c_idx == 0))
            if r_idx % 2 == 1:
                set_cell_shading(cell, "EDF2F9")

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w

    return table


def add_rich_text_paragraph(doc, text, size=Pt(10.5), first_line_indent=Cm(0.74)):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = first_line_indent
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(text)
    set_run_font(run, size=size)
    return p


# ── 主流程 ────────────────────────────────────────────

doc = Document()

# ── 全局默认字体 ──
style = doc.styles["Normal"]
style.font.name = FONT_NAME
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

# ── 页面设置 ──
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

# ── 页眉 ──
header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
hr = hp.add_run("AFAC2026\u91d1\u878d\u667a\u80fd\u521b\u65b0\u5927\u8d5b  \u62a5\u540d\u4fe1\u606f\u8868")
set_run_font(hr, size=Pt(9), color=DARK_BLUE)
pPr = hp._element.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="1A3C6E"/>'
    f'</w:pBdr>'
)
pPr.append(pBdr)

# ── 页脚（页码） ──
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
fp._element.append(fldChar1)
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
fp._element.append(instrText)
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
fp._element.append(fldChar2)

# ══════════════════════════════════════════════════════
# 封面页
# ══════════════════════════════════════════════════════

p_line = doc.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_line.paragraph_format.space_before = Pt(80)
run_line = p_line.add_run("\u2501" * 30)
set_run_font(run_line, size=Pt(14), color=DARK_BLUE)

add_styled_paragraph(doc, "AFAC2026", size=Pt(36), bold=True, color=DARK_BLUE,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(30), space_after=Pt(4))
add_styled_paragraph(doc, "\u91d1\u878d\u667a\u80fd\u521b\u65b0\u5927\u8d5b", size=Pt(28), bold=True, color=DARK_BLUE,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(20))

add_styled_paragraph(doc, "\u521d\u521b\u7ec4", size=Pt(22), bold=True, color=DARK_BLUE,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(30))

p_line2 = doc.add_paragraph()
p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_line2 = p_line2.add_run("\u2501" * 30)
set_run_font(run_line2, size=Pt(14), color=DARK_BLUE)

add_styled_paragraph(doc, "\u62a5 \u540d \u4fe1 \u606f \u8868", size=Pt(26), bold=True, color=DARK_BLUE,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(40), space_after=Pt(60))

add_styled_paragraph(doc, "\u56e2\u961f\u540d\u79f0\uff1a\u667a\u878d\u5148\u950b", size=Pt(14), color=DARK_BLUE,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))
add_styled_paragraph(doc, "\u9879\u76ee\u540d\u79f0\uff1aFinAgent Pro \u91d1\u878d\u81ea\u4e3b\u667a\u80fd\u4f53\u5e73\u53f0", size=Pt(14), color=DARK_BLUE,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))
add_styled_paragraph(doc, "\u8d1f\u8d23\u4eba\uff1a\u51af\u4ea6\u6839", size=Pt(14), color=DARK_BLUE,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(60))

add_styled_paragraph(doc, "2026\u5e746\u6708", size=Pt(12), color=RGBColor(0x66, 0x66, 0x66),
                     alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 一、基本信息
# ══════════════════════════════════════════════════════

add_section_title(doc, "\u4e00\u3001\u57fa\u672c\u4fe1\u606f")

basic_info = [
    ("\u8d5b\u9053", "\u521d\u521b\u7ec4"),
    ("\u65b9\u5411", "\u524d\u6cbf\u6280\u672f - \u81ea\u4e3b\u667a\u80fd\u4f53\uff08Agentic AI\uff09"),
    ("\u9879\u76ee\u540d\u79f0", "FinAgent Pro \u91d1\u878d\u81ea\u4e3b\u667a\u80fd\u4f53\u5e73\u53f0"),
    ("\u56e2\u961f\u540d\u79f0", "\u667a\u878d\u5148\u950b"),
    ("\u8d1f\u8d23\u4eba", "\u51af\u4ea6\u6839"),
    ("\u5b66\u5386", "\u672c\u79d1\uff08\u6d59\u6c5f\u5927\u5b66 \u8ba1\u7b97\u673a\u901a\u4fe1\u5de5\u7a0b\uff09"),
    ("\u56e2\u961f\u6a21\u5f0f", "OPC\u4e00\u4eba\u56e2\u961f"),
    ("GitHub", "https://github.com/yigenfeng0707-netizen/finagent-pro-afac2026"),
]

table1 = doc.add_table(rows=len(basic_info), cols=2)
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
table1.style = "Table Grid"

for i, (key, val) in enumerate(basic_info):
    cell_key = table1.rows[i].cells[0]
    cell_key.text = ""
    p = cell_key.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(key)
    set_run_font(run, size=Pt(11), bold=True, color=WHITE)
    set_cell_shading(cell_key, "1A3C6E")
    cell_key.width = Cm(4)

    cell_val = table1.rows[i].cells[1]
    cell_val.text = ""
    p = cell_val.paragraphs[0]
    run = p.add_run(val)
    set_run_font(run, size=Pt(10.5))
    if i % 2 == 1:
        set_cell_shading(cell_val, "EDF2F9")
    cell_val.width = Cm(12)

doc.add_paragraph()

# ══════════════════════════════════════════════════════
# 二、项目简介
# ══════════════════════════════════════════════════════

add_section_title(doc, "\u4e8c\u3001\u9879\u76ee\u7b80\u4ecb\uff08500\u5b57\u5185\uff09")

proj_intro = (
    'FinAgent Pro\u89e3\u51b3\u91d1\u878d\u884c\u4e1a\u4e09\u5927\u6838\u5fc3\u75db\u70b9\uff1a'
    '\u4e00\u662f\u5206\u6790\u5e08\u8fc7\u52b3\uff0c\u65e5\u5747\u4ec5\u8986\u76d63-5\u53ea\u6807\u7684\uff0c'
    '\u7814\u62a5\u64b0\u5199\u8017\u65f64-6\u5c0f\u65f6\uff1b'
    '\u4e8c\u662f\u98ce\u63a7\u6ede\u540e\uff0c80%\u98ce\u9669\u4e8b\u4ef6\u4e8b\u540e\u624d\u53d1\u73b0\uff1b'
    '\u4e09\u662f\u73b0\u6709AI\u5de5\u5177\u88ab\u52a8\u54cd\u5e94\uff0c'
    '\u9700\u4eba\u5de5\u63d0\u95ee\u624d\u5de5\u4f5c\uff0c\u65e0\u6cd5\u4e3b\u52a8\u53d1\u73b0\u673a\u4f1a\u548c\u98ce\u9669\u3002'
)
add_rich_text_paragraph(doc, proj_intro)

proj_target = (
    '\u76ee\u6807\u5ba2\u6237\u8986\u76d6\u56db\u7c7b\uff1a'
    '\u5238\u5546\u7814\u7a76\u6240\uff08\u7814\u62a5\u64b0\u5199\u3001\u6807\u7684\u8ddf\u8e2a\uff09\u3001'
    '\u57fa\u91d1\u516c\u53f8\uff08\u7ec4\u5408\u76d1\u63a7\u3001\u98ce\u9669\u9884\u8b66\uff09\u3001'
    '\u94f6\u884c\u7406\u8d22\uff08\u4ea7\u54c1\u5206\u6790\u3001\u5408\u89c4\u5ba1\u67e5\uff09\u3001'
    '\u4e2a\u4eba\u6295\u8d44\u8005\uff08\u667a\u80fd\u6295\u987e\u3001\u5e02\u573a\u89e3\u8bfb\uff09\u3002'
)
add_rich_text_paragraph(doc, proj_target)

proj_innovation = (
    '\u6838\u5fc3\u521b\u65b0\u70b9\uff1a'
    '\uff081\uff09ReAct\u63a8\u7406\u5f15\u64ce\u2014\u2014\u6bcf\u4e2a\u667a\u80fd\u4f53\u5177\u5907'
    '\u201c\u601d\u8003\u2192\u884c\u52a8\u2192\u89c2\u5bdf\u201d\u5faa\u73af\uff0c'
    '\u4e0d\u662f\u7b80\u5355\u5de5\u5177\u8c03\u7528\uff0c\u800c\u662f\u771f\u6b63\u81ea\u4e3b\u63a8\u7406\uff1b'
    '\uff082\uff09\u6301\u4e45\u8bb0\u5fc6\u7cfb\u7edf\u2014\u2014\u5de5\u4f5c\u8bb0\u5fc6+\u60c5\u8282\u8bb0\u5fc6+\u8bed\u4e49\u8bb0\u5fc6\u4e09\u5c42\u67b6\u6784\uff0c'
    '\u8ba9\u6570\u5b57\u5458\u5de5\u201c\u8bb0\u4f4f\u201d\u4e0a\u4e0b\u6587\u548c\u5386\u53f2\u7ecf\u9a8c\uff1b'
    '\uff083\uff09\u4e3b\u52a8\u667a\u80fd\u2014\u2014\u5b9a\u65f6\u5de1\u68c0+\u4e8b\u4ef6\u9a71\u52a8+\u9608\u503c\u9884\u8b66\uff0c'
    '\u63a8\u52a8\u91d1\u878d\u670d\u52a1\u4ece\u88ab\u52a8\u54cd\u5e94\u8fc8\u5411\u4e3b\u52a8\u667a\u80fd\uff1b'
    '\uff084\uff09\u5408\u89c4\u5185\u5d4c\u2014\u2014\u5185\u7f6e\u94f6\u4fdd\u76d1\u4f1a/\u8bc1\u76d1\u4f1a\u76d1\u7ba1\u89c4\u5219\u5f15\u64ce\uff0c'
    '\u6bcf\u6b65\u64cd\u4f5c\u53ef\u5ba1\u8ba1\u53ef\u8ffd\u6eaf\uff1b'
    '\uff085\uff09\u591a\u667a\u80fd\u4f53\u534f\u5546\u2014\u20146\u5927\u4e13\u4e1a\u667a\u80fd\u4f53'
    '\uff08\u5e02\u573a\u5206\u6790/\u65b0\u95fb\u8206\u60c5/\u98ce\u63a7\u5408\u89c4/'
    '\u6295\u8d44\u7b56\u7565/\u62a5\u544a\u751f\u6210/\u6267\u884c\u76d1\u63a7\uff09\u534f\u540c\u5de5\u4f5c\uff0c'
    '\u52a0\u6743\u534f\u5546\u751f\u6210\u6700\u7ec8\u5efa\u8bae\uff1b'
    '\uff086\uff09\u601d\u7ef4\u94fe\u53ef\u89c6\u5316\u2014\u2014'
    '\u5b9e\u65f6\u5c55\u793aAI\u63a8\u7406\u8fc7\u7a0b\uff0c\u8ba9\u51b3\u7b56\u201c\u770b\u5f97\u89c1\u201d\u3002'
    '\u4ee5\u56fd\u4ea7\u5927\u6a21\u578b\u4e3a\u57fa\u5ea7\uff0c\u6570\u636e\u4e0d\u51fa\u5883\uff0c\u5b89\u5168\u53ef\u4fe1\u3002'
)
add_rich_text_paragraph(doc, proj_innovation)

# ══════════════════════════════════════════════════════
# 三、团队简介
# ══════════════════════════════════════════════════════

add_section_title(doc, "\u4e09\u3001\u56e2\u961f\u7b80\u4ecb\uff08200\u5b57\u5185\uff09")

team_intro = (
    '\u667a\u878d\u5148\u950b\u7531\u6d59\u6c5f\u5927\u5b66\u8ba1\u7b97\u673a\u901a\u4fe1\u5de5\u7a0b\u4e13\u4e1a\u672c\u79d1\u751f\u51af\u4ea6\u6839\u72ec\u7acb\u521b\u7acb\uff0c'
    '\u662f\u4e00\u652f\u805a\u7126\u91d1\u878d\u81ea\u4e3b\u667a\u80fd\u4f53\uff08Agentic AI\uff09\u7684OPC\u4e00\u4eba\u56e2\u961f\u3002'
    '\u51af\u4ea6\u6839\u517c\u5177AI\u4e0e\u901a\u4fe1\u5de5\u7a0b\u590d\u5408\u80cc\u666f\uff0c'
    '\u72ec\u7acb\u5b8c\u6210FinAgent Pro\u91d1\u878d\u667a\u80fd\u4f53\u5e73\u53f0\u4ece\u67b6\u6784\u8bbe\u8ba1\u5230\u5168\u6808\u5f00\u53d1\u7684\u5168\u6d41\u7a0b'
    '\u2014\u2014\u6db5\u76d6ReAct\u63a8\u7406\u5f15\u64ce\u3001\u4e09\u5c42\u8bb0\u5fc6\u7cfb\u7edf\u30016\u5927\u4e13\u4e1a'
    '\u667a\u80fd\u4f53\u534f\u540c\u8c03\u5ea6\u3001\u5408\u89c4\u89c4\u5219\u5f15\u64ce\u53caVue3\u524d\u7aef\uff0c'
    '\u4ee5\u4e00\u4eba\u4e4b\u529b\u6784\u5efa\u4e86\u5b8c\u6574\u7684\u91d1\u878d\u6570\u5b57\u5458\u5de5\u7cfb\u7edf\u3002'
    '\u56e2\u961f\u575a\u6301\u201c\u6280\u672f\u521b\u65b0+\u5408\u89c4\u5185\u5d4c\u201d\u53cc\u8f6e\u9a71\u52a8\uff0c'
    '\u4ee5\u56fd\u4ea7\u5927\u6a21\u578b\u4e3a\u57fa\u5ea7\uff0c\u6253\u9020\u5b89\u5168\u53ef\u4fe1\u7684\u91d1\u878d\u667a\u80fd\u4f53\u5e73\u53f0\uff0c\u63a8\u52a8'
    '\u91d1\u878d\u670d\u52a1\u4ece\u88ab\u52a8\u54cd\u5e94\u8fc8\u5411\u4e3b\u52a8\u667a\u80fd\u3002'
    '\u4e00\u4e2a\u4eba\uff0c\u4e00\u652f\u961f\uff0c\u4ee5\u4e00\u5f53\u5341\u3002'
)
add_rich_text_paragraph(doc, team_intro)

# ══════════════════════════════════════════════════════
# 四、团队概况
# ══════════════════════════════════════════════════════

add_section_title(doc, "\u56db\u3001\u56e2\u961f\u6982\u51b5\uff08\u8be6\u7ec6\u7248\uff09")

team_detail = (
    '\u667a\u878d\u5148\u950b\u662f\u4e00\u652fOPC\u4e00\u4eba\u56e2\u961f\uff0c\u7531\u521b\u59cb\u4eba\u51af\u4ea6\u6839\u72ec\u7acb\u8fd0\u8425\u3002'
    '\u51af\u4ea6\u6839\u6bd5\u4e1a\u4e8e\u6d59\u6c5f\u5927\u5b66\u8ba1\u7b97\u673a\u901a\u4fe1\u5de5\u7a0b'
    '\u4e13\u4e1a\uff0c\u517c\u5177\u4eba\u5de5\u667a\u80fd\u4e0e\u901a\u4fe1\u5de5\u7a0b\u590d\u5408\u6280\u672f\u80cc\u666f\uff0c'
    '\u719f\u7ec3\u638c\u63e1\u5927\u8bed\u8a00\u6a21\u578b\u5e94\u7528\u5f00\u53d1\u3001\u591a\u667a\u80fd\u4f53\u7cfb\u7edf\u67b6\u6784\u8bbe\u8ba1\u3001'
    '\u5168\u6808\u5de5\u7a0b\u7b49\u6838\u5fc3\u80fd\u529b\u3002'
)
add_rich_text_paragraph(doc, team_detail)

team_detail2 = (
    '\u5728FinAgent Pro\u9879\u76ee\u4e2d\uff0c\u51af\u4ea6\u6839\u4ee5\u4e00\u4eba\u4e4b\u529b\u72ec\u7acb\u5b8c\u6210\u4e86\u4ece\u4ea7\u54c1\u5b9a\u4e49\u3001\u67b6\u6784\u8bbe\u8ba1\u5230\u524d\u540e\u7aef\u5168\u6808\u5f00\u53d1\u7684'
    '\u5168\u90e8\u5de5\u4f5c\u2014\u2014\u6db5\u76d6ReAct\u63a8\u7406\u5f15\u64ce\u3001\u4e09\u5c42\u8bb0\u5fc6\u7cfb\u7edf\u30016\u5927\u4e13\u4e1a\u667a\u80fd\u4f53\u534f\u540c\u8c03\u5ea6\u3001\u5408\u89c4\u89c4\u5219\u5f15\u64ce\u53caVue3\u524d\u7aef'
    '\u754c\u9762\uff0c\u4ee3\u7801\u91cf\u8d857000\u884c\u3002'
)
add_rich_text_paragraph(doc, team_detail2)

team_detail3 = (
    '\u56e2\u961f\u4ee5\u201c\u4e00\u4e2a\u4eba\uff0c\u4e00\u652f\u961f\uff0c\u4ee5\u4e00\u5f53\u5341\u201d\u4e3a\u4fe1\u6761\uff0c'
    '\u5145\u5206\u53d1\u6325AI\u539f\u751f\u5f00\u53d1\u4f18\u52bf\uff1a\u7528\u667a\u80fd\u4f53\u67b6\u6784\u500d\u589e\u4e2a\u4eba'
    '\u751f\u4ea7\u529b\uff0c\u7528\u81ea\u52a8\u5316\u5de5\u5177\u66ff\u4ee3\u91cd\u590d\u52b3\u52a8\uff0c\u7528\u6781\u7b80\u4e3b\u4e49\u805a\u7126\u6838\u5fc3\u4ef7\u503c\uff0c'
    '\u5b9e\u73b0\u4e86\u4e00\u4eba\u56e2\u961f\u7684\u9ad8\u6548\u8fd0\u8f6c\u3002\u672a\u6765'
    '\u56e2\u961f\u5c06\u6839\u636e\u4e1a\u52a1\u53d1\u5c55\u9700\u8981\uff0c\u4f18\u5148\u5728\u91d1\u878d\u5408\u89c4\u548c\u5546\u52a1\u62d3\u5c55\u65b9\u5411\u5f15\u5165\u5408\u4f19\u4eba\uff0c'
    '\u6301\u7eed\u4fdd\u6301\u6280\u672f\u56e2\u961f\u7684\u7cbe\u7b80\u9ad8\u6548\u3002'
)
add_rich_text_paragraph(doc, team_detail3)

# ══════════════════════════════════════════════════════
# 五、提交材料清单
# ══════════════════════════════════════════════════════

add_section_title(doc, "\u4e94\u3001\u63d0\u4ea4\u6750\u6599\u6e05\u5355")

materials = [
    ("1", "\u9879\u76ee\u7533\u62a5\u4e66", "FinAgentPro\u9879\u76ee\u7533\u62a5\u4e66.docx"),
    ("2", "\u5546\u4e1a\u8ba1\u5212\u4e66", "FinAgentPro\u5546\u4e1a\u8ba1\u5212\u4e66.docx"),
    ("3", "\u6280\u672f\u65b9\u6848", "FinAgentPro\u6280\u672f\u65b9\u6848.docx"),
    ("4", "\u8def\u6f14PPT", "FinAgentPro\u8def\u6f14PPT.pptx"),
    ("5", "\u7528\u6237\u624b\u518c", "FinAgentPro\u7528\u6237\u4f7f\u7528\u624b\u518c.docx"),
    ("6", "\u6e90\u4ee3\u7801", "GitHub\u4ed3\u5e93"),
]

create_info_table(doc, ["\u5e8f\u53f7", "\u6750\u6599\u540d\u79f0", "\u6587\u4ef6\u540d"], materials,
                  col_widths=[Cm(2), Cm(5), Cm(9)])

doc.add_paragraph()

# ══════════════════════════════════════════════════════
# 六、关键时间节点
# ══════════════════════════════════════════════════════

add_section_title(doc, "\u516d\u3001\u5173\u952e\u65f6\u95f4\u8282\u70b9")

timeline = [
    ("2026\u5e746\u6708", "\u62a5\u540d\u901a\u9053\u5f00\u542f"),
    ("2026\u5e747\u6708", "\u6210\u679c\u521d\u9009"),
    ("2026\u5e748\u6708\u4e2d\u65ec", "\u6210\u679c\u8def\u6f14"),
    ("2026\u5e749\u6708", "\u4e0a\u6d77Inclusion\u00b7\u5916\u6ee9\u5927\u4f1a\u9881\u5956"),
]

create_info_table(doc, ["\u65f6\u95f4", "\u4e8b\u9879"], timeline,
                  col_widths=[Cm(5), Cm(11)])

# ── 保存 ──
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"文档已生成: {OUTPUT_PATH}")
